import asyncio
import random
import string
import logging
from typing import Dict, Any
from pathlib import Path
from datetime import datetime
import io

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from modern_bot.config import TEMPLATE_PATH, DOCS_DIR
from modern_bot.utils.files import sanitize_filename
from modern_bot.database.db import load_user_data

logger = logging.getLogger(__name__)

def replace_placeholders_in_document(doc: Document, placeholders: Dict[str, str]) -> None:
    """
    Replaces placeholders inside paragraphs and tables, preserving original formatting.
    """
    def _replace_in_paragraph(paragraph):
        if not placeholders:
            return
        
        # 1. Fast path: replace inside individual runs first to preserve style splits
        for run in paragraph.runs:
            for key, value in placeholders.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)
                    
        # 2. Slow path: if placeholders are split across multiple runs by MS Word,
        # merge runs, perform replacement, and restore font attributes to the first run.
        full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        has_unreplaced = any(key in full_text for key in placeholders)
        
        if has_unreplaced:
            updated_text = full_text
            for key, value in placeholders.items():
                updated_text = updated_text.replace(key, value)
                
            if paragraph.runs:
                # Save first run formatting
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                bold = first_run.font.bold
                italic = first_run.font.italic
                color_rgb = first_run.font.color.rgb if first_run.font.color else None
                
                # Clear all other runs
                p_element = paragraph._p
                for r in list(paragraph.runs[1:]):
                    p_element.remove(r._r)
                    
                first_run.text = updated_text
                first_run.font.name = font_name
                first_run.font.size = font_size
                first_run.font.bold = bold
                first_run.font.italic = italic
                if color_rgb:
                    first_run.font.color.rgb = color_rgb
            else:
                paragraph.text = updated_text

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph)

def get_city_by_region(region: str) -> str:
    """
    Resolves Russian region name to its main city (e.g. Челябинская область -> Челябинск).
    """
    region_clean = region.strip()
    mapping = {
        "Челябинская область": "Челябинск",
        "Свердловская область": "Екатеринбург",
        "Башкирия": "Уфа",
        "ХМАО-Югра": "Сургут",
        "Ростовская область": "Ростов-на-Дону",
        "Краснодарский край": "Краснодар",
        "Нижний Новгород": "Нижний Новгород",
        "Санкт-Петербург": "Санкт-Петербург",
        "Екатеринбург": "Екатеринбург",
        "Тюмень": "Тюмень",
        "Челябинск": "Челябинск",
        "Магнитогорск": "Магнитогорск",
        "Курган": "Курган"
    }
    for k, v in mapping.items():
        if k.lower() in region_clean.lower() or region_clean.lower() in k.lower():
            return v
            
    city = region_clean.replace(" область", "").replace(" край", "")
    if city.endswith("ая"): # e.g. Челябинская -> Челябинск
        city = city[:-2] + "ск"
    return city

def insert_elegant_header(doc: Document, placeholders: Dict[str, str]) -> None:
    """
    Creates a premium formatted header matching the official annex standard
    at the top of the document, if the template does not already have placeholders.
    """
    # Check if template already has placeholders in XML
    xml_str = doc._element.xml
    if any(key in xml_str for key in placeholders):
        logger.info("Template already contains placeholders, skipping dynamic header generation")
        return

    logger.info("Generating premium dynamic header matching official annex standard")
    
    # Remove first three legacy paragraphs if present (к договору..., empty, г. Челябинск)
    for _ in range(min(3, len(doc.paragraphs))):
        p_to_del = doc.paragraphs[0]
        p_element = p_to_del._p
        p_element.getparent().remove(p_element)

    p_ref = doc.paragraphs[0] if doc.paragraphs else None

    def add_para(text, bold=False, size_pt=10, alignment=None, italic=False, space_after_pt=0, space_before_pt=0):
        nonlocal p_ref
        if p_ref:
            p = p_ref.insert_paragraph_before()
        else:
            p = doc.add_paragraph()
            
        p.paragraph_format.space_after = Pt(space_after_pt)
        p.paragraph_format.space_before = Pt(space_before_pt)
        p.paragraph_format.line_spacing = 1.15
        
        if alignment is not None:
            p.alignment = alignment
            
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            run.font.italic = italic
        return p

    # Insert header paragraphs in correct top-to-bottom order.
    # Note: Since insert_paragraph_before() inserts a paragraph BEFORE p_ref, 
    # to insert paragraphs P1, P2, P3... in this exact sequence, we can simply 
    # call insert_paragraph_before() on p_ref sequentially! All of them will be 
    # inserted before p_ref, in the exact sequence they are called.
    
    # 1. Annex and Contract metadata
    app_text = (
        f"Приложение № 1\n"
        f"к Договору на оказание услуг по оценке № {placeholders.get('{department_number}', 'Не указано')}\n"
        f"от 01.04.2019 г."
    )
    add_para(app_text, bold=False, size_pt=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT, italic=True)

    # 2. Spacer paragraph
    add_para("", size_pt=10, space_after_pt=6)

    # 3. Document Title
    add_para(f"ЗАКЛЮЧЕНИЕ № {placeholders.get('{issue_number}', 'Не указано')}", bold=True, size_pt=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=2)

    # 4. Assessment Subject
    add_para("об оценке предметов старины (антиквариата)", bold=True, size_pt=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=2)

    # 5. Ticket Number
    add_para(f"по залоговому билету № {placeholders.get('{ticket_number}', 'Не указано')}", bold=True, size_pt=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=6)

    # 6. Location and Date (г. Регион, дата)
    try:
        date_obj = datetime.strptime(placeholders.get('{date}', ''), '%d.%m.%Y')
    except Exception:
        date_obj = datetime.now()
        
    months = ["января", "февраля", "марта", "апреля", "мая", "июня", "июля", "августа", "сентября", "октября", "ноября", "декабря"]
    date_str = f"«{date_obj.day:02d}» {months[date_obj.month - 1]} {date_obj.year} г."
    
    region_raw = placeholders.get('{region}', 'Не указано')
    region_clean = get_city_by_region(region_raw)
    
    p_meta = add_para(f"г. {region_clean}\t{date_str}", bold=False, size_pt=10, space_after_pt=18)
    tab_stops = p_meta.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)

    # 7. Spacer paragraph before start of text
    add_para("", size_pt=10, space_after_pt=6)

def set_cell_text_with_style(cell, text, font_name="Times New Roman", font_size_pt=9, bold=False, italic=False, alignment=None):
    """
    Safely populates a cell's text while preserving core font styling.
    """
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.text = ""
    # Explicitly clear any remaining empty runs to ensure clean XML structure
    p_element = p._p
    for r in list(p.runs):
        if not r.text:
            p_element.remove(r._r)
            
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    
    if alignment is not None:
        p.alignment = alignment
        
    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        run.font.bold = bold
        run.font.italic = italic

def populate_table_with_data(doc: Document, data: Dict[str, Any]) -> None:
    """
    Populates the main photo description table with data using premium Times New Roman styling.
    """
    if not doc.tables:
        logger.error("No tables found in document.")
        return
    table = doc.tables[0]
    photo_desc = data.get('photo_desc', [])
    logger.info(f"Populating table with {len(photo_desc)} photo entries")
    
    for i, item in enumerate(photo_desc, 1):
        try:
            new_row = table.add_row()
            row_cells = new_row.cells
            if len(row_cells) < 8:
                logger.error("Table structure mismatch (less than 8 columns).")
                continue

            photo_path = Path(item.get('photo', ""))
            logger.info(f"Processing item {i}: photo_path={photo_path}, exists={photo_path.is_file()}")
            
            # Col 0: Index
            set_cell_text_with_style(row_cells[0], str(i), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Col 1: Description
            description = item.get('description') or 'Нет описания'
            set_cell_text_with_style(row_cells[1], description, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            
            # Col 2: Photo
            if photo_path.is_file():
                logger.info(f"Adding photo to document: {photo_path}")
                p = row_cells[2].paragraphs[0] if row_cells[2].paragraphs else row_cells[2].add_paragraph()
                p.text = ""
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.add_run().add_picture(str(photo_path), width=Inches(1.0))
                logger.info("Photo added successfully")
            else:
                logger.warning(f"Photo file not found: {photo_path}")
                set_cell_text_with_style(row_cells[2], 'Фото отсутствует', italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # Col 3 & 4: Characteristics (Material / Weight Size)
            set_cell_text_with_style(row_cells[3], '-', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text_with_style(row_cells[4], '-', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Col 5 & 6: Evaluation
            evaluation_value = item.get('evaluation') or 'Нет данных'
            set_cell_text_with_style(row_cells[5], evaluation_value, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text_with_style(row_cells[6], evaluation_value, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Col 7: Antique value
            set_cell_text_with_style(row_cells[7], 'да', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Col 8: Note (if present)
            set_cell_text_with_style(row_cells[8], '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
        except Exception as e:
            logger.error(f"Error populating table row {i}: {e}", exc_info=True)

# Dynamic reloadable template caching based on mtime
_template_cache = None
_template_mtime = 0.0

def _get_template_stream() -> io.BytesIO:
    global _template_cache, _template_mtime
    try:
        current_mtime = TEMPLATE_PATH.stat().st_mtime
    except Exception:
        current_mtime = 0.0
        
    if _template_cache is None or current_mtime != _template_mtime:
        logger.info(f"Loading/Reloading template '{TEMPLATE_PATH}' from disk (mtime={current_mtime})")
        with open(TEMPLATE_PATH, "rb") as f:
            _template_cache = f.read()
        _template_mtime = current_mtime
    return io.BytesIO(_template_cache)

async def create_document(user_id: int, user_name: str, db_data_override: Dict[str, Any] = None) -> Path:
    """
    Generates the DOCX document based on user inputs and the premium template.
    """
    try:
        if db_data_override:
            data = db_data_override
        else:
            data = await load_user_data(user_id)
            
        if not data:
            raise ValueError("No data found for user")
    except Exception as e:
        raise RuntimeError(f"Error loading user data: {e}") from e

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template '{TEMPLATE_PATH}' not found.")

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    selected_date = data.get('date') or datetime.now().strftime('%d.%m.%Y')
    timestamp = datetime.now().strftime('%H-%M-%S')
    placeholders = {
        '{date}': selected_date,
        '{issue_number}': data.get('issue_number', 'Не указано'),
        '{department_number}': data.get('department_number', 'Не указано'),
        '{region}': data.get('region', 'Не указано'),
        '{ticket_number}': data.get('ticket_number', 'Не указано'),
        '{username}': user_name
    }
    
    base_filename = (f"{placeholders['{department_number}']}, Заключение антиквариат № "
                     f"{placeholders['{issue_number}']} (билет {placeholders['{ticket_number}']}), "
                     f"{placeholders['{region}']}, от {selected_date} {timestamp}.docx")
    
    safe_filename_str = sanitize_filename(base_filename)
    if not safe_filename_str:
        safe_filename_str = f"Заключение_{timestamp}.docx"
    filepath = DOCS_DIR / safe_filename_str

    suffix = Path(safe_filename_str).suffix or ".docx"
    stem = Path(safe_filename_str).stem or "Заключение"
    while filepath.exists():
        unique_suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=4))
        candidate_name = sanitize_filename(f"{stem}_{unique_suffix}{suffix}")
        if not candidate_name:
            candidate_name = f"Заключение_{timestamp}_{unique_suffix}.docx"
        filepath = DOCS_DIR / candidate_name
        safe_filename_str = candidate_name

    def _build_document():
        try:
            # Load template stream (handles dynamic reload on disk changes)
            doc = Document(_get_template_stream())
            
            # 1. Insert elegant formatted header standard
            insert_elegant_header(doc, placeholders)
            
            # 2. Perform placeholder replacing with style-preservation
            replace_placeholders_in_document(doc, placeholders)
            
            # 3. Populate data table elegantly
            populate_table_with_data(doc, data)
            
            # 4. Save result
            doc.save(filepath)
        except Exception as doc_error:
            logger.error(f"Failed to build document {filepath}: {doc_error}", exc_info=True)
            raise

    try:
        await asyncio.to_thread(_build_document)
    except Exception as exc:
        raise RuntimeError("Error generating document.") from exc
    logger.info(f"Document saved: {filepath}")
    return filepath
