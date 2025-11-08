import os
import random
import string
import logging
import re
import json
import shutil
import zipfile
import math
import io
import copy
import statistics
from pathlib import Path
from dataclasses import dataclass, field
from enum import Enum, auto
from sqlite3 import IntegrityError
from typing import Dict, Any, List, Optional, Set, Tuple, Union, Sequence
from PIL import Image, ImageOps
from telegram import (
    Update,
    ReplyKeyboardMarkup,
    ReplyKeyboardRemove,
    KeyboardButton,
    InputMediaPhoto,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    BotCommand,
    BotCommandScopeAllPrivateChats,
    BotCommandScopeChat,
    WebAppInfo,
)
from telegram.constants import ChatAction
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    CallbackContext,
    ConversationHandler,
    CallbackQueryHandler
)
from telegram.error import RetryAfter, TimedOut, NetworkError, TelegramError
from datetime import datetime, timedelta
from calendar import monthrange
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time
import asyncio
import nest_asyncio
import aiosqlite
from openpyxl import Workbook, load_workbook
import matplotlib
from gettext import gettext as _

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.dates import DateFormatter
from logging.handlers import RotatingFileHandler

# Применяем nest_asyncio для возможности вложенного запуска цикла событий
nest_asyncio.apply()

# --- НАСТРОЙКИ БОТА ---
# При наличии переменной окружения она использует её, иначе берёт токен из кода.
BOT_TOKEN_ENV_VAR = "BOT_TOKEN"
BOT_TOKEN = "7514668293:AAHcxAooMsd1oDeoHOWkwbnTUD0BPUWePOY"
MAIN_GROUP_CHAT_ID: int = -1002381542769
DEFAULT_ADMIN_IDS: List[int] = [2064900]
WEBAPP_URL_ENV_VAR = "BOT_WEBAPP_URL"
DEFAULT_WEBAPP_URL = "https://146b0d031e1f.ngrok-free.app/?ngrok-skip-browser-warning=1"

REGION_TOPICS: Dict[str, int] = {
    "Санкт-Петербург": 11, "Свердловская область": 8, "Челябинская область": 6,
    "Екатеринбург": 4, "Башкирия": 12, "Тюмень": 13, "ХМАО-Югра": 15,
    "Нижний Новгород": 9, "Ростовская область": 17, "Челябинск": 2,
    "Магнитогорск": 7, "Курган": 16, "Краснодарский край": 14,
}

# Пути к файлам и папкам
TEMPLATE_PATH = Path("template.docx")
TEMP_PHOTOS_DIR = Path("photos")
DOCS_DIR = Path("documents")
ARCHIVE_DIR = Path("documents_archive")
ARCHIVE_INDEX_FILE = ARCHIVE_DIR / "index.json"
ADMIN_FILE = Path("config") / "admins.json"
DATABASE_FILE = Path("user_data.db")
EXCEL_FILE = Path("conclusions.xlsx")
LOG_DIR = Path("logs")
LOG_FILE = LOG_DIR / "bot.log"
BACKUP_ROOT = Path("backups")
MAX_PHOTOS: int = 30
MAX_PHOTO_SIZE_MB: int = 5
MIN_TICKET_DIGITS: int = 11
MAX_TICKET_DIGITS: int = 11
PROGRESS_STEPS: Dict[str, int] = {
    "department": 1,
    "issue": 2,
    "ticket": 3,
    "date": 4,
    "region": 5,
    "photo": 6,
    "description": 7,
    "evaluation": 8,
    "summary": 9,
    "mode": 10,
}
# вычисляем верхнюю границу для отображения прогресса
TOTAL_STEPS: int = max(PROGRESS_STEPS.values())
PREVIEW_MAX_ITEMS: int = 2
NETWORK_RECOVERY_INTERVAL: float = 45.0
MAX_PENDING_RESENDS: int = 20
MENU_BUTTON_LABEL = "/menu 📋"
PHOTO_REQUIREMENTS_MESSAGE = (
    "Требования к фото:\n"
    "• Формат JPG/PNG\n"
    f"• Размер до {MAX_PHOTO_SIZE_MB} МБ\n"
    "• Минимальное разрешение 800×600"
)
EXCEL_HEADERS = ["Ticket Number", "Conclusion Number", "Department Number", "Date", "Region", "Item Number", "Description", "Evaluation"]
PROGRESS_BAR_SEGMENTS = 5
PROGRESS_SYMBOL_FILLED = "●"
PROGRESS_SYMBOL_EMPTY = "○"
VOID_CALLBACK_PREFIX = "void:"
CONFIRM_CALLBACK_PREFIX = "confirm:"
MODE_CALLBACK_PREFIX = "mode:"
ADD_PHOTO_PREFIX = "photo:"
REGION_CALLBACK_PREFIX = "region:"
REPORT_REGION_CALLBACK_PREFIX = "report_region:"
REPORT_ACTION_PREFIX = "report:"
NAVIGATION_CALLBACK_PREFIX = "nav:"
ACHIEVEMENTS_CALLBACK_PREFIX = "achv:"
ACHIEVEMENTS_DEFAULT_VIEW = "main"
ACHIEVEMENTS_PANEL_KEY = "achievements_panel"
EDIT_CALLBACK_PREFIX = "edit:"
DRAFT_CALLBACK_PREFIX = "draft:"
HISTORY_CALLBACK_PREFIX = "history:"
ANALYTICS_CALLBACK_PREFIX = "analytics:"
ADMIN_CALLBACK_PREFIX = "admin:"
BACK_NAV_CALLBACK_PREFIX = "backnav:"
PENDING_BACK_DECISION_KEY = "pending_back_decision"
LEVEL_TARGET_COUNT: int = 80
LEVEL_BASE_XP: int = 120
LEVEL_GROWTH_RATE: float = 1.08
LEVEL_STEP_BONUS: int = 35
LEVEL_EMOJIS: List[str] = ["🌱", "🌿", "🌳", "🌼", "🌟", "🚀", "🛡", "🏆", "💎", "👑"]
LEVEL_TITLES: List[str] = [
    "Новичок",
    "Исследователь",
    "Мастер",
    "Эксперт",
    "Наставник",
    "Стратег",
    "Визионер",
    "Легенда",
    "Архитектор",
    "Куратор",
]


@dataclass
class ConclusionData:
    department_number: str = ""
    issue_number: str = ""
    ticket_number: str = ""
    date: str = ""
    region: str = ""
    photo_desc: List[Dict[str, Any]] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "department_number": self.department_number,
            "issue_number": self.issue_number,
            "ticket_number": self.ticket_number,
            "date": self.date,
            "region": self.region,
            "photo_desc": list(self.photo_desc),
        }

@dataclass
class ValidationResult:
    ok: bool
    message: Optional[str] = None
    hint: Optional[str] = None


def _compose_level_title(level: int) -> str:
    emoji = LEVEL_EMOJIS[(level - 1) % len(LEVEL_EMOJIS)]
    descriptor = LEVEL_TITLES[((level - 1) // len(LEVEL_EMOJIS)) % len(LEVEL_TITLES)]
    return f"{emoji} {descriptor}"


def generate_level_catalog(
    total_levels: int,
    base_increment: int = LEVEL_BASE_XP,
    growth: float = LEVEL_GROWTH_RATE,
    bonus_step: int = LEVEL_STEP_BONUS,
) -> List[Dict[str, Any]]:
    catalog: List[Dict[str, Any]] = []
    cumulative = 0
    for level in range(1, total_levels + 1):
        increment = int(round(base_increment * (growth ** (level - 1)) + bonus_step * math.log(level + 1)))
        increment = max(40, increment)
        cumulative += increment
        title = _compose_level_title(level)
        code = f"level_{level}"
        catalog.append(
            {
                "level": level,
                "code": code,
                "xp_required": cumulative,
                "title": f"{title} • уровень {level}",
                "label": title,
            }
        )
    return catalog


LEVEL_CATALOG: List[Dict[str, Any]] = generate_level_catalog(LEVEL_TARGET_COUNT)
# Основные клавиатуры
USER_MAIN_MENU_ROWS: List[List[str]] = [
    ["/start 🚀"],
    ["/help 📚"],
    ["/leaders 🏆"],
    ["/achievements 🏅"],
    ["/void_ticket ♻️"],
    ["/cancel ❌"],
]
ADMIN_MAIN_MENU_ROWS: List[List[str]] = USER_MAIN_MENU_ROWS + [
    ["/reports 📦"],
    ["/history 🕑"],
    ["/download_month ⚙️"],
    ["/stats_period 📈"],
    ["/stats 📊"],
    ["/add_admin 👥"],
    ["/help_admin 🔧"],
]
# Предопределённый порядок регионов для инлайн-выбора
REGION_CHOICES: List[str] = list(REGION_TOPICS.keys())
# Настройки ачивок и рейтингов
ACHIEVEMENT_MEDIA_DIR = Path("media") / "achievements"

ACHIEVEMENT_TIERS: Dict[str, List[Dict[str, Any]]] = {}


def _build_achievement_tiers() -> Dict[str, List[Dict[str, Any]]]:
    total_tiers: List[Dict[str, Any]] = [
        {
            "code": "total_1",
            "threshold": 1,
            "title": "Старт карьеры",
            "personal_template": "Первое окончательное заключение оформлено.",
            "media": ACHIEVEMENT_MEDIA_DIR / "total_1.png",
        },
        {
            "code": "total_3",
            "threshold": 3,
            "title": "Лёгкий разгон",
            "personal_template": "{current_total} окончательных заключений — вы быстро набираете обороты.",
        },
        {
            "code": "total_5",
            "threshold": 5,
            "title": "Уверенный шаг",
            "personal_template": "Вы оформили {current_total} окончательных заключений. Хороший темп.",
            "media": ACHIEVEMENT_MEDIA_DIR / "total_5.png",
        },
        {
            "code": "total_10",
            "threshold": 10,
            "title": "Топ 10",
            "personal_template": "Уже {current_total} окончательных заключений. Команда может на вас рассчитывать.",
            "media": ACHIEVEMENT_MEDIA_DIR / "total_10.png",
        },
        {
            "code": "total_15",
            "threshold": 15,
            "title": "Командный боец",
            "personal_template": "{current_total} окончательных заключений. Команда чувствует вашу поддержку.",
        },
        {
            "code": "total_20",
            "threshold": 20,
            "title": "Темповик",
            "personal_template": "Продолжайте держать темп: уже {current_total} окончательных заключений.",
        },
        {
            "code": "total_25",
            "threshold": 25,
            "title": "Чемпион сервиса",
            "personal_template": "{current_total} окончательных заключений. Вы в числе лидеров.",
            "media": ACHIEVEMENT_MEDIA_DIR / "total_25.png",
        },
        {
            "code": "total_30",
            "threshold": 30,
            "title": "Опорная точка",
            "personal_template": "{current_total} окончательных заключений. Ваш опыт становится референсом.",
        },
        {
            "code": "total_40",
            "threshold": 40,
            "title": "Стабильный результат",
            "personal_template": "За плечами {current_total} заключений. Это уровень эксперта.",
        },
        {
            "code": "total_50",
            "threshold": 50,
            "title": "Эксперт",
            "personal_template": "{current_total} окончательных заключений — серьёзный опыт.",
            "media": ACHIEVEMENT_MEDIA_DIR / "total_50.png",
        },
        {
            "code": "total_75",
            "threshold": 75,
            "title": "Топ-участник",
            "personal_template": "Целых {current_total} окончательных заключений. Вы входите в ядро команды.",
        },
        {
            "code": "total_100",
            "threshold": 100,
            "title": "Легенда сервиса",
            "personal_template": "{current_total} окончательных заключений. Вы задаёте стандарты.",
            "media": ACHIEVEMENT_MEDIA_DIR / "total_100.png",
        },
        {
            "code": "total_150",
            "threshold": 150,
            "title": "Гранд-мастер",
            "personal_template": "На счету {current_total} окончательных заключений. Уровень истинного наставника.",
        },
        {
            "code": "total_200",
            "threshold": 200,
            "title": "Абсолютный лидер",
            "personal_template": "{current_total} окончательных заключений. Вы формируете историю сервиса.",
        },
    ]

    items_tiers: List[Dict[str, Any]] = [
        {
            "code": "items_10",
            "threshold": 10,
            "title": "Коллекционер",
            "personal_template": "В оформленных заключениях уже {current_items_total} предметов. Отличное начало.",
        },
        {
            "code": "items_25",
            "threshold": 25,
            "title": "Ассортимент",
            "personal_template": "{current_items_total} предметов обработано. Вы расширяете каталог.",
        },
        {
            "code": "items_50",
            "threshold": 50,
            "title": "Хранитель фонда",
            "personal_template": "Пять десятков предметов ({current_items_total}) прошли через ваши руки.",
        },
        {
            "code": "items_100",
            "threshold": 100,
            "title": "Смотритель коллекции",
            "personal_template": "Уже {current_items_total} предметов оформлено. Это уровень музейного куратора.",
        },
        {
            "code": "items_200",
            "threshold": 200,
            "title": "Арсенал",
            "personal_template": "{current_items_total} предметов. Вы уверенно управляете большим запасом.",
        },
        {
            "code": "items_400",
            "threshold": 400,
            "title": "Хранитель легенд",
            "personal_template": "Коллекция насчитывает {current_items_total} предметов. Феноменальный охват.",
        },
    ]

    value_tiers: List[Dict[str, Any]] = [
        {
            "code": "value_50k",
            "threshold": 50_000,
            "title": "50 000 руб. оценок",
            "personal_template": "Суммарная стоимость достигла {current_value_total_fmt} руб. Отличная работа.",
        },
        {
            "code": "value_100k",
            "threshold": 100_000,
            "title": "Шестизначный вклад",
            "personal_template": "Команда доверяет {current_value_total_fmt} руб. оценок вашему профессионализму.",
        },
        {
            "code": "value_250k",
            "threshold": 250_000,
            "title": "Четверть миллиона",
            "personal_template": "Совокупная оценка достигла {current_value_total_fmt} руб. Вы держите крупные обороты.",
        },
        {
            "code": "value_500k",
            "threshold": 500_000,
            "title": "Полмиллиона",
            "personal_template": "Полмиллиона рублей ({current_value_total_fmt}) оценено при вашем участии.",
        },
        {
            "code": "value_1m",
            "threshold": 1_000_000,
            "title": "Миллионник",
            "personal_template": "Суммарная стоимость оформленных предметов превысила {current_value_total_fmt} руб.",
        },
    ]

    monthly_tiers: List[Dict[str, Any]] = [
        {
            "code": "monthly_5",
            "threshold": 5,
            "title": "Бронзовый темп",
            "personal_template": "{month_label}: оформлено {current_month} заключений.",
            "media": ACHIEVEMENT_MEDIA_DIR / "monthly_5.png",
        },
        {
            "code": "monthly_10",
            "threshold": 10,
            "title": "Серебряный темп",
            "personal_template": "{month_label}: уже {current_month} заключений. Держите скорость.",
            "media": ACHIEVEMENT_MEDIA_DIR / "monthly_10.png",
        },
        {
            "code": "monthly_20",
            "threshold": 20,
            "title": "Золотой темп",
            "personal_template": "{month_label}: {current_month} заключений — отличная динамика.",
            "media": ACHIEVEMENT_MEDIA_DIR / "monthly_20.png",
        },
        {
            "code": "monthly_40",
            "threshold": 40,
            "title": "Месячный рекорд",
            "personal_template": "{month_label}: {current_month} заключений. Это уровень рекорда.",
            "media": ACHIEVEMENT_MEDIA_DIR / "monthly_40.png",
        },
    ]

    daily_tiers: List[Dict[str, Any]] = [
        {
            "code": "daily_3",
            "threshold": 3,
            "title": "Разгон",
            "personal_template": "{day_label}: {current_day} заключения за день. Отличный темп.",
            "media": ACHIEVEMENT_MEDIA_DIR / "daily_3.png",
        },
        {
            "code": "daily_5",
            "threshold": 5,
            "title": "Супердень",
            "personal_template": "{day_label}: {current_day} заключений за день. Вы в фокусе задач.",
            "media": ACHIEVEMENT_MEDIA_DIR / "daily_5.png",
        },
        {
            "code": "daily_8",
            "threshold": 8,
            "title": "Дневной максимум",
            "personal_template": "{day_label}: {current_day} заключений. Максимальная продуктивность.",
            "media": ACHIEVEMENT_MEDIA_DIR / "daily_8.png",
        },
    ]

    streak_tiers: List[Dict[str, Any]] = [
        {
            "code": "streak_3",
            "threshold": 3,
            "title": "Серия 3",
            "personal_template": "{current_streak} дней подряд с оформленными заключениями.",
            "media": ACHIEVEMENT_MEDIA_DIR / "streak_3.png",
        },
        {
            "code": "streak_7",
            "threshold": 7,
            "title": "Серия 7",
            "personal_template": "{current_streak} дней подряд держите темп.",
            "media": ACHIEVEMENT_MEDIA_DIR / "streak_7.png",
        },
        {
            "code": "streak_14",
            "threshold": 14,
            "title": "Серия 14",
            "personal_template": "{current_streak} дней без пауз. Это высший уровень стабильности.",
            "media": ACHIEVEMENT_MEDIA_DIR / "streak_14.png",
        },
    ]

    level_tiers: List[Dict[str, Any]] = []
    for entry in LEVEL_CATALOG:
        level_tiers.append(
            {
                "code": entry["code"],
                "threshold": entry["xp_required"],
                "title": entry["title"],
                "personal_template": "Открыт уровень {current_level} — {current_level_label}.",
                "label": entry["label"],
                "level": entry["level"],
            }
        )

    return {
        "total": total_tiers,
        "items_total": items_tiers,
        "value_total": value_tiers,
        "monthly": monthly_tiers,
        "daily": daily_tiers,
        "streak": streak_tiers,
        "level": level_tiers,
    }


ACHIEVEMENT_TIERS = _build_achievement_tiers()
ACHIEVEMENT_LOOKUP: Dict[str, Dict[str, Any]] = {
    tier["code"]: {"metric": metric, **tier}
    for metric, tiers in ACHIEVEMENT_TIERS.items()
    for tier in tiers
}
LEADERBOARD_SIZE: int = 5
# --- КОНЕЦ НАСТРОЕК ---


def load_bot_token() -> str:
    """Возвращает токен бота (с приоритетом переменной окружения)."""
    token = os.getenv(BOT_TOKEN_ENV_VAR, "").strip() or BOT_TOKEN
    if not token:
        raise RuntimeError(
            f"Не указан токен бота. Установите переменную окружения {BOT_TOKEN_ENV_VAR}."
        )
    return token


def load_webapp_url() -> str:
    """Возвращает URL мини-приложения Telegram."""
    url = os.getenv(WEBAPP_URL_ENV_VAR, "").strip()
    if url:
        return url
    return DEFAULT_WEBAPP_URL

# Этапы диалога


class DialogState(Enum):
    DEPARTMENT = auto()
    ISSUE_NUMBER = auto()
    TICKET_NUMBER = auto()
    DATE = auto()
    REGION = auto()
    PHOTO = auto()
    DESCRIPTION = auto()
    EVALUATION = auto()
    MORE_PHOTO = auto()
    CONFIRMATION = auto()
    TESTING = auto()


class ReportState(Enum):
    ACTION = auto()
    MONTH_INPUT = auto()
    MONTH_REGION = auto()
    PERIOD_START = auto()
    PERIOD_END = auto()
    PERIOD_REGION = auto()

TEXTUAL_BACK_STATES: Set[DialogState] = {
    DialogState.DEPARTMENT,
    DialogState.ISSUE_NUMBER,
    DialogState.TICKET_NUMBER,
    DialogState.DATE,
    DialogState.REGION,
}

# Настройка логирования


def setup_logging(level: int = logging.INFO) -> None:
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    root_logger = logging.getLogger()
    root_logger.setLevel(level)
    for handler in list(root_logger.handlers):
        root_logger.removeHandler(handler)

    console_handler = logging.StreamHandler()
    console_handler.setFormatter(formatter)

    file_handler = RotatingFileHandler(LOG_FILE, maxBytes=2_000_000, backupCount=5, encoding="utf-8")
    file_handler.setFormatter(formatter)

    root_logger.addHandler(console_handler)
    root_logger.addHandler(file_handler)


setup_logging()
logger = logging.getLogger(__name__)

# Блокировки для безопасного доступа к ресурсам
db_lock = asyncio.Lock()
excel_lock = asyncio.Lock()
archive_lock = asyncio.Lock()
admin_ids: Set[int] = set()
network_recovery_lock = asyncio.Lock()
network_recovery_pending: Dict[int, Dict[str, Any]] = {}


async def mark_network_issue(chat_id: int, text: str, kwargs: Dict[str, Any], skip_notice: bool = False) -> None:
    async with network_recovery_lock:
        entry = network_recovery_pending.setdefault(
            chat_id,
            {"timestamp": time.time() - NETWORK_RECOVERY_INTERVAL, "messages": []}
        )
        messages: List[Tuple[str, Dict[str, Any]]] = entry.setdefault("messages", [])
        messages.append((text, kwargs))
        if len(messages) > MAX_PENDING_RESENDS:
            entry["messages"] = messages[-MAX_PENDING_RESENDS:]
        entry["timestamp"] = time.time() - NETWORK_RECOVERY_INTERVAL
        entry["skip_notice"] = skip_notice


async def process_network_recovery(bot, min_interval: float = NETWORK_RECOVERY_INTERVAL) -> None:
    async with network_recovery_lock:
        snapshot = {
            chat_id: {
                "timestamp": payload.get("timestamp", 0.0),
                "messages": list(payload.get("messages", [])),
                "skip_notice": payload.get("skip_notice", False),
            }
            for chat_id, payload in network_recovery_pending.items()
        }

    if not snapshot:
        return

    now = time.time()
    for chat_id, payload in snapshot.items():
        messages: List[Tuple[str, Dict[str, Any]]] = payload.get("messages", [])
        timestamp = payload.get("timestamp", 0.0)
        skip_notice = bool(payload.get("skip_notice", False))

        if not messages:
            async with network_recovery_lock:
                network_recovery_pending.pop(chat_id, None)
            continue

        if now - timestamp < min_interval:
            continue

        remaining: List[Tuple[str, Dict[str, Any]]] = []
        sent_count = 0
        failure = False

        for idx, (msg_text, msg_kwargs) in enumerate(messages):
            try:
                await bot.send_message(chat_id, msg_text, **msg_kwargs)
                logger.info(f"Повторно отправлено сохранённое сообщение в чат {chat_id}.")
                sent_count += 1
            except RetryAfter as retry_error:
                delay = getattr(retry_error, "retry_after", min_interval)
                remaining = messages[idx:]
                async with network_recovery_lock:
                    network_recovery_pending[chat_id] = {
                        "timestamp": now + delay,
                        "messages": remaining,
                        "skip_notice": skip_notice,
                    }
                failure = True
                break
            except (NetworkError, asyncio.TimeoutError) as net_error:
                logger.warning(f"Сеть по-прежнему недоступна для чата {chat_id}: {net_error}")
                remaining = messages[idx:]
                async with network_recovery_lock:
                    network_recovery_pending[chat_id] = {
                        "timestamp": now,
                        "messages": remaining,
                        "skip_notice": skip_notice,
                    }
                failure = True
                break
            except TelegramError as tg_error:
                logger.error(f"Не удалось доставить сохранённое сообщение в чат {chat_id}: {tg_error}")
                continue

        if failure:
            continue

        async with network_recovery_lock:
            network_recovery_pending.pop(chat_id, None)

        if sent_count and not skip_notice:
            recovery_text = (
                f"✅ Связь с ботом восстановлена. Доставлено {sent_count} сохранённых сообщений."
            )
            message = await safe_bot_send_message(
                bot,
                chat_id,
                recovery_text,
                skip_notice_on_retry=True,
            )
            if message:
                logger.info(f"Чат {chat_id}: доставлено {sent_count} сохранённых сообщений после восстановления сети.")


async def network_recovery_job(context: CallbackContext) -> None:
    await process_network_recovery(context.application.bot)


async def error_handler(update: object, context: CallbackContext) -> None:
    err = context.error
    if isinstance(err, NetworkError):
        logger.warning(f"Сетевая ошибка: {err}. Попробуем восстановить соединение.")
        await asyncio.sleep(5)
        job_queue = getattr(context.application, "job_queue", None)
        if job_queue and not job_queue.running:
            job_queue.start()
        await process_network_recovery(context.application.bot)
        return

    logger.error("Ошибка при обработке обновления:", exc_info=err)

# -------------------- Работа с базой данных (aiosqlite) --------------------
db: aiosqlite.Connection = None


def _is_db_ready() -> bool:
    if db is None:
        logger.error("База данных не инициализирована. Вызов init_db() ещё не выполнен.")
        return False
    return True


async def init_db() -> None:
    """Инициализирует базу данных и создает таблицу, если её нет."""
    global db
    if db is not None:
        return
    db = await aiosqlite.connect(DATABASE_FILE)
    await db.execute("PRAGMA journal_mode=WAL;")
    await db.execute("PRAGMA synchronous=NORMAL;")
    await db.execute('''CREATE TABLE IF NOT EXISTS user_data (
        user_id INTEGER PRIMARY KEY, department_number TEXT, issue_number TEXT,
        date TEXT, photo_desc TEXT, region TEXT, ticket_number TEXT
    )''')
    await db.execute('''CREATE TABLE IF NOT EXISTS user_meta (
        user_id INTEGER PRIMARY KEY,
        department_number TEXT,
        region TEXT,
        issue_sequence INTEGER,
        issue_date TEXT,
        recent_departments TEXT,
        recent_regions TEXT,
        recent_pairs TEXT
    )''')
    await db.execute('''CREATE TABLE IF NOT EXISTS completions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        username TEXT,
        completed_at TEXT NOT NULL,
        item_count INTEGER NOT NULL,
        total_evaluation REAL NOT NULL,
        region TEXT,
        ticket_number TEXT,
        issue_number TEXT,
        department_number TEXT,
        date TEXT,
        group_chat_id INTEGER,
        group_message_id INTEGER,
        thread_id INTEGER,
        archive_path TEXT,
        items_json TEXT
    )''')
    await db.execute('''CREATE TABLE IF NOT EXISTS achievement_log (
        user_id INTEGER NOT NULL,
        achievement_key TEXT NOT NULL,
        achieved_at TEXT NOT NULL,
        PRIMARY KEY (user_id, achievement_key)
    )''')
    await db.execute('''CREATE TABLE IF NOT EXISTS drafts (
        user_id INTEGER PRIMARY KEY,
        data TEXT NOT NULL,
        state TEXT,
        title TEXT,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL,
        metrics TEXT
    )''')
    await db.execute('''CREATE TABLE IF NOT EXISTS user_flags (
        user_id INTEGER PRIMARY KEY,
        is_blocked INTEGER DEFAULT 0,
        notes TEXT,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )''')
    await _ensure_table_columns("completions", {
        "ticket_number": "TEXT",
        "issue_number": "TEXT",
        "department_number": "TEXT",
        "date": "TEXT",
        "group_chat_id": "INTEGER",
        "group_message_id": "INTEGER",
        "thread_id": "INTEGER",
        "archive_path": "TEXT",
        "items_json": "TEXT",
        "xp_value": "INTEGER",
        "processing_time_seconds": "REAL",
        "step_metrics": "TEXT",
        "is_deleted": "INTEGER DEFAULT 0",
        "deleted_at": "TEXT",
        "deleted_by": "INTEGER",
        "deletion_note": "TEXT",
    })
    await db.execute('CREATE INDEX IF NOT EXISTS idx_completions_user_date ON completions(user_id, completed_at)')
    await db.execute('CREATE INDEX IF NOT EXISTS idx_completions_completed_at ON completions(completed_at)')
    await db.execute('CREATE INDEX IF NOT EXISTS idx_completions_ticket ON completions(ticket_number)')
    await db.execute('CREATE INDEX IF NOT EXISTS idx_completions_issue ON completions(issue_number)')
    await _ensure_table_columns("user_meta", {
        "recent_departments": "TEXT",
        "recent_regions": "TEXT",
        "recent_pairs": "TEXT",
        "issue_sequence": "INTEGER",
        "issue_date": "TEXT",
        "department_number": "TEXT",
        "region": "TEXT",
    })
    await db.commit()

async def close_db() -> None:
    """Корректно закрывает соединение с БД при остановке бота."""
    global db
    if db:
        await db.close()
        db = None
        logger.info("Соединение с базой данных закрыто.")

async def save_user_data_to_db(user_id: int, data: ConclusionData) -> None:
    """Сохраняет данные пользователя в БД."""
    if not _is_db_ready():
        return
    current_streak = 0
    previous_streak = 0

    async with db_lock:
        try:
            await db.execute(
                '''INSERT OR REPLACE INTO user_data (user_id, department_number, issue_number, date, region, ticket_number, photo_desc)
                   VALUES (?, ?, ?, ?, ?, ?, ?)''',
                (
                    user_id,
                    data.department_number,
                    data.issue_number,
                    data.date,
                    data.region,
                    data.ticket_number,
                    json.dumps(data.photo_desc),
                )
            )
            await db.commit()
        except Exception as e:
            logger.error(f"Ошибка БД при сохранении данных пользователя {user_id}: {e}")
            return

    await update_meta_defaults(user_id, data)

async def load_user_data_from_db(user_id: int) -> ConclusionData:
    """Загружает данные пользователя из БД."""
    if not _is_db_ready():
        data = ConclusionData()
    else:
        async with db_lock:
            row = None
            try:
                async with db.execute(
                    'SELECT department_number, issue_number, date, region, ticket_number, photo_desc FROM user_data WHERE user_id = ?',
                    (user_id,),
                ) as cursor:
                    row = await cursor.fetchone()
            except Exception as error:
                logger.error(f"Ошибка БД при загрузке данных пользователя {user_id}: {error}")

        if row:
            data = ConclusionData(
                department_number=row[0] or "",
                issue_number=row[1] or "",
                date=row[2] or "",
                region=row[3] or "",
                ticket_number=row[4] or "",
                photo_desc=json.loads(row[5] or '[]'),
            )
        else:
            data = ConclusionData()

    meta = await load_user_meta(user_id)
    if not data.department_number and meta.get("department_number"):
        data.department_number = meta["department_number"]
    if not data.region and meta.get("region"):
        data.region = meta["region"]
    if not data.date:
        data.date = _today_display()

    if not data.issue_number:
        next_issue = _compute_next_issue(meta, data.date)
        data.issue_number = str(next_issue)

    return data

async def delete_user_data_from_db(user_id: int) -> None:
    """Удаляет данные пользователя из БД."""
    if not _is_db_ready():
        return
    async with db_lock:
        try:
            await db.execute('DELETE FROM user_data WHERE user_id = ?', (user_id,))
            await db.commit()
        except Exception as e:
            logger.error(f"Ошибка БД при удалении данных пользователя {user_id}: {e}")


def _now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _today_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d")


def _today_display() -> str:
    return datetime.now().strftime("%d.%m.%Y")


def _display_to_iso(date_str: Optional[str]) -> Optional[str]:
    if not date_str:
        return None
    try:
        return datetime.strptime(date_str, "%d.%m.%Y").strftime("%Y-%m-%d")
    except ValueError:
        try:
            return datetime.strptime(date_str, "%Y-%m-%d").strftime("%Y-%m-%d")
        except ValueError:
            return None


def _iso_to_display(date_str: Optional[str]) -> Optional[str]:
    if not date_str:
        return None
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d.%m.%Y")
    except ValueError:
        return None


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _conclusion_from_payload(payload: Dict[str, Any]) -> ConclusionData:
    return ConclusionData(
        department_number=payload.get("department_number", "") if payload else "",
        issue_number=payload.get("issue_number", "") if payload else "",
        ticket_number=payload.get("ticket_number", "") if payload else "",
        date=payload.get("date", "") if payload else "",
        region=payload.get("region", "") if payload else "",
        photo_desc=list(payload.get("photo_desc") or []) if payload else [],
    )


async def load_user_meta(user_id: int) -> Dict[str, Any]:
    if not _is_db_ready():
        return {}
    async with db_lock:
        try:
            async with db.execute(
                "SELECT department_number, region, issue_sequence, issue_date, recent_departments, recent_regions, recent_pairs "
                "FROM user_meta WHERE user_id = ?",
                (user_id,),
            ) as cursor:
                row = await cursor.fetchone()
        except Exception as error:
            logger.error(f"Не удалось загрузить user_meta для {user_id}: {error}")
            return {}

    if not row:
        return {}

    def _loads(blob: Optional[str]) -> List[Any]:
        if not blob:
            return []
        try:
            return json.loads(blob)
        except json.JSONDecodeError:
            return []

    return {
        "department_number": row[0] or "",
        "region": row[1] or "",
        "issue_sequence": _safe_int(row[2]),
        "issue_date": row[3] or "",
        "recent_departments": _loads(row[4]),
        "recent_regions": _loads(row[5]),
        "recent_pairs": _loads(row[6]),
    }


async def upsert_user_meta(user_id: int, **fields: Any) -> None:
    if not _is_db_ready() or not fields:
        return

    allowed = {
        "department_number",
        "region",
        "issue_sequence",
        "issue_date",
        "recent_departments",
        "recent_regions",
        "recent_pairs",
    }
    payload: Dict[str, Any] = {}
    for key, value in fields.items():
        if key not in allowed or value is None:
            continue
        if key.startswith("recent_"):
            payload[key] = json.dumps(value, ensure_ascii=False)
        else:
            payload[key] = value

    if not payload:
        return

    columns = ", ".join(payload.keys())
    placeholders = ", ".join(["?"] * len(payload))
    update_clause = ", ".join(f"{key}=excluded.{key}" for key in payload.keys())
    values = list(payload.values())

    async with db_lock:
        try:
            await db.execute(
                f"INSERT INTO user_meta (user_id, {columns}) VALUES (?, {placeholders}) "
                f"ON CONFLICT(user_id) DO UPDATE SET {update_clause}",
                [user_id, *values],
            )
            await db.commit()
        except Exception as error:
            logger.error(f"Не удалось обновить user_meta для {user_id}: {error}")


def _merge_recent(entries: List[str], new_value: str, limit: int = 5) -> List[str]:
    if not new_value:
        return entries[:limit]
    combined = [new_value] + [item for item in entries if item and item != new_value]
    return combined[:limit]


def _merge_recent_pairs(pairs: List[Dict[str, Any]], department: str, region: str, date_iso: Optional[str], limit: int = 5) -> List[Dict[str, Any]]:
    if not (department or region):
        return pairs[:limit]
    normalized_date = date_iso or _today_iso()
    new_entry = {
        "department_number": department or "",
        "region": region or "",
        "date": normalized_date,
        "date_iso": normalized_date,
    }
    filtered = [pair for pair in pairs if not (
        (pair.get("department_number") == new_entry["department_number"]) and
        (pair.get("region") == new_entry["region"]) and
        ((pair.get("date_iso") or pair.get("date")) == new_entry["date"])
    )]
    return [new_entry] + filtered[: limit - 1]


async def update_meta_defaults(user_id: int, data: ConclusionData) -> None:
    if not isinstance(data, ConclusionData):
        return
    meta = await load_user_meta(user_id)
    department = data.department_number or meta.get("department_number") or ""
    region = data.region or meta.get("region") or ""
    recent_departments = _merge_recent(meta.get("recent_departments", []), data.department_number)
    recent_regions = _merge_recent(meta.get("recent_regions", []), data.region)
    recent_pairs = _merge_recent_pairs(meta.get("recent_pairs", []), data.department_number, data.region, _display_to_iso(data.date))

    await upsert_user_meta(
        user_id,
        department_number=department or None,
        region=region or None,
        recent_departments=recent_departments,
        recent_regions=recent_regions,
        recent_pairs=recent_pairs,
    )


async def register_completion_meta(user_id: int, data: ConclusionData) -> None:
    if not isinstance(data, ConclusionData):
        return
    meta = await load_user_meta(user_id)
    issue_number = _safe_int(data.issue_number, default=0)
    issue_date_iso = _display_to_iso(data.date) or _today_iso()
    issue_sequence = _safe_int(meta.get("issue_sequence"), default=0)
    current_date = meta.get("issue_date") or ""
    if issue_date_iso != current_date:
        issue_sequence = 0
    if issue_number <= 0:
        issue_number = issue_sequence

    new_sequence = max(issue_sequence, issue_number)
    recent_departments = _merge_recent(meta.get("recent_departments", []), data.department_number)
    recent_regions = _merge_recent(meta.get("recent_regions", []), data.region)
    recent_pairs = _merge_recent_pairs(meta.get("recent_pairs", []), data.department_number, data.region, issue_date_iso)

    await upsert_user_meta(
        user_id,
        department_number=(data.department_number or meta.get("department_number") or None),
        region=(data.region or meta.get("region") or None),
        issue_sequence=new_sequence,
        issue_date=issue_date_iso,
        recent_departments=recent_departments,
        recent_regions=recent_regions,
        recent_pairs=recent_pairs,
    )


def _compute_next_issue(meta: Dict[str, Any], target_date_display: Optional[str] = None) -> int:
    if not meta:
        return 1
    issue_sequence = _safe_int(meta.get("issue_sequence"), default=0)
    issue_date = meta.get("issue_date") or ""
    target_iso = _display_to_iso(target_date_display) or _today_iso()
    if issue_date == target_iso:
        return max(1, issue_sequence + 1)
    return 1


async def build_draft_meta(user_id: int, data: ConclusionData) -> Dict[str, Any]:
    meta = await load_user_meta(user_id)
    next_issue = _compute_next_issue(meta, data.date)
    today_iso = _today_iso()
    today_display = _today_display()

    recent_departments = meta.get("recent_departments", [])
    recent_regions = meta.get("recent_regions", [])
    raw_pairs = meta.get("recent_pairs", [])
    recent_pairs = []
    for pair in raw_pairs:
        if not isinstance(pair, dict):
            continue
        iso_date = pair.get("date")
        recent_pairs.append(
            {
                "department_number": pair.get("department_number", ""),
                "region": pair.get("region", ""),
                "date": _iso_to_display(iso_date) or iso_date,
                "date_iso": iso_date or "",
            }
        )

    if not data.issue_number:
        data.issue_number = str(next_issue)
    if not data.date:
        data.date = today_display

    return {
        "next_issue_number": next_issue,
        "issue_sequence": _safe_int(meta.get("issue_sequence"), default=0),
        "issue_date_iso": meta.get("issue_date") or "",
        "recent_departments": recent_departments,
        "recent_regions": recent_regions,
        "recent_pairs": recent_pairs,
        "last_department": meta.get("department_number") or "",
        "last_region": meta.get("region") or "",
        "today_iso": today_iso,
        "today_display": today_display,
    }


def validate_conclusion_data(data: ConclusionData) -> List[str]:
    errors: List[str] = []
    department = (data.department_number or "").strip()
    issue_number = (data.issue_number or "").strip()
    ticket_number = (data.ticket_number or "").strip()
    region = (data.region or "").strip()
    date_display = (data.date or "").strip()

    if not re.fullmatch(r"\d{3}", department):
        errors.append("Подразделение должно содержать ровно 3 цифры.")

    if not re.fullmatch(r"1\d{0,2}", issue_number):
        errors.append("Номер заключения начинается с 1 и содержит не более 3 цифр.")

    if not re.fullmatch(r"\d{11}", ticket_number):
        errors.append("Номер билета должен содержать 11 цифр.")

    if not region:
        errors.append("Укажите регион из справочника.")

    date_iso = _display_to_iso(date_display)
    if not date_iso:
        errors.append("Некорректная дата заключения.")

    return errors


async def save_draft_snapshot(
    user_id: int,
    data: ConclusionData,
    next_state: Optional[DialogState],
    title: Optional[str] = None,
    metrics: Optional[Dict[str, Any]] = None,
) -> None:
    """Сохраняет промежуточные данные пользователя как черновик."""
    if not _is_db_ready():
        return
    if not isinstance(data, ConclusionData):
        return

    payload = data.to_dict()
    state_value: Optional[str] = None
    if isinstance(next_state, DialogState):
        state_value = next_state.name
    elif isinstance(next_state, str):
        state_value = next_state

    timestamp = _now_iso()

    metrics_blob = json.dumps(metrics, ensure_ascii=False) if metrics else None

    async with db_lock:
        existing_row = None
        try:
            async with db.execute(
                "SELECT created_at, title FROM drafts WHERE user_id = ?",
                (user_id,),
            ) as cursor:
                existing_row = await cursor.fetchone()
        except Exception as db_error:
            logger.error(f"Не удалось прочитать черновик пользователя {user_id}: {db_error}")
            return

        try:
            if existing_row:
                created_at, existing_title = existing_row
                title_to_save = title if title is not None else existing_title
                await db.execute(
                    "UPDATE drafts SET data = ?, state = ?, title = ?, updated_at = ?, metrics = ? WHERE user_id = ?",
                    (
                        json.dumps(payload, ensure_ascii=False),
                        state_value,
                        title_to_save,
                        timestamp,
                        metrics_blob,
                        user_id,
                    ),
                )
            else:
                await db.execute(
                    "INSERT INTO drafts (user_id, data, state, title, created_at, updated_at, metrics) VALUES (?, ?, ?, ?, ?, ?, ?)",
                    (
                        user_id,
                        json.dumps(payload, ensure_ascii=False),
                        state_value,
                        title,
                        timestamp,
                        timestamp,
                        metrics_blob,
                    ),
                )
            await db.commit()
        except Exception as db_error:
            logger.error(f"Не удалось сохранить черновик пользователя {user_id}: {db_error}")


async def fetch_draft(user_id: int) -> Optional[Dict[str, Any]]:
    """Возвращает сохранённый черновик пользователя."""
    if not _is_db_ready():
        return None
    async with db_lock:
        try:
            async with db.execute(
                "SELECT data, state, title, created_at, updated_at, metrics FROM drafts WHERE user_id = ?",
                (user_id,),
            ) as cursor:
                row = await cursor.fetchone()
        except Exception as db_error:
            logger.error(f"Не удалось получить черновик пользователя {user_id}: {db_error}")
            return None

    if not row:
        return None

    data_blob, state_value, title, created_at, updated_at, metrics_blob = row
    try:
        payload = json.loads(data_blob) if data_blob else {}
    except json.JSONDecodeError:
        payload = {}
    metrics_payload = None
    if metrics_blob:
        try:
            metrics_payload = json.loads(metrics_blob)
        except json.JSONDecodeError:
            metrics_payload = None
    return {
        "user_id": user_id,
        "data": payload,
        "conclusion": _conclusion_from_payload(payload),
        "state": state_value,
        "title": title,
        "created_at": created_at,
        "updated_at": updated_at,
        "metrics": metrics_payload,
    }


async def clear_draft(user_id: int) -> None:
    """Удаляет черновик пользователя."""
    if not _is_db_ready():
        return
    async with db_lock:
        try:
            await db.execute("DELETE FROM drafts WHERE user_id = ?", (user_id,))
            await db.commit()
        except Exception as db_error:
            logger.error(f"Не удалось удалить черновик пользователя {user_id}: {db_error}")


async def persist_draft(
    context: CallbackContext,
    user_id: int,
    data: ConclusionData,
    next_state: Optional[DialogState],
    title: Optional[str] = None,
) -> None:
    metrics = metrics_snapshot(context)
    await save_draft_snapshot(user_id, data, next_state, title=title, metrics=metrics)


async def get_user_flags(user_id: int) -> Dict[str, Any]:
    default = {
        "user_id": user_id,
        "is_blocked": False,
        "notes": "",
        "created_at": None,
        "updated_at": None,
    }
    if not _is_db_ready():
        return default
    async with db_lock:
        try:
            async with db.execute(
                "SELECT is_blocked, notes, created_at, updated_at FROM user_flags WHERE user_id = ?",
                (user_id,),
            ) as cursor:
                row = await cursor.fetchone()
        except Exception as error:
            logger.error(f"Не удалось получить флаги пользователя {user_id}: {error}")
            return default
    if not row:
        return default
    is_blocked, notes, created_at, updated_at = row
    default.update(
        {
            "is_blocked": bool(is_blocked),
            "notes": notes or "",
            "created_at": created_at,
            "updated_at": updated_at,
        }
    )
    return default


async def set_user_block_status(user_id: int, blocked: bool, notes: Optional[str] = None) -> None:
    if not _is_db_ready():
        return
    timestamp = _now_iso()
    if notes is None:
        existing_flags = await get_user_flags(user_id)
        notes_value = existing_flags.get("notes", "")
    else:
        notes_value = notes
    async with db_lock:
        try:
            await db.execute(
                "INSERT INTO user_flags (user_id, is_blocked, notes, created_at, updated_at) "
                "VALUES (?, ?, ?, ?, ?) "
                "ON CONFLICT(user_id) DO UPDATE SET is_blocked = excluded.is_blocked, notes = excluded.notes, updated_at = excluded.updated_at",
                (
                    user_id,
                    1 if blocked else 0,
                    notes_value,
                    timestamp,
                    timestamp,
                ),
            )
            await db.commit()
        except Exception as error:
            logger.error(f"Не удалось обновить статус блокировки пользователя {user_id}: {error}")


async def update_user_notes(user_id: int, notes: str) -> None:
    if not _is_db_ready():
        return
    timestamp = _now_iso()
    async with db_lock:
        try:
            await db.execute(
                "INSERT INTO user_flags (user_id, is_blocked, notes, created_at, updated_at) "
                "VALUES (?, 0, ?, ?, ?) "
                "ON CONFLICT(user_id) DO UPDATE SET notes = excluded.notes, updated_at = excluded.updated_at",
                (
                    user_id,
                    notes,
                    timestamp,
                    timestamp,
                ),
            )
            await db.commit()
        except Exception as error:
            logger.error(f"Не удалось обновить заметку пользователя {user_id}: {error}")


async def is_user_blocked(user_id: int) -> bool:
    flags = await get_user_flags(user_id)
    return flags.get("is_blocked", False)


async def ensure_user_not_blocked_message(update: Update, context: CallbackContext) -> bool:
    user = update.effective_user
    if not user:
        return True
    if await is_user_blocked(user.id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return False
    return True


async def ensure_user_not_blocked_query(query, context: CallbackContext) -> bool:
    user = getattr(query, "from_user", None)
    if not user:
        return True
    if await is_user_blocked(user.id):
        try:
            await query.answer("Ваш доступ к боту ограничен.", show_alert=True)
        except TelegramError:
            pass
        return False
    return True


async def fetch_recent_users(limit: int = 10) -> List[Dict[str, Any]]:
    if not _is_db_ready():
        return []
    async with db_lock:
        async with db.execute(
            "SELECT user_id, MAX(username) as username, MAX(completed_at) as last_completed"
            " FROM completions WHERE is_deleted IS NULL OR is_deleted = 0"
            " GROUP BY user_id ORDER BY last_completed DESC LIMIT ?",
            (limit,),
        ) as cursor:
            rows = await cursor.fetchall()

    users: List[Dict[str, Any]] = []
    for user_id, username, last_completed in rows:
        flags = await get_user_flags(user_id)
        users.append(
            {
                "user_id": user_id,
                "username": username or "—",
                "last_completed": last_completed,
                "is_blocked": flags.get("is_blocked", False),
                "notes": flags.get("notes", ""),
            }
        )
    return users


async def fetch_soft_deleted_completions(limit: int = 5) -> List[Dict[str, Any]]:
    if not _is_db_ready():
        return []
    async with db_lock:
        async with db.execute(
            "SELECT id, user_id, username, completed_at, ticket_number, issue_number, department_number, date, region, deleted_at "
            "FROM completions WHERE is_deleted = 1 ORDER BY deleted_at DESC LIMIT ?",
            (limit,),
        ) as cursor:
            rows = await cursor.fetchall()

    entries: List[Dict[str, Any]] = []
    for row in rows:
        (
            comp_id,
            user_id,
            username,
            completed_at,
            ticket_number,
            issue_number,
            department_number,
            date_val,
            region,
            deleted_at,
        ) = row
        entries.append(
            {
                "id": comp_id,
                "user_id": user_id,
                "username": username,
                "completed_at": completed_at,
                "ticket_number": ticket_number,
                "issue_number": issue_number,
                "department_number": department_number,
                "date": date_val,
                "region": region,
                "deleted_at": deleted_at,
            }
        )
    return entries


async def _ensure_table_columns(table: str, declarations: Dict[str, str]) -> None:
    """Гарантирует наличие указанных колонок в таблице."""
    if not _is_db_ready():
        return
    async with db_lock:
        try:
            async with db.execute(f"PRAGMA table_info({table})") as cursor:
                rows = await cursor.fetchall()
        except Exception as e:
            logger.error(f"Не удалось получить информацию о таблице {table}: {e}")
            return

        existing = {row[1] for row in rows}
        pending = {name: ddl for name, ddl in declarations.items() if name not in existing}
        for column_name, column_definition in pending.items():
            try:
                await db.execute(f"ALTER TABLE {table} ADD COLUMN {column_name} {column_definition}")
                logger.info(f"В таблицу {table} добавлена колонка {column_name}.")
            except Exception as e:
                logger.error(f"Не удалось добавить колонку {column_name} в {table}: {e}")
        if pending:
            await db.commit()


def _current_month_key(target: Optional[datetime] = None) -> str:
    base = target or datetime.now()
    return base.strftime("%Y-%m")


def _format_month_label(month_key: str) -> str:
    try:
        dt = datetime.strptime(f"{month_key}-01", "%Y-%m-%d")
        return dt.strftime("%m.%Y")
    except ValueError:
        return month_key


def _normalize_day(dt: datetime) -> datetime:
    return dt.replace(hour=0, minute=0, second=0, microsecond=0)


def _period_bounds(period: str, reference: Optional[datetime] = None) -> Tuple[datetime, datetime, str]:
    now = reference or datetime.now()
    if period == "week":
        start = _normalize_day(now - timedelta(days=now.weekday()))
        end = start + timedelta(days=7)
        label = f"{start.strftime('%d.%m')}–{(end - timedelta(days=1)).strftime('%d.%m')}"
        return start, end, label
    start = _normalize_day(now.replace(day=1))
    if start.month == 12:
        end = start.replace(year=start.year + 1, month=1)
    else:
        end = start.replace(month=start.month + 1)
    label = start.strftime("%m.%Y")
    return start, end, label


def _collect_completion_metrics(data: Union[ConclusionData, Dict[str, Any]]) -> Tuple[int, float]:
    if isinstance(data, ConclusionData):
        items = data.photo_desc
    else:
        items = data.get("photo_desc", []) or []
    item_count = len(items)
    total_evaluation = 0.0
    for item in items:
        try:
            total_evaluation += float(item.get("evaluation") or 0)
        except (TypeError, ValueError):
            continue
    return item_count, total_evaluation


async def record_completion_entry(
    user_id: int,
    username: str,
    data: Union[ConclusionData, Dict[str, Any]],
    group_chat_id: Optional[int] = None,
    group_message_id: Optional[int] = None,
    thread_id: Optional[int] = None,
    archive_path: Optional[Path] = None,
    processing_time: Optional[float] = None,
    step_metrics: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Фиксирует готовое заключение для подсчёта ачивок и рейтингов."""
    if not _is_db_ready():
        return {}

    if isinstance(data, ConclusionData):
        data_dict = data.to_dict()
    else:
        data_dict = data or {}

    completed_at = datetime.now()
    month_key = _current_month_key(completed_at)
    day_key = completed_at.strftime("%Y-%m-%d")
    month_label = _format_month_label(month_key)
    day_label = completed_at.strftime("%d.%m.%Y")

    region = data_dict.get("region")
    ticket_number = data_dict.get("ticket_number")
    issue_number = data_dict.get("issue_number")
    department_number = data_dict.get("department_number")
    conclusion_date = data_dict.get("date")

    item_count, total_evaluation = _collect_completion_metrics(data_dict)
    completion_xp = calculate_completion_xp(item_count, total_evaluation)
    items_json = json.dumps(data_dict.get("photo_desc", []) or [])
    archive_path_str = str(archive_path) if archive_path else None
    chat_id_to_store = group_chat_id or MAIN_GROUP_CHAT_ID
    processing_value = float(processing_time or 0.0)
    metrics_json = json.dumps(step_metrics, ensure_ascii=False) if step_metrics else None
    previous_items_total = 0
    previous_value_total = 0.0
    previous_xp_total = 0

    async with db_lock:
        async with db.execute(
            "SELECT COUNT(*) FROM completions WHERE user_id = ? AND substr(completed_at, 1, 7) = ? AND (is_deleted IS NULL OR is_deleted = 0)",
            (user_id, month_key),
        ) as cursor:
            row = await cursor.fetchone()
            previous_month_count = row[0] if row else 0

        async with db.execute(
            "SELECT COUNT(*), COALESCE(SUM(item_count), 0), COALESCE(SUM(total_evaluation), 0) "
            "FROM completions WHERE user_id = ? AND (is_deleted IS NULL OR is_deleted = 0)",
            (user_id,),
        ) as cursor:
            row = await cursor.fetchone()
            if row:
                previous_total_count = row[0] or 0
                previous_items_total = int(row[1] or 0)
                previous_value_total = float(row[2] or 0.0)
            else:
                previous_total_count = 0
                previous_items_total = 0
                previous_value_total = 0.0
        async with db.execute(
            "SELECT COALESCE(SUM(xp_value), 0) FROM completions WHERE user_id = ? AND (is_deleted IS NULL OR is_deleted = 0)",
            (user_id,),
        ) as cursor:
            row = await cursor.fetchone()
            if row:
                previous_xp_total = int(row[0] or 0)

        async with db.execute(
            "SELECT COUNT(*) FROM completions WHERE user_id = ? AND substr(completed_at, 1, 10) = ? AND (is_deleted IS NULL OR is_deleted = 0)",
            (user_id, day_key),
        ) as cursor:
            row = await cursor.fetchone()
            previous_day_count = row[0] if row else 0

        cursor = await db.execute(
            """
            INSERT INTO completions (
                user_id, username, completed_at, item_count, total_evaluation, region,
                ticket_number, issue_number, department_number, date,
                group_chat_id, group_message_id, thread_id, archive_path, items_json, xp_value,
                processing_time_seconds, step_metrics
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                user_id,
                username,
                completed_at.isoformat(),
                item_count,
                total_evaluation,
                region,
                ticket_number,
                issue_number,
                department_number,
                conclusion_date,
                chat_id_to_store,
                group_message_id,
                thread_id,
                archive_path_str,
                items_json,
                completion_xp,
                processing_value,
                metrics_json,
            ),
        )
        completion_id = cursor.lastrowid or 0
        async with db.execute(
            "SELECT substr(completed_at, 1, 10) as day_key FROM completions WHERE user_id = ? AND (is_deleted IS NULL OR is_deleted = 0) GROUP BY day_key ORDER BY day_key DESC",
            (user_id,),
        ) as cursor:
            streak_rows = await cursor.fetchall()
        day_strings = [row[0] for row in streak_rows if row[0]]
        current_streak = _calculate_streak_from_days(day_strings)
        previous_streak = current_streak
        if previous_day_count == 0 and current_streak > 0:
            previous_streak = current_streak - 1
        await db.commit()

    current_items_total = previous_items_total + item_count
    current_value_total = previous_value_total + total_evaluation
    current_xp_total = previous_xp_total + completion_xp
    previous_level_info = level_profile_from_xp(previous_xp_total)
    current_level_info = level_profile_from_xp(current_xp_total)

    return {
        "completion_id": completion_id,
        "month_key": month_key,
        "month_label": month_label,
        "day_key": day_key,
        "day_label": day_label,
        "previous_month": previous_month_count,
        "current_month": previous_month_count + 1,
        "previous_total": previous_total_count,
        "current_total": previous_total_count + 1,
        "previous_day": previous_day_count,
        "current_day": previous_day_count + 1,
        "region": region,
        "ticket_number": ticket_number,
        "issue_number": issue_number,
        "department_number": department_number,
        "date": conclusion_date,
        "group_chat_id": chat_id_to_store,
        "group_message_id": group_message_id,
        "thread_id": thread_id,
        "item_count": item_count,
        "total_evaluation": total_evaluation,
        "previous_items_total": previous_items_total,
        "current_items_total": current_items_total,
        "previous_value_total": previous_value_total,
        "current_value_total": current_value_total,
        "previous_xp": previous_xp_total,
        "current_xp": current_xp_total,
        "completion_xp": completion_xp,
        "processing_time_seconds": processing_value,
        "step_metrics": step_metrics or {},
        "previous_level": previous_level_info.get("level", 0),
        "current_level": current_level_info.get("level", 0),
        "current_level_label": current_level_info.get("label"),
        "xp_to_next": current_level_info.get("xp_to_next"),
        "next_level_threshold": current_level_info.get("xp_next"),
        "previous_streak": max(previous_streak, 0),
        "current_streak": current_streak,
    }


async def fetch_leaderboard(period: str, limit: int = LEADERBOARD_SIZE) -> Tuple[str, List[Dict[str, Any]]]:
    """Возвращает строку-подпись и список лидеров за указанный период."""
    if not _is_db_ready():
        return "", []

    start, end, label = _period_bounds(period)
    start_iso = start.isoformat()
    end_iso = end.isoformat()

    async with db_lock:
        async with db.execute(
            """
            SELECT
                user_id,
                COALESCE(MAX(username), 'Неизвестно') AS display_name,
                COUNT(*) AS completions_count,
                COALESCE(SUM(item_count), 0) AS items_total,
                COALESCE(SUM(total_evaluation), 0) AS total_evaluation,
                MAX(completed_at) AS last_completed_at
            FROM completions
            WHERE completed_at >= ? AND completed_at < ? AND (is_deleted IS NULL OR is_deleted = 0)
            GROUP BY user_id
            ORDER BY completions_count DESC, last_completed_at ASC
            LIMIT ?
            """,
            (start_iso, end_iso, limit),
        ) as cursor:
            rows = await cursor.fetchall()

    leaderboard = []
    for user_id_value, display_name, completions_count, items_total, total_evaluation, _ in rows:
        leaderboard.append(
            {
                "user_id": user_id_value,
                "username": display_name or "Неизвестно",
                "completions": int(completions_count),
                "items": int(items_total),
                "total_evaluation": float(total_evaluation or 0.0),
            }
        )
    return label, leaderboard




def _calculate_streak_from_days(day_strings: List[str]) -> int:
    if not day_strings:
        return 0
    today = datetime.now().date()
    streak = 0
    expected = today
    for day_str in day_strings:
        try:
            day_date = datetime.strptime(day_str, "%Y-%m-%d").date()
        except ValueError:
            continue
        if day_date == expected:
            streak += 1
            expected -= timedelta(days=1)
        elif day_date < expected:
            break
    return streak

async def send_achievement_notification(bot, chat_id: int, text: str, media_path: Optional[Path]) -> None:
    """Отправляет медиа-награду (фото/видео) или текст при неудаче."""
    if media_path and media_path.is_file():
        suffix = media_path.suffix.lower()
        try:
            with media_path.open("rb") as media_file:
                if suffix in {".mp4", ".mov", ".m4v"}:
                    await bot.send_video(chat_id=chat_id, video=media_file, caption=text)
                elif suffix in {".gif"}:
                    await bot.send_animation(chat_id=chat_id, animation=media_file, caption=text)
                elif suffix in {".png", ".jpg", ".jpeg", ".webp"}:
                    await bot.send_photo(chat_id=chat_id, photo=media_file, caption=text)
                else:
                    await bot.send_document(chat_id=chat_id, document=media_file, caption=text)
            return
        except (RetryAfter, TelegramError, NetworkError, asyncio.TimeoutError) as error:
            logger.warning(f"Не удалось отправить медиа ачивки {media_path}: {error}")
    await safe_bot_send_message(bot, chat_id, text, skip_notice_on_retry=True)


async def _register_achievement(user_id: int, achievement_key: str) -> bool:
    if not _is_db_ready():
        return False
    async with db_lock:
        try:
            await db.execute(
                "INSERT INTO achievement_log (user_id, achievement_key, achieved_at) VALUES (?, ?, ?)",
                (user_id, achievement_key, datetime.now().isoformat()),
            )
        except IntegrityError:
            return False
        except Exception as error:
            logger.error(f"Не удалось записать ачивку {achievement_key} для пользователя {user_id}: {error}")
            return False
        else:
            await db.commit()
            return True


async def refresh_achievements_for_user(user_id: int) -> None:
    if not _is_db_ready():
        return

    async with db_lock:
        try:
            async with db.execute("SELECT completed_at FROM completions WHERE user_id = ?", (user_id,)) as cursor:
                completion_rows = await cursor.fetchall()
            async with db.execute("SELECT achievement_key FROM achievement_log WHERE user_id = ?", (user_id,)) as cursor:
                existing_rows = await cursor.fetchall()
            async with db.execute(
                "SELECT substr(completed_at, 1, 7) as month_key, COUNT(*) FROM completions WHERE user_id = ? GROUP BY month_key",
                (user_id,),
            ) as cursor:
                month_rows = await cursor.fetchall()
            async with db.execute(
                "SELECT substr(completed_at, 1, 10) as day_key, COUNT(*) FROM completions WHERE user_id = ? GROUP BY day_key",
                (user_id,),
            ) as cursor:
                day_rows = await cursor.fetchall()
            async with db.execute(
                "SELECT substr(completed_at, 1, 10) as day_key FROM completions WHERE user_id = ? GROUP BY day_key ORDER BY day_key DESC",
                (user_id,),
            ) as cursor:
                streak_rows = await cursor.fetchall()
            async with db.execute(
                "SELECT COALESCE(SUM(item_count), 0), COALESCE(SUM(total_evaluation), 0), COALESCE(SUM(xp_value), 0) FROM completions WHERE user_id = ?",
                (user_id,),
            ) as cursor:
                totals_row = await cursor.fetchone()
        except Exception as error:
            logger.error(f"Не удалось обновить ачивки пользователя {user_id}: {error}")
            return

        total_count = len(completion_rows)
        monthly_counts = {row[0]: row[1] for row in month_rows if row[0]}
        daily_counts = {row[0]: row[1] for row in day_rows if row[0]}
        day_strings = [row[0] for row in streak_rows if row[0]]
        current_streak = _calculate_streak_from_days(day_strings)
        items_total = int((totals_row[0] if totals_row else 0) or 0)
        value_total = float((totals_row[1] if totals_row else 0.0) or 0.0)
        xp_total = int((totals_row[2] if totals_row else 0) or 0)

        existing_keys = {row[0] for row in existing_rows}
        required_keys: Set[str] = set()

        for tier in ACHIEVEMENT_TIERS.get("total", []):
            if total_count >= tier["threshold"]:
                required_keys.add(tier["code"])

        for tier in ACHIEVEMENT_TIERS.get("items_total", []):
            if items_total >= tier["threshold"]:
                required_keys.add(tier["code"])

        for tier in ACHIEVEMENT_TIERS.get("value_total", []):
            if value_total >= tier["threshold"]:
                required_keys.add(tier["code"])

        for tier in ACHIEVEMENT_TIERS.get("level", []):
            if xp_total >= tier["threshold"]:
                required_keys.add(tier["code"])

        for month_key, count in monthly_counts.items():
            for tier in ACHIEVEMENT_TIERS.get("monthly", []):
                if count >= tier["threshold"]:
                    required_keys.add(f"{tier['code']}:{month_key}")

        for day_key, count in daily_counts.items():
            for tier in ACHIEVEMENT_TIERS.get("daily", []):
                if count >= tier["threshold"]:
                    required_keys.add(f"{tier['code']}:{day_key}")

        for tier in ACHIEVEMENT_TIERS.get("streak", []):
            if current_streak >= tier["threshold"]:
                required_keys.add(tier["code"])

        keys_to_remove = existing_keys - required_keys
        if keys_to_remove:
            await db.executemany(
                "DELETE FROM achievement_log WHERE user_id = ? AND achievement_key = ?",
                [(user_id, key) for key in keys_to_remove],
            )
        await db.commit()


def build_progress_lines(stats: Dict[str, Any]) -> List[str]:
    lines: List[str] = []

    def _progress_line(metric: str, current: float, context_label: str = "", unit_label: str = ""):
        tiers = ACHIEVEMENT_TIERS.get(metric, [])
        for tier in tiers:
            threshold = tier["threshold"]
            if current < threshold:
                remaining_value = threshold - current
                if remaining_value <= 0:
                    continue
                remaining_text = format_number(math.ceil(remaining_value))
                if unit_label:
                    remaining_text = f"{remaining_text}{unit_label}"
                label = tier["title"]
                if context_label:
                    return f"{context_label}: до уровня «{label}» осталось {remaining_text}."
                return f"До уровня «{label}» осталось {remaining_text}."
        return None

    total_line = _progress_line("total", stats.get("current_total", 0))
    if total_line:
        lines.append(total_line)

    level_line = _progress_line("level", stats.get("current_xp", 0), "Уровень", " очков опыта")
    if level_line:
        lines.append(level_line)

    items_line = _progress_line("items_total", stats.get("current_items_total", 0), "Предметов всего")
    if items_line:
        lines.append(items_line)

    value_line = _progress_line("value_total", stats.get("current_value_total", 0.0), "Сумма оценок", " руб.")
    if value_line:
        lines.append(value_line)

    month_key = stats.get("month_key")
    if month_key:
        month_line = _progress_line("monthly", stats.get("current_month", 0), stats.get("month_label", "Текущий месяц"))
        if month_line:
            lines.append(month_line)

    day_key = stats.get("day_key")
    if day_key:
        day_line = _progress_line("daily", stats.get("current_day", 0), stats.get("day_label", "Сегодня"))
        if day_line:
            lines.append(day_line)

    streak_line = _progress_line("streak", stats.get("current_streak", 0), "Текущая серия")
    if streak_line:
        lines.append(streak_line)

    return lines


async def send_progress_overview(bot, user_id: int, stats: Dict[str, Any]) -> None:
    lines = build_progress_lines(stats)
    if lines:
        bullets = "\n".join(f"• {line}" for line in lines)
        await safe_bot_send_message(
            bot,
            user_id,
            f"Следующие цели:\n{bullets}",
            skip_notice_on_retry=True,
        )
    else:
        await safe_bot_send_message(
            bot,
            user_id,
            "Все текущие цели закрыты. Продолжайте в том же духе.",
            skip_notice_on_retry=True,
        )


def _split_achievement_key(key: str) -> Tuple[str, Optional[str]]:
    if ":" in key:
        base, suffix = key.split(":", 1)
        return base, suffix
    return key, None


def _render_achievement_title(key: str) -> str:
    base, suffix = _split_achievement_key(key)
    info = ACHIEVEMENT_LOOKUP.get(base, {})
    title = info.get("title", base)
    metric = info.get("metric")
    if suffix:
        if metric == "monthly":
            title = f"{title} ({_format_month_label(suffix)})"
        elif metric == "daily":
            try:
                day_label = datetime.strptime(suffix, "%Y-%m-%d").strftime("%d.%m.%Y")
            except ValueError:
                day_label = suffix
            title = f"{title} ({day_label})"
        else:
            title = f"{title} ({suffix})"
    return title


async def _fetch_user_achievement_log(user_id: int) -> List[Tuple[str, Optional[str]]]:
    if not _is_db_ready():
        return []
    async with db_lock:
        async with db.execute(
            "SELECT achievement_key, achieved_at FROM achievement_log WHERE user_id = ? ORDER BY achieved_at DESC",
            (user_id,),
        ) as cursor:
            rows = await cursor.fetchall()
    return rows or []


async def _fetch_user_totals(user_id: int) -> Dict[str, Any]:
    if not _is_db_ready():
        return {
            "count": 0,
            "items": 0,
            "value": 0.0,
            "xp": 0,
            "streak": 0,
        }

    async with db_lock:
        async with db.execute(
            "SELECT COUNT(*), COALESCE(SUM(item_count), 0), COALESCE(SUM(total_evaluation), 0), COALESCE(SUM(xp_value), 0) "
            "FROM completions WHERE user_id = ? AND (is_deleted IS NULL OR is_deleted = 0)",
            (user_id,),
        ) as cursor:
            row = await cursor.fetchone()
        async with db.execute(
            "SELECT substr(completed_at, 1, 10) FROM completions WHERE user_id = ? AND (is_deleted IS NULL OR is_deleted = 0) "
            "GROUP BY substr(completed_at, 1, 10) ORDER BY substr(completed_at, 1, 10) DESC",
            (user_id,),
        ) as cursor:
            streak_rows = await cursor.fetchall()

    total_count = int(row[0] or 0) if row else 0
    items_total = int(row[1] or 0) if row else 0
    value_total = float(row[2] or 0.0) if row else 0.0
    xp_total = int(row[3] or 0) if row else 0
    day_strings = [value[0] for value in streak_rows if value and value[0]] if streak_rows else []
    streak = _calculate_streak_from_days(day_strings)

    return {
        "count": total_count,
        "items": items_total,
        "value": value_total,
        "xp": xp_total,
        "streak": streak,
    }


async def gather_user_achievement_summary(user_id: int) -> Dict[str, Any]:
    totals = await _fetch_user_totals(user_id)
    stats_period = await fetch_user_stats(user_id)
    achievement_rows = await _fetch_user_achievement_log(user_id)

    level_info = level_profile_from_xp(totals.get("xp", 0))
    unlocked_keys = [row[0] for row in achievement_rows]
    static_unlocked = {
        _split_achievement_key(key)[0]
        for key in unlocked_keys
        if _split_achievement_key(key)[0] in ACHIEVEMENT_LOOKUP
    }

    recent_entries: List[str] = []
    for key, timestamp in achievement_rows[:5]:
        title = _render_achievement_title(key)
        if timestamp:
            try:
                ts_label = datetime.fromisoformat(timestamp).strftime("%d.%m")
                recent_entries.append(f"{ts_label} • {title}")
            except ValueError:
                recent_entries.append(title)
        else:
            recent_entries.append(title)

    today = datetime.now()
    summary: Dict[str, Any] = {
        "current_total": totals.get("count", 0),
        "current_items_total": totals.get("items", 0),
        "current_value_total": totals.get("value", 0.0),
        "current_xp": totals.get("xp", 0),
        "previous_total": totals.get("count", 0),
        "previous_items_total": totals.get("items", 0),
        "previous_value_total": totals.get("value", 0.0),
        "previous_xp": totals.get("xp", 0),
        "current_level": level_info.get("level", 0),
        "current_level_label": level_info.get("label", "🌱 Новичок"),
        "xp_to_next": level_info.get("xp_to_next", 0),
        "next_level_threshold": level_info.get("xp_next", 0),
        "current_streak": totals.get("streak", 0),
        "previous_streak": max(totals.get("streak", 0) - 1, 0),
        "month_key": _current_month_key(today),
        "month_label": stats_period.get("monthly", {}).get("label", today.strftime("%m.%Y")),
        "current_month": stats_period.get("monthly", {}).get("count", 0),
        "day_key": today.strftime("%Y-%m-%d"),
        "day_label": stats_period.get("daily", {}).get("label", today.strftime("%d.%m.%Y")),
        "current_day": stats_period.get("daily", {}).get("count", 0),
        "unlocked_keys": unlocked_keys,
        "unlocked_count": len(unlocked_keys),
        "static_unlocked_count": len(static_unlocked),
        "available_static_count": len(ACHIEVEMENT_LOOKUP),
        "recent_entries": recent_entries,
        "has_records": totals.get("count", 0) > 0,
    }
    summary.update(stats_period)
    return summary


def build_achievements_keyboard(view: str = ACHIEVEMENTS_DEFAULT_VIEW) -> InlineKeyboardMarkup:
    if view == "main":
        buttons = [
            [
                InlineKeyboardButton("🎯 Цели", callback_data=f"{ACHIEVEMENTS_CALLBACK_PREFIX}goals"),
                InlineKeyboardButton("🏆 Уровни", callback_data=f"{ACHIEVEMENTS_CALLBACK_PREFIX}levels"),
            ],
            [InlineKeyboardButton("✖ Закрыть", callback_data=f"{ACHIEVEMENTS_CALLBACK_PREFIX}close")],
        ]
        return InlineKeyboardMarkup(buttons)

    rows: List[List[InlineKeyboardButton]] = [
        [InlineKeyboardButton("⬅️ Назад", callback_data=f"{ACHIEVEMENTS_CALLBACK_PREFIX}main")]
    ]
    secondary: List[InlineKeyboardButton] = []
    if view != "goals":
        secondary.append(InlineKeyboardButton("🎯 Цели", callback_data=f"{ACHIEVEMENTS_CALLBACK_PREFIX}goals"))
    if view != "levels":
        secondary.append(InlineKeyboardButton("🏆 Уровни", callback_data=f"{ACHIEVEMENTS_CALLBACK_PREFIX}levels"))
    if secondary:
        rows.append(secondary)
    rows.append([InlineKeyboardButton("✖ Закрыть", callback_data=f"{ACHIEVEMENTS_CALLBACK_PREFIX}close")])
    return InlineKeyboardMarkup(rows)


def build_achievements_main_text(summary: Dict[str, Any]) -> str:
    lines = [
        "🏅 Профиль достижений",
        f"{summary.get('current_level_label', '🌱 Новичок')} — уровень {summary.get('current_level', 0)}",
    ]
    xp_total = summary.get("current_xp", 0)
    xp_to_next = summary.get("xp_to_next", 0)
    if xp_to_next > 0:
        lines.append(
            f"Опыт: {format_number(xp_total)} очков (до следующего уровня {format_number(xp_to_next)} очков)"
        )
    else:
        lines.append(f"Опыт: {format_number(xp_total)} очков — максимальный уровень достигнут.")

    lines.append("")
    lines.append(f"Наград открыто: {summary.get('unlocked_count', 0)}")
    lines.append(
        f"Основных: {summary.get('static_unlocked_count', 0)} из {summary.get('available_static_count', 0)}"
    )
    lines.append("")
    lines.append("Очки опыта начисляются за оформленные окончательные заключения: учитывается число предметов и их оценка.")

    if summary.get("recent_entries"):
        lines.append("")
        lines.append("Последние награды:")
        for entry in summary["recent_entries"]:
            lines.append(f"• {entry}")
    else:
        lines.append("")
        lines.append("Последние награды: пока нет.")

    goal_lines = build_progress_lines(summary)
    if goal_lines:
        lines.append("")
        lines.append("Ближайшие цели:")
        for goal in goal_lines[:3]:
            lines.append(f"• {goal}")

    return "\n".join(lines)


def build_achievements_goals_text(summary: Dict[str, Any]) -> str:
    lines = ["🎯 Ближайшие цели"]
    goal_lines = build_progress_lines(summary)
    if goal_lines:
        lines.extend(f"• {goal}" for goal in goal_lines)
    else:
        lines.append("Все цели закрыты — можно отдыхать.")
    return "\n".join(lines)


def build_achievements_levels_text(summary: Dict[str, Any]) -> str:
    lines = [
        "🏆 Лестница уровней",
        f"Текущий уровень: {summary.get('current_level_label', '🌱 Новичок')} — {summary.get('current_level', 0)}",
    ]
    xp_to_next = summary.get("xp_to_next", 0)
    if xp_to_next > 0:
        lines.append(f"До следующего уровня: {format_number(xp_to_next)} очков опыта")
    else:
        lines.append("Следующая ступень недоступна — вы на вершине.")

    lines.append("")
    lines.append("Ближайшие ступени:")
    current_level = summary.get("current_level", 0)
    current_xp = summary.get("current_xp", 0)
    start_index = max(0, current_level - 3)
    preview = LEVEL_CATALOG[start_index:start_index + 6]
    for entry in preview:
        status = "✅" if current_xp >= entry["xp_required"] else "•"
        lines.append(f"{status} {entry['title']} — {format_number(entry['xp_required'])} очков опыта")
    return "\n".join(lines)


async def build_achievements_view(user_id: int, view: str = ACHIEVEMENTS_DEFAULT_VIEW) -> Optional[Dict[str, Any]]:
    summary = await gather_user_achievement_summary(user_id)
    if not summary:
        return None

    view_key = view if view in {"main", "goals", "levels"} else ACHIEVEMENTS_DEFAULT_VIEW
    if view_key == "goals":
        text = build_achievements_goals_text(summary)
    elif view_key == "levels":
        text = build_achievements_levels_text(summary)
    else:
        text = build_achievements_main_text(summary)
        view_key = ACHIEVEMENTS_DEFAULT_VIEW

    keyboard = build_achievements_keyboard(view_key)
    return {
        "text": text,
        "keyboard": keyboard,
        "view": view_key,
    }

async def handle_new_achievements(
    bot,
    user_id: int,
    username: str,
    stats: Dict[str, Any],
    region: Optional[str],
) -> None:
    metric_contexts = {
        "total": {
            "previous": stats.get("previous_total", 0),
            "current": stats.get("current_total", 0),
            "suffix": None,
        },
        "level": {
            "previous": stats.get("previous_xp", 0),
            "current": stats.get("current_xp", 0),
            "suffix": None,
        },
        "items_total": {
            "previous": stats.get("previous_items_total", 0),
            "current": stats.get("current_items_total", 0),
            "suffix": None,
        },
        "value_total": {
            "previous": stats.get("previous_value_total", 0.0),
            "current": stats.get("current_value_total", 0.0),
            "suffix": None,
        },
        "monthly": {
            "previous": stats.get("previous_month", 0),
            "current": stats.get("current_month", 0),
            "suffix": stats.get("month_key"),
        },
        "daily": {
            "previous": stats.get("previous_day", 0),
            "current": stats.get("current_day", 0),
            "suffix": stats.get("day_key"),
        },
        "streak": {
            "previous": stats.get("previous_streak", 0),
            "current": stats.get("current_streak", 0),
            "suffix": None,
        },
    }
    unlocked_titles: List[str] = []
    format_context = {
        "username": username,
        "region": region or "",
        **stats,
    }
    format_context.update(
        {
            "current_items_total": stats.get("current_items_total", 0),
            "previous_items_total": stats.get("previous_items_total", 0),
            "current_value_total": stats.get("current_value_total", 0.0),
            "previous_value_total": stats.get("previous_value_total", 0.0),
            "current_value_total_fmt": format_number(stats.get("current_value_total", 0.0)),
            "current_level": stats.get("current_level", 0),
            "previous_level": stats.get("previous_level", 0),
            "current_level_label": stats.get("current_level_label", ""),
            "current_xp": stats.get("current_xp", 0),
            "previous_xp": stats.get("previous_xp", 0),
            "xp_to_next": stats.get("xp_to_next", 0),
        }
    )

    for metric, info in metric_contexts.items():
        suffix = info.get("suffix")
        tiers = ACHIEVEMENT_TIERS.get(metric, [])
        if not tiers:
            continue
        previous_value = info.get("previous", 0) or 0
        current_value = info.get("current", 0) or 0
        if suffix is None and metric in {"monthly", "daily"}:
            # если нечему фиксироваться (нет месяца/дня), пропускаем
            if metric == "monthly" and not stats.get("month_key"):
                continue
            if metric == "daily" and not stats.get("day_key"):
                continue
        for tier in tiers:
            threshold = tier["threshold"]
            if previous_value < threshold <= current_value:
                key = tier["code"] if suffix is None else f"{tier['code']}:{suffix}"
                registered = await _register_achievement(user_id, key)
                if not registered:
                    continue
                personal_text = tier.get("personal_template", "").format(**format_context)
                message_text = f"🏅 Вы получили награду «{tier['title']}»!\n{personal_text}".strip()
                await send_achievement_notification(bot, user_id, message_text, tier.get("media"))
                unlocked_titles.append(tier["title"])

    if unlocked_titles:
        summary_lines = "\n".join(f"• {title}" for title in unlocked_titles)
        await safe_bot_send_message(
            bot,
            user_id,
            f"Новые награды:\n{summary_lines}",
            skip_notice_on_retry=True,
        )

    await send_progress_overview(bot, user_id, stats)


async def _aggregate_user_stats(user_id: int, start: Optional[datetime] = None, end: Optional[datetime] = None) -> Dict[str, Any]:
    if not _is_db_ready():
        return {"count": 0, "items": 0, "total": 0.0}

    query = (
        "SELECT COUNT(*), COALESCE(SUM(item_count), 0), COALESCE(SUM(total_evaluation), 0) "
        "FROM completions WHERE user_id = ? AND (is_deleted IS NULL OR is_deleted = 0)"
    )
    params: List[Any] = [user_id]
    if start:
        query += " AND completed_at >= ?"
        params.append(start.isoformat())
    if end:
        query += " AND completed_at < ?"
        params.append(end.isoformat())

    async with db_lock:
        async with db.execute(query, tuple(params)) as cursor:
            row = await cursor.fetchone()

    if not row:
        return {"count": 0, "items": 0, "total": 0.0}
    count, items, total = row
    return {
        "count": int(count or 0),
        "items": int(items or 0),
        "total": float(total or 0.0),
    }


async def fetch_user_stats(user_id: int) -> Dict[str, Any]:
    if not _is_db_ready():
        return {}
    now = datetime.now()
    day_start = _normalize_day(now)
    day_end = day_start + timedelta(days=1)
    week_start, week_end, week_label = _period_bounds("week", now)
    month_start, month_end, month_label = _period_bounds("month", now)

    daily = await _aggregate_user_stats(user_id, day_start, day_end)
    weekly = await _aggregate_user_stats(user_id, week_start, week_end)
    monthly = await _aggregate_user_stats(user_id, month_start, month_end)
    overall = await _aggregate_user_stats(user_id)

    daily["label"] = day_start.strftime("%d.%m.%Y")
    weekly["label"] = week_label
    monthly["label"] = month_label

    return {
        "daily": daily,
        "weekly": weekly,
        "monthly": monthly,
        "overall": overall,
    }


async def fetch_completion_by_id(completion_id: int) -> Optional[Dict[str, Any]]:
    if not _is_db_ready():
        return None

    async with db_lock:
        async with db.execute(
            "SELECT id, user_id, username, completed_at, item_count, total_evaluation, region, "
            "ticket_number, issue_number, department_number, date, archive_path, items_json, group_chat_id, group_message_id, thread_id, "
            "is_deleted, deleted_at, deleted_by, deletion_note "
            "FROM completions WHERE id = ?",
            (completion_id,),
        ) as cursor:
            row = await cursor.fetchone()

    if not row:
        return None

    (
        comp_id,
        user_id,
        username,
        completed_at,
        item_count,
        total_evaluation,
        region,
        ticket_val,
        issue_val,
        department_val,
        date_val,
        archive_path,
        items_json,
        group_chat_id,
        group_message_id,
        thread_id,
        is_deleted,
        deleted_at,
        deleted_by,
        deletion_note,
    ) = row
    return {
        "id": comp_id,
        "user_id": user_id,
        "username": username,
        "completed_at": completed_at,
        "item_count": item_count,
        "total_evaluation": total_evaluation,
        "region": region,
        "ticket_number": ticket_val,
        "issue_number": issue_val,
        "department_number": department_val,
        "date": date_val,
        "archive_path": archive_path,
        "items_json": items_json,
        "group_chat_id": group_chat_id,
        "group_message_id": group_message_id,
        "thread_id": thread_id,
        "is_deleted": bool(is_deleted),
        "deleted_at": deleted_at,
        "deleted_by": deleted_by,
        "deletion_note": deletion_note,
    }


async def fetch_completion_by_message(
    chat_id: int,
    message_id: int,
    thread_id: Optional[int] = None,
) -> Optional[Dict[str, Any]]:
    if not _is_db_ready():
        return None

    query = (
        "SELECT id, user_id, username, completed_at, item_count, total_evaluation, region, "
        "ticket_number, issue_number, department_number, date, archive_path, items_json, group_chat_id, group_message_id, thread_id, "
        "is_deleted, deleted_at, deleted_by, deletion_note "
        "FROM completions WHERE group_chat_id = ? AND group_message_id = ? AND (is_deleted IS NULL OR is_deleted = 0)"
    )
    params: List[Any] = [chat_id, message_id]
    if thread_id is not None:
        query += " AND (thread_id = ? OR thread_id IS NULL)"
        params.append(thread_id)

    async with db_lock:
        async with db.execute(query, tuple(params)) as cursor:
            row = await cursor.fetchone()

    if not row:
        return None

    (
        completion_id,
        user_id,
        username,
        completed_at,
        item_count,
        total_evaluation,
        region,
        ticket_val,
        issue_val,
        department_val,
        date_val,
        archive_path,
        items_json,
        group_chat_id,
        group_message_id,
        stored_thread_id,
        is_deleted,
        deleted_at,
        deleted_by,
        deletion_note,
    ) = row
    return {
        "id": completion_id,
        "user_id": user_id,
        "username": username,
        "completed_at": completed_at,
        "item_count": item_count,
        "total_evaluation": total_evaluation,
        "region": region,
        "ticket_number": ticket_val,
        "issue_number": issue_val,
        "department_number": department_val,
        "date": date_val,
        "archive_path": archive_path,
        "items_json": items_json,
        "group_chat_id": group_chat_id,
        "group_message_id": group_message_id,
        "thread_id": stored_thread_id,
        "is_deleted": bool(is_deleted),
        "deleted_at": deleted_at,
        "deleted_by": deleted_by,
        "deletion_note": deletion_note,
    }


async def fetch_completions_by_ticket(
    ticket_number: str,
    date_text: Optional[str] = None,
    issue_number: Optional[str] = None,
) -> List[Dict[str, Any]]:
    if not _is_db_ready():
        return []

    query = (
        "SELECT id, user_id, username, completed_at, item_count, total_evaluation, region, "
        "ticket_number, issue_number, department_number, date, archive_path, items_json, group_chat_id, group_message_id, thread_id, "
        "is_deleted, deleted_at, deleted_by, deletion_note "
        "FROM completions WHERE ticket_number = ? AND (is_deleted IS NULL OR is_deleted = 0)"
    )
    params: List[Any] = [ticket_number]
    if date_text:
        query += " AND date = ?"
        params.append(date_text)
    if issue_number:
        query += " AND issue_number = ?"
        params.append(issue_number)
    query += " ORDER BY completed_at DESC"

    async with db_lock:
        async with db.execute(query, tuple(params)) as cursor:
            rows = await cursor.fetchall()

    records: List[Dict[str, Any]] = []
    for row in rows:
        (
            completion_id,
            user_id,
            username,
            completed_at,
            item_count,
            total_evaluation,
            region,
            ticket_val,
            issue_val,
            department_val,
            date_val,
            archive_path,
            items_json,
            group_chat_id,
            group_message_id,
            thread_id,
            is_deleted,
            deleted_at,
            deleted_by,
            deletion_note,
        ) = row
        if is_deleted:
            continue
        records.append(
            {
                "id": completion_id,
                "user_id": user_id,
                "username": username,
                "completed_at": completed_at,
                "item_count": item_count,
                "total_evaluation": total_evaluation,
                "region": region,
                "ticket_number": ticket_val,
                "issue_number": issue_val,
                "department_number": department_val,
                "date": date_val,
                "archive_path": archive_path,
                "items_json": items_json,
                "group_chat_id": group_chat_id,
                "group_message_id": group_message_id,
                "thread_id": thread_id,
                "is_deleted": bool(is_deleted),
                "deleted_at": deleted_at,
                "deleted_by": deleted_by,
                "deletion_note": deletion_note,
            }
        )
    return records


async def build_personal_stats_message(user_id: int) -> Optional[str]:
    stats = await fetch_user_stats(user_id)
    if not stats:
        return None

    daily = stats["daily"]
    weekly = stats["weekly"]
    monthly = stats["monthly"]
    overall = stats["overall"]

    lines = [
        "📈 Ваша статистика:",
        f"Сегодня ({daily['label']}): {daily['count']} заключ., {daily['items']} предметов, {format_number(daily['total'])} руб.",
        f"Неделя ({weekly['label']}): {weekly['count']} заключ., {weekly['items']} предметов, {format_number(weekly['total'])} руб.",
        f"Месяц ({monthly['label']}): {monthly['count']} заключ., {monthly['items']} предметов, {format_number(monthly['total'])} руб.",
        f"Всего: {overall['count']} заключ., {overall['items']} предметов, {format_number(overall['total'])} руб.",
        "Команда /leaders покажет общие рейтинги команды.",
    ]
    return "\n".join(lines)


async def soft_delete_completion_record(
    conclusion: Dict[str, Any],
    initiator_id: Optional[int] = None,
    note: Optional[str] = None,
) -> Dict[str, Any]:
    """Помечает заключение как удалённое без физического удаления данных."""
    result = {
        "db_marked": False,
        "archive_marked": False,
        "already_deleted": False,
    }

    completion_id = conclusion.get("id")
    if not completion_id:
        return result

    timestamp = _now_iso()

    if _is_db_ready():
        async with db_lock:
            try:
                cursor = await db.execute(
                    "UPDATE completions SET is_deleted = 1, deleted_at = ?, deleted_by = ?, deletion_note = ? "
                    "WHERE id = ? AND (is_deleted IS NULL OR is_deleted = 0)",
                    (timestamp, initiator_id, note, completion_id),
                )
                await db.commit()
                if cursor.rowcount and cursor.rowcount > 0:
                    result["db_marked"] = True
                else:
                    result["already_deleted"] = True
            except Exception as error:
                logger.error(f"Не удалось пометить запись {completion_id} как удалённую: {error}")

    archive_path = conclusion.get("archive_path")
    if archive_path:
        archive_info = await set_archive_entry_status(archive_path, deleted=True, initiator_id=initiator_id, note=note)
        result["archive_marked"] = archive_info.get("updated", False)

    user_id = conclusion.get("user_id")
    if result["db_marked"] and user_id:
        await refresh_achievements_for_user(user_id)

    return result


async def restore_completion_record(completion_id: int, restorer_id: Optional[int] = None) -> Dict[str, Any]:
    result = {
        "restored": False,
        "reason": None,
        "record": None,
    }
    record = await fetch_completion_by_id(completion_id)
    if not record:
        result["reason"] = "not_found"
        return result
    if not record.get("is_deleted"):
        result["reason"] = "not_deleted"
        return result

    if _is_db_ready():
        async with db_lock:
            try:
                await db.execute(
                    "UPDATE completions SET is_deleted = 0, deleted_at = NULL, deleted_by = NULL, deletion_note = NULL "
                    "WHERE id = ?",
                    (completion_id,),
                )
                await db.commit()
                result["restored"] = True
            except Exception as error:
                logger.error(f"Не удалось восстановить запись {completion_id}: {error}")
                result["reason"] = "db_error"
                return result

    if result["restored"]:
        record = await fetch_completion_by_id(completion_id) or record
        result["record"] = record

    archive_path = record.get("archive_path")
    if archive_path:
        await restore_archived_document(archive_path, restorer_id)

    user_id = record.get("user_id")
    if result["restored"] and user_id:
        await refresh_achievements_for_user(user_id)

    return result

async def send_personal_stats(bot, user_id: int) -> None:
    message = await build_personal_stats_message(user_id)
    if not message:
        return
    await safe_bot_send_message(bot, user_id, message, skip_notice_on_retry=True)

# -------------------- Утилиты --------------------
def generate_unique_filename() -> str:
    return ''.join(random.choices(string.ascii_letters + string.digits, k=16)) + ".jpg"

def sanitize_filename(filename: str) -> str:
    """Очищает имя файла от запрещённых символов, имён и ограничивает длину."""
    cleaned = re.sub(r'[\/:*?"<>|]', '_', filename)
    reserved_names = {"CON", "PRN", "AUX", "NUL", "COM1", "COM2", "COM3", "COM4", "COM5", "COM6", "COM7", "COM8", "COM9", "LPT1", "LPT2", "LPT3", "LPT4", "LPT5", "LPT6", "LPT7", "LPT8", "LPT9"}
    if cleaned.upper() in reserved_names:
        cleaned = f"_{cleaned}_"
    return cleaned[:150]


def format_number(value: Any) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        numeric = 0.0
    if numeric.is_integer():
        numeric = int(numeric)
    else:
        numeric = int(round(numeric))
    return f"{numeric:,}".replace(",", " ")


def format_filesize(size_bytes: Optional[int]) -> str:
    try:
        size = float(size_bytes or 0)
    except (TypeError, ValueError):
        size = 0.0
    units = ["Б", "КБ", "МБ", "ГБ"]
    for unit in units:
        if size < 1024 or unit == units[-1]:
            if unit == "Б":
                return f"{int(size)} {unit}"
            return f"{size:.1f} {unit}"
        size /= 1024.0
    return f"{size:.1f} ГБ"


def get_state_label(state_name: str) -> str:
    if not state_name:
        return ""
    return STATE_LABELS.get(
        state_name,
        state_name.title().replace("_", " "),
    )


def level_profile_from_xp(xp: int) -> Dict[str, Any]:
    xp = max(0, int(xp or 0))
    level = 0
    xp_next = LEVEL_CATALOG[0]["xp_required"] if LEVEL_CATALOG else 0
    label = "🌱 Новичок"
    for entry in LEVEL_CATALOG:
        if xp >= entry["xp_required"]:
            level = entry["level"]
            label = entry["label"]
            xp_next = entry["xp_required"]
        else:
            xp_next = entry["xp_required"]
            break
    if level >= LEVEL_CATALOG[-1]["level"]:
        xp_next = LEVEL_CATALOG[-1]["xp_required"]
    xp_previous = 0
    if level > 0:
        xp_previous = LEVEL_CATALOG[level - 1]["xp_required"] if level <= len(LEVEL_CATALOG) else LEVEL_CATALOG[-1]["xp_required"]
    xp_to_next = max(0, xp_next - xp) if level < LEVEL_CATALOG[-1]["level"] else 0
    return {
        "level": level,
        "label": label,
        "xp_previous": xp_previous,
        "xp_next": xp_next,
        "xp_to_next": xp_to_next,
    }


def calculate_completion_xp(item_count: int, total_evaluation: float) -> int:
    base = 90
    item_bonus = item_count * 25
    value_bonus = int(min(max(total_evaluation, 0.0), 500_000) // 750)
    streak_bonus = 15 if item_count >= 3 else 0
    return max(80, base + item_bonus + value_bonus + streak_bonus)

def is_digit(value: str) -> bool:
    return value.isdigit()

def is_valid_ticket_number(value: str) -> bool:
    return value.isdigit() and MIN_TICKET_DIGITS <= len(value) <= MAX_TICKET_DIGITS

def match_region_name(text: str) -> Optional[str]:
    cleaned = (text or "").strip().lower()
    for region in REGION_TOPICS.keys():
        if region.lower() == cleaned:
            return region
    return None

def normalize_region_input(text: str) -> Optional[str]:
    if not text:
        return None
    cleaned = text.strip()
    if cleaned.startswith("🌍"):
        parts = cleaned.split(" ", 1)
        if len(parts) > 1:
            cleaned = parts[1]
    matched = match_region_name(cleaned)
    if matched:
        return matched
    return cleaned if cleaned in REGION_TOPICS else None

def parse_date_str(date_text: str) -> Optional[datetime]:
    try:
        return datetime.strptime(date_text, "%d.%m.%Y")
    except (ValueError, TypeError):
        return None

def get_month_bounds(month_text: str) -> Optional[tuple[datetime, datetime]]:
    try:
        month_date = datetime.strptime(month_text, "%m.%Y")
    except ValueError:
        return None
    last_day = monthrange(month_date.year, month_date.month)[1]
    start = month_date.replace(day=1)
    end = month_date.replace(day=last_day)
    return start, end


async def filter_records(start_date: Optional[datetime] = None, end_date: Optional[datetime] = None, region: Optional[str] = None) -> List[List[Any]]:
    records = await read_excel_data()
    if not records:
        return []

    filtered = []
    for row in records:
        date_obj = parse_date_str(row[3])
        if start_date and (not date_obj or date_obj < start_date):
            continue
        if end_date and (not date_obj or date_obj > end_date):
            continue
        if region and (row[4] or "").strip() != region:
            continue
        filtered.append(row)
    return filtered

def is_image_too_large(image_path: Path, max_size_mb: int = 5) -> bool:
    file_size_mb = image_path.stat().st_size / (1024 * 1024)
    return file_size_mb > max_size_mb

def compress_image(input_path: Path, output_path: Path, quality: int = 70) -> None:
    """Надёжно сжимает изображение, исправляя ориентацию и конвертируя в RGB."""
    with Image.open(input_path) as img:
        img = ImageOps.exif_transpose(img)
        if img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGB")
        img.save(output_path, "JPEG", quality=quality, optimize=True)

def clean_temp_files(max_age_seconds: int = 3600) -> None:
    """Удаляет устаревшие временные файлы."""
    if TEMP_PHOTOS_DIR.exists():
        now = time.time()
        for file in TEMP_PHOTOS_DIR.iterdir():
            if not file.is_file():
                continue
            if file.stat().st_mtime < now - max_age_seconds:
                try:
                    file.unlink()
                    logger.info(f"Удалён временный файл: {file.name}")
                except Exception as e:
                    logger.error(f"Ошибка при удалении файла {file.name}: {e}")

async def clean_temp_files_job(context: CallbackContext):
    """Асинхронная обёртка для вызова из JobQueue."""
    logger.info("Запуск периодической очистки временных файлов...")
    await asyncio.to_thread(clean_temp_files, 3600)


def cleanup_user_photos(data: Union[ConclusionData, Dict[str, Any]]) -> None:
    """Удаляет временные фото пользователя."""
    if isinstance(data, ConclusionData):
        photos = data.photo_desc
    else:
        photos = data.get('photo_desc', []) if data else []

    for item in photos:
        photo_path = Path(item.get('photo', ""))
        if photo_path.is_file():
            try:
                photo_path.unlink()
                logger.info(f"Удалено временное фото: {photo_path.name}")
            except OSError as e:
                logger.warning(f"Не удалось удалить временное фото {photo_path}: {e}")
            parent_dir = photo_path.parent
            if parent_dir != TEMP_PHOTOS_DIR and parent_dir.is_dir():
                try:
                    if not any(parent_dir.iterdir()):
                        parent_dir.rmdir()
                        logger.info(f"Удалена пустая папка {parent_dir}")
                except OSError as cleanup_error:
                    logger.debug(f"Не удалось удалить папку {parent_dir}: {cleanup_error}")


def format_progress(stage: str) -> str:
    step = PROGRESS_STEPS.get(stage)
    if not step:
        return PROGRESS_SYMBOL_EMPTY * PROGRESS_BAR_SEGMENTS
    filled = math.ceil(step / TOTAL_STEPS * PROGRESS_BAR_SEGMENTS)
    filled = max(0, min(PROGRESS_BAR_SEGMENTS, filled))
    empty = PROGRESS_BAR_SEGMENTS - filled
    return f"{PROGRESS_SYMBOL_FILLED * filled}{PROGRESS_SYMBOL_EMPTY * empty}"


def ticket_digits_hint() -> str:
    if MIN_TICKET_DIGITS == MAX_TICKET_DIGITS:
        return str(MIN_TICKET_DIGITS)
    return f"{MIN_TICKET_DIGITS}-{MAX_TICKET_DIGITS}"


def ticket_digits_phrase() -> str:
    if MIN_TICKET_DIGITS == MAX_TICKET_DIGITS:
        return f"ровно {MIN_TICKET_DIGITS} цифр"
    return f"от {MIN_TICKET_DIGITS} до {MAX_TICKET_DIGITS} цифр"


EDIT_FIELD_TARGETS: Dict[str, Dict[str, Any]] = {
    "department": {
        "state": DialogState.DEPARTMENT,
        "stage": "department",
        "prompt": "Введите номер подразделения (например 385).",
    },
    "issue": {
        "state": DialogState.ISSUE_NUMBER,
        "stage": "issue",
        "prompt": "Введите порядковый номер заключения за день (например 1).",
    },
    "ticket": {
        "state": DialogState.TICKET_NUMBER,
        "stage": "ticket",
        "prompt": f"Введите номер залогового билета (нужно {ticket_digits_phrase()}).",
    },
    "date": {
        "state": DialogState.DATE,
        "stage": "date",
        "prompt": "Введите дату заключения в формате 01.03.2025.",
    },
    "region": {
        "state": DialogState.REGION,
        "stage": "region",
        "prompt": "Выберите регион из списка:",
    },
}

STATE_LABELS: Dict[str, str] = {
    DialogState.DEPARTMENT.name: "Подразделение",
    DialogState.ISSUE_NUMBER.name: "№ заключения",
    DialogState.TICKET_NUMBER.name: "Билет",
    DialogState.DATE.name: "Дата",
    DialogState.REGION.name: "Регион",
    DialogState.PHOTO.name: "Фото",
    DialogState.DESCRIPTION.name: "Описание",
    DialogState.EVALUATION.name: "Оценка",
    DialogState.MORE_PHOTO.name: "Доп. фото",
    DialogState.CONFIRMATION.name: "Подтверждение",
    DialogState.TESTING.name: "Выбор режима",
}


def get_state_label(state_name: str) -> str:
    if not state_name:
        return ""
    return STATE_LABELS.get(state_name, state_name.title().replace("_", " "))


def extract_state_value(data: ConclusionData, state: DialogState) -> str:
    """Return the currently stored value for a given dialog state."""
    if data is None:
        return ""
    if state == DialogState.DEPARTMENT:
        return data.department_number
    if state == DialogState.ISSUE_NUMBER:
        return data.issue_number
    if state == DialogState.TICKET_NUMBER:
        return data.ticket_number
    if state == DialogState.DATE:
        return data.date
    if state == DialogState.REGION:
        return data.region
    return ""


def format_duration(seconds: Optional[float]) -> str:
    try:
        total_seconds = int(round(float(seconds or 0.0)))
    except (TypeError, ValueError):
        total_seconds = 0
    if total_seconds < 60:
        return f"{total_seconds} с"
    minutes, sec = divmod(total_seconds, 60)
    if minutes < 60:
        return f"{minutes} мин {sec} с"
    hours, minutes = divmod(minutes, 60)
    return f"{hours} ч {minutes} мин"


def validate_department(value: str) -> ValidationResult:
    cleaned = (value or "").strip()
    if not cleaned:
        return ValidationResult(False, "Введите номер подразделения — поле не может быть пустым.", "Например: 385")
    if not cleaned.isdigit():
        return ValidationResult(False, "Номер подразделения должен содержать только цифры.", "Используйте клавиатуру: 385")
    if len(cleaned) > 6:
        return ValidationResult(False, "Слишком длинный номер подразделения.", "Максимум 6 цифр, например 123456")
    return ValidationResult(True)


def validate_issue_number(value: str) -> ValidationResult:
    cleaned = (value or "").strip()
    if not cleaned:
        return ValidationResult(False, "Укажите порядковый номер заключения за день.", "Например: 1")
    if not cleaned.isdigit():
        return ValidationResult(False, "Номер заключения должен содержать только цифры.", "Например: 1 или 12")
    if len(cleaned) > 4:
        return ValidationResult(False, "Слишком длинный номер заключения.", "Используйте до 4 цифр, например 1234")
    return ValidationResult(True)


def validate_ticket_number(value: str) -> ValidationResult:
    cleaned = (value or "").strip()
    if not cleaned:
        return ValidationResult(False, "Номер билета отсутствует.", f"Введите {ticket_digits_phrase()} — например 01234567890")
    if not cleaned.isdigit():
        return ValidationResult(False, "В номере билета допустимы только цифры.", f"Используйте {ticket_digits_phrase()} — например 01234567890")
    if not (MIN_TICKET_DIGITS <= len(cleaned) <= MAX_TICKET_DIGITS):
        return ValidationResult(False, f"Номер билета содержит {len(cleaned)} цифр.", f"Нужно {ticket_digits_phrase()} — например 01234567890")
    return ValidationResult(True)


def validate_date_text(value: str) -> ValidationResult:
    cleaned = (value or "").strip()
    if not cleaned:
        return ValidationResult(False, "Дата заключения не указана.", "Формат: 01.03.2025")
    try:
        date_obj = datetime.strptime(cleaned, "%d.%m.%Y")
    except ValueError:
        return ValidationResult(False, "Дата не распознана.", "Используйте формат ДД.ММ.ГГГГ — например 01.03.2025")
    if date_obj > datetime.now() + timedelta(days=1):
        return ValidationResult(False, "Дата выглядит будущей.", "Проверьте календарь и корректность года.")
    return ValidationResult(True)


def validate_evaluation(value: str) -> ValidationResult:
    cleaned = (value or "").strip().replace(" ", "")
    if not cleaned:
        return ValidationResult(False, "Оценка не указана.", "Введите сумму в рублях, например 1500")
    if not cleaned.isdigit():
        return ValidationResult(False, "Оценка должна быть числом без символов.", "Пример: 1500")
    amount = int(cleaned)
    if amount <= 0:
        return ValidationResult(False, "Оценка должна быть положительной.", "Минимум 1 рубль.")
    if amount > 5_000_000:
        return ValidationResult(False, "Оценка превышает лимит 5 000 000 руб.", "Проверьте значение или обсудите с руководителем.")
    return ValidationResult(True)


async def emit_validation_error(update: Update, context: CallbackContext, result: ValidationResult, stage: str, keyboard=None) -> int:
    hint = f"\n{result.hint}" if result.hint else ""
    message = f"{format_progress(stage)} {result.message}{hint}"
    reply_markup = keyboard if keyboard is not None else build_step_inline_keyboard(context=context)
    await safe_reply(update, message, reply_markup=reply_markup)
    return -1


def ensure_menu_button(rows: List[List[str]]) -> List[List[str]]:
    has_menu = any(MENU_BUTTON_LABEL in row for row in rows)
    new_rows = [list(row) for row in rows]
    if not has_menu:
        new_rows.append([MENU_BUTTON_LABEL])
    return new_rows


def build_keyboard(rows: List[List[str]], one_time: bool = False) -> ReplyKeyboardMarkup:
    """Создаёт клавиатуру с крупными кнопками (без уменьшения размера)."""
    button_rows = [[KeyboardButton(text=label) for label in row] for row in rows]
    return ReplyKeyboardMarkup(button_rows, one_time_keyboard=one_time, resize_keyboard=False)


def build_keyboard_with_menu(rows: List[List[str]], one_time: bool = False) -> ReplyKeyboardMarkup:
    return build_keyboard(ensure_menu_button(rows), one_time=one_time)


def build_main_menu(user_id: int) -> ReplyKeyboardMarkup:
    rows = ADMIN_MAIN_MENU_ROWS if is_admin(user_id) else USER_MAIN_MENU_ROWS
    return build_keyboard_with_menu(rows, one_time=False)


def get_state_stack(context: CallbackContext) -> List[DialogState]:
    return context.user_data.setdefault("state_stack", [])


def clear_state_stack(context: CallbackContext) -> None:
    context.user_data.pop("state_stack", None)


def push_state(context: CallbackContext, state: DialogState) -> None:
    stack = get_state_stack(context)
    if stack and stack[-1] == state:
        return
    stack.append(state)


def pop_state(context: CallbackContext) -> Optional[DialogState]:
    stack = context.user_data.get("state_stack")
    if stack:
        return stack.pop()
    return None


def has_previous_state(context: CallbackContext) -> bool:
    stack = context.user_data.get("state_stack")
    return bool(stack)


def set_current_state(context: CallbackContext, state: DialogState) -> None:
    context.user_data["current_state"] = state


def get_current_state(context: CallbackContext) -> Optional[DialogState]:
    state = context.user_data.get("current_state")
    return state if isinstance(state, DialogState) else None


def clear_pending_items(context: CallbackContext) -> None:
    """Удаляет вспомогательные очереди/указатели для незавершённых пунктов."""
    context.user_data.pop("pending_item_queue", None)
    context.user_data.pop("current_item_index", None)


def set_resume_state(context: CallbackContext, state: Optional[DialogState]) -> None:
    if state is None:
        context.user_data.pop("resume_state", None)
    else:
        context.user_data["resume_state"] = state


def get_resume_state(context: CallbackContext) -> Optional[DialogState]:
    state = context.user_data.get("resume_state")
    return state if isinstance(state, DialogState) else None


def build_step_inline_keyboard(
    button_rows: Optional[List[List[InlineKeyboardButton]]] = None,
    *,
    include_cancel: bool = True,
    include_menu: bool = False,
    include_back: Optional[bool] = None,
    context: Optional[CallbackContext] = None,
) -> InlineKeyboardMarkup:
    rows: List[List[InlineKeyboardButton]] = []
    if button_rows:
        rows.extend([list(row) for row in button_rows])
    navigation_row: List[InlineKeyboardButton] = []
    if include_back is None and context is not None:
        include_back = has_previous_state(context)
    if include_back:
        navigation_row.append(
            InlineKeyboardButton("⬅️ Назад", callback_data=f"{NAVIGATION_CALLBACK_PREFIX}back")
        )
    if include_menu:
        navigation_row.append(
            InlineKeyboardButton("📋 Меню", callback_data=f"{NAVIGATION_CALLBACK_PREFIX}menu")
        )
    if include_cancel:
        navigation_row.append(
            InlineKeyboardButton("❌ Отмена", callback_data=f"{NAVIGATION_CALLBACK_PREFIX}cancel")
        )
    if navigation_row:
        rows.append(navigation_row)
    if not rows:
        rows.append(
            [InlineKeyboardButton("❌ Отмена", callback_data=f"{NAVIGATION_CALLBACK_PREFIX}cancel")]
        )
    return InlineKeyboardMarkup(rows)


def _format_human_dt(value: Optional[str]) -> str:
    if not value:
        return "—"
    try:
        dt = datetime.fromisoformat(value)
    except ValueError:
        return value
    return dt.strftime("%d.%m.%Y %H:%M")


def build_draft_keyboard(discard_label: str = "Удалить черновик") -> InlineKeyboardMarkup:
    rows = [
        [InlineKeyboardButton("Продолжить черновик", callback_data=f"{DRAFT_CALLBACK_PREFIX}resume")],
        [InlineKeyboardButton(discard_label, callback_data=f"{DRAFT_CALLBACK_PREFIX}discard")],
    ]
    return InlineKeyboardMarkup(rows)


def build_draft_summary_text(draft: Dict[str, Any]) -> str:
    conclusion: Optional[ConclusionData] = draft.get("conclusion")
    updated_label = _format_human_dt(draft.get("updated_at"))
    created_label = _format_human_dt(draft.get("created_at"))
    lines = [
        "💾 Черновик сохранён.",
        f"Создан: {created_label}",
        f"Обновлён: {updated_label}",
    ]
    if conclusion:
        lines.append("")
        lines.append(f"Подразделение: {conclusion.department_number or '—'}")
        lines.append(f"№ заключения: {conclusion.issue_number or '—'}")
        lines.append(f"Билет: {conclusion.ticket_number or '—'}")
        lines.append(f"Дата: {conclusion.date or '—'}")
        lines.append(f"Регион: {conclusion.region or '—'}")
        lines.append(f"Предметов: {len(conclusion.photo_desc)}")
    return "\n".join(lines)


def _draft_discard_label(context: CallbackContext) -> str:
    return context.user_data.get("draft_discard_label") or "Удалить черновик"


async def remind_draft_decision(target: Any, context: CallbackContext) -> None:
    draft: Optional[Dict[str, Any]] = context.user_data.get("pending_draft")
    if draft:
        text = build_draft_summary_text(draft) + "\n\nВыберите действие с черновиком:"
    else:
        text = "Выберите действие с черновиком."
    await _send_via_target(
        target,
        context,
        text,
        reply_markup=build_draft_keyboard(discard_label=_draft_discard_label(context))
    )


def truncate_text(value: Any, limit: int = 60) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 1)].rstrip() + "…"


def _paginate(items: Sequence[Any], page: int, page_size: int) -> Tuple[List[Any], int, int]:
    total_items = len(items)
    if total_items == 0:
        return [], 1, 1
    total_pages = max(1, math.ceil(total_items / page_size))
    current_page = max(1, min(page, total_pages))
    start = (current_page - 1) * page_size
    end = start + page_size
    return list(items[start:end]), current_page, total_pages


def build_history_page(records: List[List[Any]], page: int, page_size: int) -> Tuple[str, InlineKeyboardMarkup, int]:
    entries, current_page, total_pages = _paginate(records, page, page_size)
    if not records:
        return (
            "📜 История заключений пуста.",
            InlineKeyboardMarkup([[InlineKeyboardButton("Закрыть", callback_data=f"{HISTORY_CALLBACK_PREFIX}close")]]),
            current_page,
        )

    lines = [
        f"📜 История заключений — страница {current_page}/{total_pages}",
        f"Всего записей: {len(records)}",
        "",
    ]
    if not entries:
        lines.append("Записей на этой странице нет.")
    else:
        for row in entries:
            ticket = (row[0] or "—") if len(row) > 0 else "—"
            issue_number = (row[1] or "—") if len(row) > 1 else "—"
            department = (row[2] or "—") if len(row) > 2 else "—"
            date_label = (row[3] or "—") if len(row) > 3 else "—"
            region = (row[4] or "—") if len(row) > 4 else "—"
            item_number = row[5] if len(row) > 5 else ""
            description = truncate_text(row[6] if len(row) > 6 else "Без описания", 70)
            evaluation = row[7] if len(row) > 7 else ""
            try:
                evaluation_value = format_number(evaluation)
            except Exception:
                evaluation_value = str(evaluation or "0")
            item_label = f"предмет {int(item_number)}" if isinstance(item_number, (int, float)) else "предмет"
            lines.append(
                f"• {date_label} • {region} • {item_label}\n  "
                f"Подразделение {department}, билет {ticket}, № {issue_number}\n  "
                f"Оценка: {evaluation_value} руб.\n  {description}"
            )
            lines.append("")

    nav_buttons: List[InlineKeyboardButton] = []
    if current_page > 1:
        nav_buttons.append(
            InlineKeyboardButton("⬅️ Назад", callback_data=f"{HISTORY_CALLBACK_PREFIX}{current_page - 1}")
        )
    nav_buttons.append(
        InlineKeyboardButton(f"{current_page}/{total_pages}", callback_data=f"{HISTORY_CALLBACK_PREFIX}{current_page}")
    )
    if current_page < total_pages:
        nav_buttons.append(
            InlineKeyboardButton("Вперёд ➡️", callback_data=f"{HISTORY_CALLBACK_PREFIX}{current_page + 1}")
        )

    keyboard_rows = [nav_buttons] if nav_buttons else []
    keyboard_rows.append([InlineKeyboardButton("Закрыть", callback_data=f"{HISTORY_CALLBACK_PREFIX}close")])
    keyboard = InlineKeyboardMarkup(keyboard_rows)
    return "\n".join(lines).strip(), keyboard, current_page


def _metrics_container(context: CallbackContext) -> Dict[str, Any]:
    metrics = context.user_data.setdefault(
        "step_metrics",
        {
            "started_at": time.time(),
            "started_at_iso": _now_iso(),
            "durations": {},
            "sequence": [],
            "active_state": None,
            "state_started_at": None,
            "finished_at": None,
            "finished_at_epoch": None,
        },
    )
    return metrics


def metrics_reset(context: CallbackContext) -> None:
    context.user_data["step_metrics"] = {
        "started_at": time.time(),
        "started_at_iso": _now_iso(),
        "durations": {},
        "sequence": [],
        "active_state": None,
        "state_started_at": None,
        "finished_at": None,
        "finished_at_epoch": None,
    }


def metrics_enter_state(context: CallbackContext, state: DialogState) -> None:
    metrics = _metrics_container(context)
    now_ts = time.time()
    metrics["active_state"] = state.name
    metrics["state_started_at"] = now_ts
    metrics.setdefault("sequence", []).append(
        {"state": state.name, "event": "start", "timestamp": _now_iso()}
    )


def metrics_complete_state(context: CallbackContext, state: DialogState) -> None:
    metrics = _metrics_container(context)
    start_at = metrics.get("state_started_at")
    active_state = metrics.get("active_state")
    if start_at is None or active_state != state.name:
        return
    duration = max(0.0, time.time() - start_at)
    durations = metrics.setdefault("durations", {})
    durations[state.name] = durations.get(state.name, 0.0) + duration
    metrics.setdefault("sequence", []).append(
        {
            "state": state.name,
            "event": "finish",
            "timestamp": _now_iso(),
            "duration": duration,
        }
    )
    metrics["active_state"] = None
    metrics["state_started_at"] = None


def metrics_snapshot(context: CallbackContext) -> Dict[str, Any]:
    metrics = _metrics_container(context)
    snapshot = {
        "started_at": metrics.get("started_at_iso"),
        "started_at_epoch": metrics.get("started_at"),
        "durations": metrics.get("durations", {}).copy(),
        "sequence": copy.deepcopy(metrics.get("sequence", [])),
        "active_state": metrics.get("active_state"),
        "finished_at": metrics.get("finished_at"),
        "finished_at_epoch": metrics.get("finished_at_epoch"),
    }
    return snapshot


def metrics_processing_time(metrics: Optional[Dict[str, Any]] = None) -> float:
    if not metrics:
        return 0.0
    started_epoch = metrics.get("started_at_epoch")
    if not started_epoch:
        return 0.0
    finished = metrics.get("finished_at_epoch")
    if finished is None:
        finished = time.time()
    return max(0.0, float(finished) - float(started_epoch))


def metrics_finalize(context: CallbackContext) -> Dict[str, Any]:
    metrics = _metrics_container(context)
    if metrics.get("active_state") and metrics.get("state_started_at") is not None:
        state_name = metrics["active_state"]
        try:
            state = DialogState[state_name]
        except KeyError:
            state = None
        if state:
            metrics_complete_state(context, state)
    metrics["finished_at_epoch"] = time.time()
    metrics["finished_at"] = _now_iso()
    return metrics_snapshot(context)


def metrics_restore(context: CallbackContext, data: Optional[Dict[str, Any]], resume_state: Optional[DialogState]) -> None:
    metrics_reset(context)
    if not data:
        return
    metrics = _metrics_container(context)
    metrics["started_at"] = data.get("started_at_epoch", time.time())
    metrics["started_at_iso"] = data.get("started_at", _now_iso())
    metrics["durations"] = data.get("durations", {}) or {}
    metrics["sequence"] = copy.deepcopy(data.get("sequence", [])) or []
    if resume_state:
        metrics.setdefault("sequence", []).append(
            {
                "state": resume_state.name,
                "event": "resume",
                "timestamp": _now_iso(),
            }
        )
    metrics["active_state"] = None
    metrics["state_started_at"] = None


def _resolve_chat_context(target: Any) -> tuple[Optional[Any], Optional[int], Optional[int], Optional[int]]:
    if isinstance(target, Update):
        message = target.effective_message
        from_user = target.effective_user
    else:
        message = getattr(target, "message", None)
        from_user = getattr(target, "from_user", None)
    if message is None:
        return None, None, None, from_user.id if from_user else None
    chat_id = getattr(message, "chat_id", None)
    thread_id = getattr(message, "message_thread_id", None)
    user_id = from_user.id if from_user else None
    return message, chat_id, thread_id, user_id


async def _send_via_target(target: Any, context: CallbackContext, text: str, *, reply_markup=None):
    message, chat_id, thread_id, _ = _resolve_chat_context(target)
    if message is None or chat_id is None:
        return None
    if isinstance(target, Update):
        return await safe_reply(target, text, reply_markup=reply_markup)
    else:
        return await safe_bot_send_message(
            context.bot,
            chat_id,
            text,
            reply_markup=reply_markup,
            message_thread_id=thread_id,
        )


def build_region_inline_keyboard(
    include_all: bool = False,
    prefix: str = REGION_CALLBACK_PREFIX,
    *,
    context: Optional[CallbackContext] = None,
) -> InlineKeyboardMarkup:
    buttons: List[List[InlineKeyboardButton]] = []
    row: List[InlineKeyboardButton] = []
    for idx, region_name in enumerate(REGION_CHOICES):
        row.append(InlineKeyboardButton(region_name, callback_data=f"{prefix}{idx}"))
        if len(row) == 2:
            buttons.append(row)
            row = []
    if row:
        buttons.append(row)
    if include_all:
        buttons.append([InlineKeyboardButton("Все регионы", callback_data=f"{prefix}all")])
    return build_step_inline_keyboard(buttons, context=context)


def build_reports_action_keyboard() -> InlineKeyboardMarkup:
    rows = [
        [InlineKeyboardButton("📦 Архив заключений", callback_data=f"{REPORT_ACTION_PREFIX}archive")],
        [InlineKeyboardButton("📥 Выгрузка за месяц", callback_data=f"{REPORT_ACTION_PREFIX}month")],
        [InlineKeyboardButton("📈 Статистика за период", callback_data=f"{REPORT_ACTION_PREFIX}period")],
        [InlineKeyboardButton("📊 Сводка по регионам", callback_data=f"{REPORT_ACTION_PREFIX}summary")],
        [InlineKeyboardButton("Отмена", callback_data=f"{REPORT_ACTION_PREFIX}cancel")],
    ]
    return build_step_inline_keyboard(rows)


def build_region_filter_keyboard(include_all: bool = True) -> ReplyKeyboardMarkup:
    rows: List[List[str]] = []
    if include_all:
        rows.append(["🌐 Все регионы"])
    rows.extend([[f"🌍 {region}"] for region in REGION_TOPICS.keys()])
    rows.append(["❌ Отмена"])
    return build_keyboard_with_menu(rows, one_time=True)


def build_void_keyboard(completion_id: int) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton(
                text="♻️ Удалить из отчётов",
                callback_data=f"{VOID_CALLBACK_PREFIX}{completion_id}"
            )
        ]
    ])


def build_confirmation_keyboard(data: ConclusionData, include_back: bool = False) -> InlineKeyboardMarkup:
    rows: List[List[InlineKeyboardButton]] = [
        [InlineKeyboardButton("Подразделение", callback_data=f"{EDIT_CALLBACK_PREFIX}field:department")],
        [InlineKeyboardButton("№ заключения", callback_data=f"{EDIT_CALLBACK_PREFIX}field:issue")],
        [InlineKeyboardButton("Залоговый билет", callback_data=f"{EDIT_CALLBACK_PREFIX}field:ticket")],
        [InlineKeyboardButton("Дата", callback_data=f"{EDIT_CALLBACK_PREFIX}field:date")],
        [InlineKeyboardButton("Регион", callback_data=f"{EDIT_CALLBACK_PREFIX}field:region")],
    ]
    if data.photo_desc:
        rows.append([InlineKeyboardButton("Удалить последнее фото", callback_data=f"{EDIT_CALLBACK_PREFIX}delete_photo")])
    rows.append([InlineKeyboardButton("Продолжить →", callback_data=f"{CONFIRM_CALLBACK_PREFIX}next")])

    action_row: List[InlineKeyboardButton] = []
    if include_back:
        action_row.append(InlineKeyboardButton("⬅️ Назад", callback_data=f"{NAVIGATION_CALLBACK_PREFIX}back"))
    action_row.append(InlineKeyboardButton("Отменить", callback_data=f"{CONFIRM_CALLBACK_PREFIX}cancel"))
    rows.append(action_row)
    return InlineKeyboardMarkup(rows)


def build_mode_keyboard() -> InlineKeyboardMarkup:
    rows = [
        [
            InlineKeyboardButton("Тестовое", callback_data=f"{MODE_CALLBACK_PREFIX}test"),
            InlineKeyboardButton("Окончательное", callback_data=f"{MODE_CALLBACK_PREFIX}final"),
        ]
    ]
    return build_step_inline_keyboard(rows)


async def send_preview_photos(
    update: Optional[Update],
    data: Union[ConclusionData, Dict[str, Any]],
    max_items: int = PREVIEW_MAX_ITEMS,
    bot=None,
    chat_id: Optional[int] = None,
    thread_id: Optional[int] = None,
) -> None:
    """Отправляет медиагруппу с последними фотографиями для визуального подтверждения."""
    if isinstance(data, ConclusionData):
        photos = data.photo_desc
    else:
        photos = data.get('photo_desc', []) if data else []
    available = []
    for item in photos:
        photo_path = Path(item.get('photo', ""))
        if photo_path.is_file():
            available.append((photo_path, item))
    if not available:
        return
    selected = available[-max_items:]
    media_items: List[InputMediaPhoto] = []
    opened_files = []
    for path, item in selected:
        try:
            file = path.open("rb")
        except OSError as err:
            logger.warning(f"Не удалось открыть файл для превью {path}: {err}")
            continue
        opened_files.append(file)
        caption_lines = [
            item.get('description', 'Без описания'),
            f"💰 {item.get('evaluation', 'Нет оценки')} руб."
        ]
        caption = "\n".join(line for line in caption_lines if line)
        media_items.append(InputMediaPhoto(file, caption=caption or None))
    if not media_items:
        for file in opened_files:
            try:
                file.close()
            except Exception:
                pass
        return
    try:
        if update and update.message:
            await update.message.reply_media_group(media_items)
        else:
            target_bot = bot or (update.get_bot() if update else None)
            if target_bot and chat_id is not None:
                await target_bot.send_media_group(
                    chat_id=chat_id,
                    media=media_items,
                    message_thread_id=thread_id,
                )
    except TelegramError as e:
        logger.warning(f"Не удалось отправить превью фотографий: {e}")
    finally:
        for file in opened_files:
            try:
                file.close()
            except Exception:
                pass


async def create_excel_snapshot(rows: List[List[Any]], filename_prefix: str) -> Path:
    """Создаёт временный Excel-файл с переданными строками и возвращает путь."""
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%d.%m.%Y_%H-%M-%S")

    def _write_snapshot() -> Path:
        wb = Workbook()
        ws = wb.active
        ws.append(EXCEL_HEADERS)
        for row in rows:
            ws.append(row)
        raw_name = f"{filename_prefix}_{timestamp}.xlsx"
        filepath = DOCS_DIR / sanitize_filename(raw_name)
        wb.save(filepath)
        wb.close()
        return filepath

    return await asyncio.to_thread(_write_snapshot)


def _read_archive_index() -> List[Dict[str, Any]]:
    if not ARCHIVE_INDEX_FILE.exists():
        return []
    try:
        with ARCHIVE_INDEX_FILE.open("r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError) as e:
        logger.warning(f"Не удалось прочитать индекс архива: {e}")
        return []


def _write_archive_index(entries: List[Dict[str, Any]]) -> None:
    ARCHIVE_INDEX_FILE.parent.mkdir(parents=True, exist_ok=True)
    with ARCHIVE_INDEX_FILE.open("w", encoding="utf-8") as f:
        json.dump(entries, f, ensure_ascii=False, indent=2)


async def set_archive_entry_status(
    rel_path: str,
    *,
    deleted: bool,
    initiator_id: Optional[int] = None,
    note: Optional[str] = None,
) -> Dict[str, Any]:
    async with archive_lock:
        entries = await asyncio.to_thread(_read_archive_index)
        updated = False
        for entry in entries:
            if entry.get("archive_path") == rel_path:
                entry["is_deleted"] = bool(deleted)
                entry["deleted_at"] = _now_iso() if deleted else None
                entry["deleted_by"] = initiator_id if deleted else None
                entry["deletion_note"] = note if deleted else None
                updated = True
                break
        if updated:
            await asyncio.to_thread(_write_archive_index, entries)
    return {"updated": updated}


async def archive_document(filepath: Path, data: Union[ConclusionData, Dict[str, Any]]) -> Optional[Path]:
    if not filepath.is_file():
        return None

    if isinstance(data, ConclusionData):
        data_dict = data.to_dict()
    else:
        data_dict = data or {}

    date_text = data_dict.get("date")
    dt = parse_date_str(date_text)
    subdir_name = dt.strftime("%Y-%m") if dt else "undated"
    month_dir = ARCHIVE_DIR / subdir_name

    description = data_dict.get('photo_desc', [])

    def _copy_and_index() -> Optional[Path]:
        month_dir.mkdir(parents=True, exist_ok=True)
        target = month_dir / filepath.name
        counter = 1
        while target.exists():
            target = month_dir / f"{filepath.stem}_{counter}{filepath.suffix}"
            counter += 1
        shutil.copy2(filepath, target)

        entry = {
            "archive_path": str(target.relative_to(ARCHIVE_DIR)),
            "date": date_text,
            "department_number": data_dict.get("department_number"),
            "issue_number": data_dict.get("issue_number"),
            "ticket_number": data_dict.get("ticket_number"),
            "region": data_dict.get("region"),
            "items": description,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "is_deleted": False,
            "deleted_at": None,
            "deleted_by": None,
            "deletion_note": None,
        }

        entries = _read_archive_index()
        entries.append(entry)
        _write_archive_index(entries)
        return target

    async with archive_lock:
        return await asyncio.to_thread(_copy_and_index)


def _cleanup_archive_dirs(start_path: Path) -> None:
    """Удаляет пустые директории архива снизу вверх до корневой."""
    current = start_path
    try:
        while current != ARCHIVE_DIR and current.exists():
            if any(current.iterdir()):
                break
            current.rmdir()
            current = current.parent
    except Exception as error:
        logger.debug(f"Не удалось удалить пустую директорию архива {current}: {error}")


async def remove_archived_document(conclusion: Dict[str, Any], initiator_id: Optional[int] = None, note: Optional[str] = None) -> Dict[str, bool]:
    """Помечает архивный файл как удалённый (soft delete)."""
    rel_path = conclusion.get("archive_path")
    if not rel_path:
        return {"file_removed": False, "index_removed": False}

    result = await set_archive_entry_status(rel_path, deleted=True, initiator_id=initiator_id, note=note)
    return {"file_removed": False, "index_removed": result.get("updated", False)}


async def restore_archived_document(rel_path: Optional[str], restorer_id: Optional[int] = None) -> bool:
    if not rel_path:
        return False
    result = await set_archive_entry_status(rel_path, deleted=False, initiator_id=restorer_id, note=None)
    return result.get("updated", False)


async def get_archive_paths(start_date: datetime, end_date: datetime, region: Optional[str]) -> List[Path]:
    async with archive_lock:
        entries = await asyncio.to_thread(_read_archive_index)

    paths: List[Path] = []
    for entry in entries:
        entry_date = parse_date_str(entry.get("date"))
        if not entry_date:
            continue
        if entry_date < start_date or entry_date > end_date:
            continue
        entry_region = entry.get("region")
        if region and entry_region != region:
            continue
        if entry.get("is_deleted"):
            continue
        rel_path = entry.get("archive_path")
        if not rel_path:
            continue
        abs_path = ARCHIVE_DIR / rel_path
        if abs_path.is_file():
            paths.append(abs_path)
    return paths


async def send_month_archive(update: Update, context: CallbackContext, month_text: str, start_date: datetime, end_date: datetime, region: Optional[str]) -> None:
    archive_paths = await get_archive_paths(start_date, end_date, region)
    if not archive_paths:
        region_note = f" и регион {region}" if region else ""
        await safe_reply(update, f"За {month_text}{region_note} архивных заключений не найдено.")
        return

    await safe_chat_action(context.bot, update.effective_chat.id, ChatAction.UPLOAD_DOCUMENT)

    region_label = region or "Все регионы"
    timestamp = datetime.now().strftime("%d.%m.%Y_%H-%M-%S")

    def _create_zip() -> Path:
        DOCS_DIR.mkdir(parents=True, exist_ok=True)
        zip_name = sanitize_filename(f"archive_{month_text}_{region_label}_{timestamp}.zip")
        zip_path = DOCS_DIR / zip_name
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for path in archive_paths:
                zf.write(path, arcname=path.name)
        return zip_path

    zip_path = await asyncio.to_thread(_create_zip)
    try:
        caption = f"Архив заключений за {month_text} ({region_label})"
        await send_document_from_path(context.bot, chat_id=update.effective_chat.id, path=zip_path, caption=caption)
        await safe_reply(update, "📦 Архив с заключениями отправлен.")
    finally:
        try:
            if zip_path.exists():
                zip_path.unlink()
        except OSError as cleanup_error:
            logger.debug(f"Не удалось удалить временный архив {zip_path}: {cleanup_error}")


def _parse_search_filters(tokens: List[str]) -> Tuple[Dict[str, Any], List[str]]:
    filters: Dict[str, Any] = {}
    errors: List[str] = []
    if not tokens:
        return filters, errors

    joined = " ".join(tokens)
    pattern = re.compile(r"(\w+)=('([^']*)'|\"([^\"]*)\"|[^\s]+)")
    matches = list(pattern.finditer(joined))

    if matches:
        for match in matches:
            key = match.group(1).lower()
            raw_value = match.group(2) or ""
            value = raw_value.strip("'\"")
            filters[key] = value
    else:
        if tokens:
            filters["ticket"] = tokens[0]
        for token in tokens[1:]:
            if "=" in token:
                key, val = token.split("=", 1)
                filters[key.lower()] = val
            else:
                errors.append(f"Не удалось распознать аргумент: {token}")

    return filters, errors


def _describe_archive_entry(entry: Dict[str, Any], index: Optional[int] = None) -> str:
    parts: List[str] = []
    prefix = f"{index}. " if index is not None else ""
    parts.append(f"{prefix}{entry.get('date') or 'дата не указана'} • {entry.get('region') or 'Регион не указан'}")
    details: List[str] = []
    if entry.get("ticket_number"):
        details.append(f"Билет {entry['ticket_number']}")
    if entry.get("issue_number"):
        details.append(f"№ {entry['issue_number']}")
    if entry.get("department_number"):
        details.append(f"Подразделение {entry['department_number']}")
    if details:
        parts.append("  " + ", ".join(details))
    items = entry.get("items") or []
    if items:
        parts.append(f"  Предметов: {len(items)}")
    if entry.get("archive_path"):
        parts.append(f"  Файл: {entry['archive_path']}")
    return "\n".join(parts)


def _build_search_label(criteria: Dict[str, Any]) -> str:
    fragments: List[str] = []
    if criteria.get("ticket_number"):
        fragments.append(f"ticket_{criteria['ticket_number']}")
    if criteria.get("issue_number"):
        fragments.append(f"issue_{criteria['issue_number']}")
    if criteria.get("region"):
        fragments.append(f"region_{criteria['region']}")
    if criteria.get("date"):
        fragments.append(f"date_{criteria['date']}")
    start = criteria.get("date_from")
    end = criteria.get("date_to")
    if isinstance(start, datetime):
        start_label = start.strftime("%d.%m.%Y")
    else:
        start_label = start
    if isinstance(end, datetime):
        end_label = end.strftime("%d.%m.%Y")
    else:
        end_label = end
    if start_label or end_label:
        fragments.append(f"period_{start_label or 'start'}-{end_label or 'end'}")
    return "_".join(fragments) or "results"


async def _send_archive_search_results(
    update: Update,
    context: CallbackContext,
    matches: List[Dict[str, Any]],
    criteria: Dict[str, Any],
) -> None:
    chat_id = update.effective_chat.id
    if not matches:
        await safe_reply(update, "По указанным условиям архивы не найдены.")
        return

    summarized: List[str] = []
    file_entries: List[Tuple[Path, Dict[str, Any]]] = []
    for entry in matches:
        rel_path = entry.get("archive_path")
        if not rel_path:
            continue
        abs_path = ARCHIVE_DIR / rel_path
        if abs_path.is_file():
            file_entries.append((abs_path, entry))
            summarized.append(_describe_archive_entry(entry, len(file_entries)))

    if not file_entries:
        await safe_reply(update, "Совпадения найдены, но файлы отсутствуют на диске.")
        return

    if len(file_entries) == 1:
        await safe_reply(update, "Найден один документ:\n" + "\n".join(summarized))
        path, entry = file_entries[0]
        caption = (
            f"🔎 Архив: {entry.get('date') or '—'}, {entry.get('region') or '—'}\n"
            f"Билет: {entry.get('ticket_number') or '—'}, №: {entry.get('issue_number') or '—'}"
        )
        await send_document_from_path(context.bot, chat_id=chat_id, path=path, caption=caption)
        return

    if len(file_entries) <= 5:
        await safe_reply(update, "Найдено несколько документов:\n\n" + "\n\n".join(summarized))
        for path, entry in file_entries:
            caption = (
                f"🔎 Архив: {entry.get('date') or '—'}, {entry.get('region') or '—'}\n"
                f"Билет: {entry.get('ticket_number') or '—'}, №: {entry.get('issue_number') or '—'}"
            )
            await send_document_from_path(context.bot, chat_id=chat_id, path=path, caption=caption)
        return

    label = _build_search_label(criteria)
    timestamp = datetime.now().strftime("%d.%м.%Y_%H-%M-%S")

    def _create_zip() -> Path:
        DOCS_DIR.mkdir(parents=True, exist_ok=True)
        zip_name = sanitize_filename(f"archive_search_{label}_{timestamp}.zip")
        zip_path = DOCS_DIR / zip_name
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for path, entry in file_entries:
                zf.write(path, arcname=path.name)
        return zip_path

    zip_path = await asyncio.to_thread(_create_zip)
    try:
        caption = f"🔎 Результаты поиска ({len(file_entries)} файлов)."
        await send_document_from_path(context.bot, chat_id=chat_id, path=zip_path, caption=caption)
    finally:
        try:
            if zip_path.exists():
                zip_path.unlink()
        except OSError as cleanup_error:
            logger.debug(f"Не удалось удалить временный архив поиска {zip_path}: {cleanup_error}")


def load_admin_ids() -> None:
    global admin_ids
    ids: Set[int] = set()
    if ADMIN_FILE.exists():
        try:
            with ADMIN_FILE.open("r", encoding="utf-8") as f:
                data = json.load(f)
            ids = {int(item) for item in data if isinstance(item, int) or (isinstance(item, str) and item.isdigit())}
        except (OSError, json.JSONDecodeError) as err:
            logger.warning(f"Не удалось прочитать список администраторов: {err}")
    if not ids:
        ids = set(DEFAULT_ADMIN_IDS)
        admin_ids = ids
        save_admin_ids()
    else:
        admin_ids = ids


def save_admin_ids() -> None:
    ADMIN_FILE.parent.mkdir(parents=True, exist_ok=True)
    with ADMIN_FILE.open("w", encoding="utf-8") as f:
        json.dump(sorted(admin_ids), f, ensure_ascii=False, indent=2)


def is_admin(user_id: int) -> bool:
    return user_id in admin_ids


async def configure_bot_commands(bot) -> None:
    user_commands = [
        BotCommand("start", "🚀 Создать новое заключение"),
        BotCommand("help", "📚 Помощь по работе с ботом"),
        BotCommand("cancel", "❌ Отменить текущее действие"),
        BotCommand("menu", "📋 Открыть главное меню"),
        BotCommand("webapp", "🧾 Мини-приложение"),
        BotCommand("leaders", "🏆 Текущий рейтинг"),
        BotCommand("achievements", "🏅 Мои достижения"),
        BotCommand("void_ticket", "♻️ Обнулить заключение"),
        BotCommand("drafts", "💾 Черновики"),
    ]

    admin_commands = user_commands + [
        BotCommand("history", "🕑 Последние 10 заключений"),
        BotCommand("stats", "📊 Статистика по всем данным"),
        BotCommand("analytics", "📈 Аналитика"),
        BotCommand("reports", "📦 Мастер отчётов"),
        BotCommand("download_month", "📥 Выгрузка за месяц"),
        BotCommand("stats_period", "📈 Статистика за период"),
        BotCommand("search_archive", "🔎 Поиск в архиве"),
        BotCommand("backup", "💾 Резервная копия"),
        BotCommand("add_admin", "👥 Добавить администратора"),
        BotCommand("help_admin", "🔧 Справка администратора"),
    ]

    try:
        await bot.set_my_commands(user_commands, scope=BotCommandScopeAllPrivateChats())
    except TelegramError as e:
        logger.warning(f"Не удалось установить команды для пользователей: {e}")

    for admin_id in admin_ids:
        try:
            await bot.set_my_commands(admin_commands, scope=BotCommandScopeChat(admin_id))
        except TelegramError as e:
            logger.warning(f"Не удалось установить команды для администратора {admin_id}: {e}")


async def safe_reply(update: Update, text: str, retries: int = 3, base_delay: float = 2.0, **kwargs):
    chat_id = update.effective_chat.id
    kwargs_copy = dict(kwargs)
    last_error: Optional[Exception] = None
    last_recoverable = False

    for attempt in range(retries):
        try:
            return await update.message.reply_text(text, **kwargs)
        except RetryAfter as error:
            last_error = error
            last_recoverable = True
            delay = getattr(error, "retry_after", base_delay * (attempt + 1))
            logger.warning(f"Telegram просит подождать {delay} сек. перед повтором отправки сообщения.")
            await asyncio.sleep(delay)
        except (NetworkError, asyncio.TimeoutError) as error:
            last_error = error
            last_recoverable = True
            delay = base_delay * (attempt + 1)
            logger.warning(f"Сетевая ошибка при отправке сообщения: {error}. Повтор через {delay} сек.")
            await asyncio.sleep(delay)
        except TelegramError as error:
            last_error = error
            last_recoverable = False
            logger.warning(f"Ошибка Telegram при отправке сообщения: {error}")
            break

    if last_recoverable:
        await mark_network_issue(chat_id, text, kwargs_copy)
        await process_network_recovery(update.get_bot())

    if last_error:
        logger.error(f"Не удалось отправить сообщение: {last_error}")
    else:
        logger.error("Не удалось отправить сообщение после нескольких попыток.")
    return None


async def safe_bot_send_message(
    bot,
    chat_id: int,
    text: str,
    retries: int = 3,
    base_delay: float = 2.0,
    skip_notice_on_retry: bool = False,
    **kwargs,
):
    """Унифицированная отправка сообщений ботом с повторными попытками и сохранением при сбое."""
    kwargs_copy = dict(kwargs)
    last_error: Optional[Exception] = None
    last_recoverable = False

    for attempt in range(retries):
        try:
            return await bot.send_message(chat_id, text, **kwargs)
        except RetryAfter as error:
            last_error = error
            last_recoverable = True
            delay = getattr(error, "retry_after", base_delay * (attempt + 1))
            logger.warning(f"Telegram просит подождать {delay} сек. перед повтором отправки ботом.")
            await asyncio.sleep(delay)
        except (NetworkError, asyncio.TimeoutError) as error:
            last_error = error
            last_recoverable = True
            delay = base_delay * (attempt + 1)
            logger.warning(f"Сетевая ошибка при отправке ботом: {error}. Повтор через {delay} сек.")
            await asyncio.sleep(delay)
        except TelegramError as error:
            last_error = error
            last_recoverable = False
            logger.warning(f"Ошибка Telegram при отправке ботом: {error}")
            break

    if last_recoverable:
        await mark_network_issue(chat_id, text, kwargs_copy, skip_notice=skip_notice_on_retry)

    if last_error:
        logger.error(f"Не удалось отправить сообщение ботом: {last_error}")
    return None


async def safe_chat_action(bot, chat_id: int, action: ChatAction, *, message_thread_id: Optional[int] = None) -> None:
    """Send a chat action, swallowing network errors so UX hints do not break main logic."""
    try:
        await bot.send_chat_action(chat_id=chat_id, action=action, message_thread_id=message_thread_id)
    except TelegramError as error:
        logger.debug(f"Не удалось отправить chat action: {error}")


# -------------------- Работа с документами и Excel --------------------
def replace_placeholders_in_document(doc: Document, placeholders: Dict[str, str]) -> None:
    """Заменяет маркеры в документе на значения из словаря."""
    def _replace_in_runs(runs):
        for run in runs:
            for key, value in placeholders.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

    for paragraph in doc.paragraphs:
        if paragraph.runs:
            _replace_in_runs(paragraph.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        _replace_in_runs(paragraph.runs)

def add_borders_to_table(table: Any) -> None:
    """Добавляет границы ко всем ячейкам таблицы."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                border_element = OxmlElement(f"w:{border}")
                border_element.set(qn('w:val'), 'single')
                border_element.set(qn('w:sz'), '8')
                border_element.set(qn('w:space'), '0')
                border_element.set(qn('w:color'), 'auto')
                borders.append(border_element)
            tcPr.append(borders)

def populate_table_with_data(doc: Document, data: ConclusionData) -> None:
    """Заполняет первую таблицу документа данными о фотографиях."""
    if not doc.tables:
        logger.error("В документе отсутствуют таблицы.")
        return
    table = doc.tables[0]
    for i, item in enumerate(data.photo_desc, 1):
        try:
            new_row = table.add_row()
            row_cells = new_row.cells
            if len(row_cells) < 8:
                logger.error("Структура таблицы не соответствует ожидаемой (менее 8 столбцов).")
                continue

            photo_path = Path(item.get('photo', ""))
            row_cells[0].text = str(i)
            if photo_path.is_file():
                p = row_cells[2].paragraphs[0] if row_cells[2].paragraphs else row_cells[2].add_paragraph()
                p.add_run().add_picture(str(photo_path), width=Inches(1.0))
            else:
                row_cells[2].text = 'Фото отсутствует'

            description = item.get('description') or 'Нет описания'
            evaluation_value = item.get('evaluation') or 'Нет данных'
            row_cells[1].text = description
            row_cells[5].text = evaluation_value
            row_cells[6].text = evaluation_value
            row_cells[7].text = 'да'
        except Exception as e:
            logger.error(f"Ошибка при заполнении таблицы: {e}")
    add_borders_to_table(table)

async def create_document(user_id: int, username: str = "") -> Path:
    """Создает документ заключения на основе введенных данных пользователя."""
    data = await load_user_data_from_db(user_id)
    if (
        not data.photo_desc
        and not any(
            (data.department_number, data.issue_number, data.ticket_number, data.date, data.region)
        )
    ):
        raise ValueError("Не найдены данные пользователя для формирования документа.")
    if not TEMPLATE_PATH.exists():
        logger.error(f"Критическая ошибка: Шаблон документа '{TEMPLATE_PATH}' не найден.")
        raise FileNotFoundError(f"Шаблон '{TEMPLATE_PATH}' не найден. Обратитесь к администратору.")

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    selected_date = data.date or datetime.now().strftime('%d.%m.%Y')
    timestamp = datetime.now().strftime('%H-%M-%S')
    placeholders = {
        '{date}': selected_date,
        '{issue_number}': data.issue_number or 'Не указано',
        '{department_number}': data.department_number or 'Не указано',
        '{region}': data.region or 'Не указано',
        '{ticket_number}': data.ticket_number or 'Не указано',
        '{username}': username
    }

    base_filename = (f"{placeholders['{department_number}']}, Заключение антиквариат № "
                     f"{placeholders['{issue_number}']} (билет {placeholders['{ticket_number}']}), "
                     f"{placeholders['{region}']}, от {selected_date} {timestamp}.docx")
    
    safe_filename_str = sanitize_filename(base_filename)
    if not safe_filename_str:
        safe_filename_str = f"Заключение_{timestamp}.docx"
    filepath = DOCS_DIR / safe_filename_str

    suffix = Path(safe_filename_str).suffix or ".docx"
    stem = Path(safe_filename_str).stem or "Заключение"
    while filepath.exists():
        unique_suffix = ''.join(random.choices(string.ascii_lowercase + string.digits, k=4))
        candidate_name = sanitize_filename(f"{stem}_{unique_suffix}{suffix}")
        if not candidate_name:
            candidate_name = f"Заключение_{timestamp}_{unique_suffix}.docx"
        filepath = DOCS_DIR / candidate_name
        safe_filename_str = candidate_name

    def _build_document():
        try:
            doc = Document(TEMPLATE_PATH)
            if doc.paragraphs:
                doc.paragraphs[0].insert_paragraph_before(filepath.stem)
            else:
                doc.add_paragraph(filepath.stem)
            replace_placeholders_in_document(doc, placeholders)
            populate_table_with_data(doc, data)
            doc.save(filepath)
        except Exception as doc_error:
            logger.error(f"Не удалось сформировать документ {filepath}: {doc_error}", exc_info=True)
            raise

    try:
        await asyncio.to_thread(_build_document)
    except Exception as exc:
        raise RuntimeError("Ошибка при генерации документа заключения.") from exc
    logger.info(f"Документ сохранён: {filepath}")
    return filepath

async def update_excel(data: Union[ConclusionData, Dict[str, Any]]) -> None:
    """Обновляет Excel-файл информацией о заключении (построчно для каждого предмета)."""
    def _write_excel():
        if isinstance(data, ConclusionData):
            payload = data.to_dict()
        else:
            payload = data or {}

        if not EXCEL_FILE.exists():
            wb = Workbook()
            ws = wb.active
            ws.append(EXCEL_HEADERS)
        else:
            wb = load_workbook(EXCEL_FILE)
            ws = wb.active

        items = payload.get("photo_desc", [])
        for idx, item in enumerate(items, 1):
            row = [
                payload.get("ticket_number", "Не указано"),
                payload.get("issue_number", "Не указано"),
                payload.get("department_number", "Не указано"),
                payload.get("date", "Не указано"),
                payload.get("region", "Не указано"),
                idx,
                item.get("description", "Нет описания"),
                item.get("evaluation", "Нет данных")
            ]
            ws.append(row)
        wb.save(EXCEL_FILE)
        wb.close()

    async with excel_lock:
        await asyncio.to_thread(_write_excel)
        logger.info("Excel-файл успешно обновлен.")


def _remove_conclusion_from_excel_sync(conclusion: Dict[str, Any]) -> int:
    if not EXCEL_FILE.exists():
        return 0
    try:
        wb = load_workbook(EXCEL_FILE)
    except Exception as error:
        logger.error(f"Не удалось открыть Excel для удаления записи: {error}")
        return 0

    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        wb.close()
        return 0

    header = rows[0]
    data_rows = rows[1:]

    def normalize(value: Any) -> str:
        if value is None:
            return ""
        return str(value).strip()

    ticket_value = normalize(conclusion.get("ticket_number") or "Не указано")
    issue_value = normalize(conclusion.get("issue_number") or "Не указано")
    department_value = normalize(conclusion.get("department_number") or "Не указано")
    date_value = normalize(conclusion.get("date") or "Не указано")

    filtered: List[List[Any]] = []
    removed = 0
    for row in data_rows:
        if not row:
            continue
        row_ticket = normalize(row[0])
        row_issue = normalize(row[1])
        row_department = normalize(row[2])
        row_date = normalize(row[3])
        if (
            row_ticket == ticket_value
            and row_issue == issue_value
            and row_department == department_value
            and row_date == date_value
        ):
            removed += 1
            continue
        filtered.append(list(row))

    if removed:
        new_wb = Workbook()
        new_ws = new_wb.active
        new_ws.append(list(header))
        for row in filtered:
            new_ws.append(list(row))
        new_wb.save(EXCEL_FILE)
        new_wb.close()
    wb.close()
    return removed


async def remove_conclusion_from_excel(conclusion: Dict[str, Any]) -> int:
    async with excel_lock:
        return await asyncio.to_thread(_remove_conclusion_from_excel_sync, conclusion)

def build_summary(data: ConclusionData) -> str:
    """Формирует расширенную сводку введенных данных для подтверждения."""
    items = data.photo_desc
    total_items = len(items)
    total_value = 0
    for item in items:
        try:
            total_value += int(item.get('evaluation', 0))
        except (ValueError, TypeError):
            continue

    summary_lines = [
        f"Номер подразделения: {data.department_number or 'Не указано'}",
        f"Порядковый номер заключения: {data.issue_number or 'Не указано'}",
        f"Номер залогового билета: {data.ticket_number or 'Не указано'}",
        f"Дата заключения: {data.date or 'Не указано'}",
        f"Регион: {data.region or 'Не указано'}",
        "---",
        f"Всего предметов: {total_items}",
        f"Суммарная оценка: {total_value}",
        "---"
    ]
    if items:
        summary_lines.append("Первые 3 предмета:")
        for i, item in enumerate(items[:3], 1):
            desc = item.get('description', 'без описания')
            eval_val = item.get('evaluation', 'б/о')
            summary_lines.append(f"  {i}. {desc[:40]}... - {eval_val} руб.")

    return "\n".join(summary_lines)


async def show_summary(target: Any, context: CallbackContext, data: ConclusionData) -> None:
    clear_pending_items(context)
    metrics_enter_state(context, DialogState.CONFIRMATION)
    message, chat_id, thread_id, _ = _resolve_chat_context(target)
    if chat_id is None:
        return
    if isinstance(target, Update):
        await send_preview_photos(target, data)
    else:
        await send_preview_photos(
            None,
            data,
            bot=context.bot,
            chat_id=chat_id,
            thread_id=thread_id,
        )
    summary_text = build_summary(data)
    keyboard = build_confirmation_keyboard(data, include_back=has_previous_state(context))
    summary_message = await _send_via_target(
        target,
        context,
        f"Проверьте данные:\n\n{summary_text}",
        reply_markup=keyboard
    )
    if summary_message:
        context.user_data["summary_message_id"] = summary_message.message_id


async def prompt_for_state(
    target: Any,
    context: CallbackContext,
    data: ConclusionData,
    state: DialogState,
) -> None:
    """Prompt the user to provide data for the requested dialog state."""
    set_current_state(context, state)
    if state == DialogState.CONFIRMATION:
        await show_summary(target, context, data)
        return
    metrics_enter_state(context, state)

    prompt_map = {
        DialogState.DEPARTMENT: (
            f"{format_progress('department')}\nВведите номер подразделения (например 385).\nНажмите «Отмена», чтобы прервать оформление.",
            build_step_inline_keyboard(context=context),
        ),
        DialogState.ISSUE_NUMBER: (
            f"{format_progress('issue')}\nВведите порядковый номер заключения за день (например 1).",
            build_step_inline_keyboard(context=context),
        ),
        DialogState.TICKET_NUMBER: (
            f"{format_progress('ticket')}\nВведите номер залогового билета ({ticket_digits_phrase()}).",
            build_step_inline_keyboard(context=context),
        ),
        DialogState.DATE: (
            f"{format_progress('date')}\nВведите дату заключения в формате 01.03.2025.",
            build_step_inline_keyboard(context=context),
        ),
    }

    if state in prompt_map:
        text, markup = prompt_map[state]
        await _send_via_target(target, context, text, reply_markup=markup)
        return

    if state == DialogState.REGION:
        message = await _send_via_target(
            target,
            context,
            f"{format_progress('region')} Выберите регион:",
            reply_markup=build_region_inline_keyboard(context=context),
        )
        if message:
            context.user_data["region_prompt"] = {
                "chat_id": message.chat_id,
                "message_id": message.message_id,
            }
        return

    if state == DialogState.PHOTO:
        photo_count = len(data.photo_desc)
        await _send_via_target(
            target,
            context,
            f"{format_progress('photo')} Отправьте фото предмета.\n{PHOTO_REQUIREMENTS_MESSAGE}\n\n(Загружено: {photo_count}/{MAX_PHOTOS})",
            reply_markup=build_step_inline_keyboard(context=context),
        )
        return

    if state == DialogState.DESCRIPTION:
        await _send_via_target(
            target,
            context,
            f"{format_progress('description')} Добавьте короткое описание предмета.",
            reply_markup=build_step_inline_keyboard(context=context),
        )
        return

    if state == DialogState.EVALUATION:
        await _send_via_target(
            target,
            context,
            f"{format_progress('evaluation')} Введите оценку предмета (целое число, например 1500).",
            reply_markup=build_step_inline_keyboard(),
        )
        return

    if state == DialogState.MORE_PHOTO:
        buttons = [[
            InlineKeyboardButton("Добавить фото", callback_data=f"{ADD_PHOTO_PREFIX}yes"),
            InlineKeyboardButton("Перейти к сводке", callback_data=f"{ADD_PHOTO_PREFIX}no"),
        ]]
        message = await _send_via_target(
            target,
            context,
            f"{format_progress('evaluation')} Оценка сохранена. Добавить ещё одно фото? ({len(data.photo_desc)}/{MAX_PHOTOS}).",
            reply_markup=build_step_inline_keyboard(buttons),
        )
        if message:
            context.user_data["add_photo_message"] = {
                "chat_id": message.chat_id,
                "message_id": message.message_id,
            }
        return

    if state == DialogState.TESTING:
        message = await _send_via_target(
            target,
            context,
            f"{format_progress('mode')} Выберите режим создания заключения.",
            reply_markup=build_mode_keyboard(),
        )
        if message:
            context.user_data["mode_prompt"] = {
                "chat_id": message.chat_id,
                "message_id": message.message_id,
            }
        return

    # Для всех непризнанных состояний повторяем ввод подразделения
    await _send_via_target(
        target,
        context,
        f"{format_progress('department')}\nВведите номер подразделения (например 385).",
        reply_markup=build_step_inline_keyboard(context=context),
    )

# -------------------- Отправка с ретраями --------------------
async def safe_send_document(bot, chat_id, **kwargs):
    """Отправляет документ с 3 попытками и экспоненциальной задержкой."""
    document_obj = kwargs.get("document")
    for attempt in range(3):
        try:
            if document_obj and hasattr(document_obj, "seek"):
                document_obj.seek(0)
            return await bot.send_document(chat_id=chat_id, **kwargs)
        except RetryAfter as e:
            logger.warning(f"Flood control: ждём {e.retry_after} сек. Попытка {attempt + 1}")
            await asyncio.sleep(e.retry_after + 1)
        except (TimedOut, NetworkError) as e:
            logger.warning(f"Сетевая ошибка: {e}. Попытка {attempt + 1}")
            await asyncio.sleep(2 ** attempt)
        except TelegramError as e:
            logger.error(f"Ошибка Telegram при отправке документа: {e}", exc_info=True)
            break
    raise RuntimeError("Не удалось отправить документ после нескольких попыток.")


async def send_document_from_path(bot, chat_id: int, path: Path, **kwargs) -> None:
    """Открывает файл и отправляет его с ретраями."""
    if not path.is_file():
        raise FileNotFoundError(f"Файл для отправки не найден: {path}")

    filename = kwargs.pop("filename", path.name)
    def _open_file():
        return path.open("rb")

    file_handle = await asyncio.to_thread(_open_file)
    try:
        return await safe_send_document(bot, chat_id=chat_id, document=file_handle, filename=filename, **kwargs)
    finally:
        try:
            file_handle.close()
        except Exception:
            pass

# -------------------- Команды для статистики и истории --------------------
async def read_excel_data() -> List[List[str]]:
    """Асинхронно и безопасно считывает данные из Excel-файла."""
    def _read_excel():
        if not EXCEL_FILE.exists():
            return []
        wb = load_workbook(EXCEL_FILE)
        ws = wb.active
        rows = [list(row) for row in ws.iter_rows(min_row=2, values_only=True)]
        wb.close()
        return rows
    
    async with excel_lock:
        return await asyncio.to_thread(_read_excel)

async def history_handler(update: Update, context: CallbackContext) -> None:
    if not is_admin(update.message.from_user.id):
        await safe_reply(update, "Недостаточно прав для доступа к истории.")
        return
    records = await read_excel_data()
    if not records:
        await safe_reply(update, "История заключений пуста.")
        return
    recent_records = list(reversed(records[-200:])) if len(records) > 200 else list(reversed(records))
    context.user_data["history_cache"] = {
        "records": recent_records,
        "page_size": 6,
        "page": 1,
    }
    text, keyboard, current_page = build_history_page(recent_records, 1, 6)
    message = await safe_reply(update, text, reply_markup=keyboard)
    if message:
        context.user_data["history_message"] = {
            "chat_id": message.chat_id,
            "message_id": message.message_id,
        }
        context.user_data["history_cache"]["page"] = current_page


async def search_archive_handler(update: Update, context: CallbackContext) -> None:
    user = update.message.from_user
    if not is_admin(user.id):
        await safe_reply(update, "Команда доступна только администраторам.")
        return

    filters, errors = _parse_search_filters(context.args)
    if errors:
        await safe_reply(update, "\n".join(errors) + "\nПример: /search_archive ticket=01234567890 date=13.08.2025")
        return

    criteria: Dict[str, Any] = {}

    ticket = filters.get("ticket") or filters.get("ticket_number")
    if ticket:
        criteria["ticket_number"] = ticket

    issue = filters.get("issue") or filters.get("issue_number")
    if issue:
        criteria["issue_number"] = issue

    date_text = filters.get("date")
    if date_text:
        if not parse_date_str(date_text):
            await safe_reply(update, "Дата должна быть в формате ДД.ММ.ГГГГ. Пример: date=13.08.2025")
            return
        criteria["date"] = date_text

    region_text = filters.get("region")
    if region_text:
        normalized_region = normalize_region_input(region_text)
        if not normalized_region:
            await safe_reply(update, "Неизвестный регион. Используйте вариант из справочника.")
            return
        criteria["region"] = normalized_region

    date_from_text = filters.get("from") or filters.get("start")
    if date_from_text:
        start_date = parse_date_str(date_from_text)
        if not start_date:
            await safe_reply(update, "Начальная дата должна быть в формате ДД.ММ.ГГГГ.")
            return
        criteria["date_from"] = start_date

    date_to_text = filters.get("to") or filters.get("end")
    if date_to_text:
        end_date = parse_date_str(date_to_text)
        if not end_date:
            await safe_reply(update, "Конечная дата должна быть в формате ДД.ММ.ГГГГ.")
            return
        criteria["date_to"] = end_date

    entries = await asyncio.to_thread(_read_archive_index)
    if not entries:
        await safe_reply(update, "Индекс архива пуст — нет доступных файлов.")
        return

    summary_parts: List[str] = []
    if criteria.get("ticket_number"):
        summary_parts.append(f"билет {criteria['ticket_number']}")
    if criteria.get("issue_number"):
        summary_parts.append(f"№ {criteria['issue_number']}")
    if criteria.get("region"):
        summary_parts.append(f"регион {criteria['region']}")
    if criteria.get("date"):
        summary_parts.append(f"дата {criteria['date']}")
    if criteria.get("date_from") or criteria.get("date_to"):
        start_lbl = criteria.get("date_from")
        end_lbl = criteria.get("date_to")
        if isinstance(start_lbl, datetime):
            start_lbl = start_lbl.strftime("%d.%m.%Y")
        if isinstance(end_lbl, datetime):
            end_lbl = end_lbl.strftime("%d.%m.%Y")
        summary_parts.append(f"период {start_lbl or '—'} — {end_lbl or '—'}")
    if summary_parts:
        await safe_reply(update, "Поиск по фильтрам: " + ", ".join(summary_parts))

    def _match(entry: Dict[str, Any]) -> bool:
        if criteria.get("ticket_number") and (entry.get("ticket_number") != criteria["ticket_number"]):
            return False
        if criteria.get("issue_number") and (entry.get("issue_number") != criteria["issue_number"]):
            return False
        if criteria.get("region") and (entry.get("region") != criteria["region"]):
            return False
        entry_date = parse_date_str(entry.get("date"))
        if criteria.get("date") and (entry.get("date") != criteria["date"]):
            return False
        if criteria.get("date_from") and (not entry_date or entry_date < criteria["date_from"]):
            return False
        if criteria.get("date_to") and (not entry_date or entry_date > criteria["date_to"]):
            return False
        return True

    matched = [entry for entry in entries if _match(entry)]
    matched.sort(key=lambda e: parse_date_str(e.get("date")) or datetime.min, reverse=True)

    await _send_archive_search_results(update, context, matched, criteria)


async def leaders_handler(update: Update, context: CallbackContext) -> None:
    if not _is_db_ready():
        await safe_reply(update, "Статистика временно недоступна.")
        return

    week_label, weekly = await fetch_leaderboard("week", LEADERBOARD_SIZE)
    month_label, monthly = await fetch_leaderboard("month", LEADERBOARD_SIZE)

    def format_section(title: str, label: str, data: List[Dict[str, Any]]) -> str:
        if not data:
            return f"{title} ({label}): пока без лидеров."
        lines = []
        for idx, entry in enumerate(data, 1):
            total_eval = int(entry["total_evaluation"])
            lines.append(
                f"{idx}. {entry['username']} — {entry['completions']} заключ., {entry['items']} предметов, {total_eval} руб."
            )
        return f"{title} ({label}):\n" + "\n".join(lines)

    text_blocks = [
        format_section("🏆 Лидеры недели", week_label or "—", weekly),
        format_section("🥇 Лидеры месяца", month_label or "—", monthly),
    ]
    await safe_reply(update, "\n\n".join(text_blocks))

async def stats_handler(update: Update, context: CallbackContext) -> None:
    if not is_admin(update.message.from_user.id):
        await safe_reply(update, "Недостаточно прав для доступа к статистике.")
        return
    records = await read_excel_data()
    if not records:
        await safe_reply(update, "Нет данных для статистики.")
        return

    total_items = len(records)
    total_eval = 0.0
    regions: Dict[str, Dict[str, float]] = {}
    for rec in records:
        region_name = rec[4] or "Не указано"
        stats = regions.setdefault(region_name, {"count": 0, "total": 0.0})
        stats["count"] += 1
        try:
            value = float(rec[7] or 0)
        except (TypeError, ValueError):
            value = 0.0
        stats["total"] += value
        total_eval += value

    avg_eval = int(total_eval / total_items) if total_items else 0
    lines = []
    for region_name, stats in sorted(regions.items(), key=lambda item: item[1]["count"], reverse=True):
        count = int(stats["count"])
        total = int(stats["total"])
        average = int(stats["total"] / stats["count"]) if stats["count"] else 0
        lines.append(f"  {region_name}: {count} предмет(ов), сумма {total}, средняя {average}")

    top_region = lines[0].strip() if lines else "Нет данных"

    stats_text = (
        "📊 Сводная статистика:\n"
        f"Всего предметов: {total_items}\n"
        f"Суммарная оценка: {int(total_eval)}\n"
        f"Средняя оценка: {avg_eval}\n"
        f"Лидер по количеству: {top_region}\n\n"
        "Разбивка по регионам:\n"
        + "\n".join(lines)
    )
    await safe_reply(update, stats_text)


async def fetch_completion_stats(
    days: Optional[int] = None,
    *,
    include_deleted: bool = False,
    return_meta: bool = False,
) -> Union[List[Dict[str, Any]], Tuple[List[Dict[str, Any]], Dict[str, Any]]]:
    if not _is_db_ready():
        return ([], {}) if return_meta else []

    conditions: List[str] = []
    params: List[Any] = []
    if not include_deleted:
        conditions.append("(is_deleted IS NULL OR is_deleted = 0)")
    if days:
        since = datetime.now() - timedelta(days=days)
        conditions.append("completed_at >= ?")
        params.append(since.isoformat())

    select_columns = (
        "completed_at, item_count, total_evaluation, processing_time_seconds, step_metrics, "
        "region, department_number, issue_number, ticket_number"
    )
    query = f"SELECT {select_columns} FROM completions"
    if conditions:
        query += " WHERE " + " AND ".join(conditions)
    query += " ORDER BY completed_at ASC"

    async with db_lock:
        async with db.execute(query, tuple(params)) as cursor:
            rows = await cursor.fetchall()

    records: List[Dict[str, Any]] = []
    for row in rows:
        completed_at_str, item_count, total_eval, processing_time, metrics_blob, region, department, issue, ticket = row
        try:
            completed_at = datetime.fromisoformat(completed_at_str)
        except (TypeError, ValueError):
            continue
        try:
            step_metrics = json.loads(metrics_blob) if metrics_blob else {}
        except (TypeError, json.JSONDecodeError):
            step_metrics = {}
        records.append(
            {
                "completed_at": completed_at,
                "item_count": int(item_count or 0),
                "total_evaluation": float(total_eval or 0.0),
                "processing_time_seconds": float(processing_time or 0.0),
                "step_metrics": step_metrics,
                "region": region or "",
                "department_number": department or "",
                "issue_number": issue or "",
                "ticket_number": ticket or "",
            }
        )

    if not return_meta:
        return records

    excluded_count = 0
    total_count = len(records)
    try:
        excluded_conditions = ["is_deleted = 1"]
        excluded_params: List[Any] = []
        if days:
            since = datetime.now() - timedelta(days=days)
            excluded_conditions.append("completed_at >= ?")
            excluded_params.append(since.isoformat())
        query_excluded = "SELECT COUNT(*) FROM completions WHERE " + " AND ".join(excluded_conditions)
        async with db_lock:
            async with db.execute(query_excluded, tuple(excluded_params)) as cursor:
                row = await cursor.fetchone()
                if row:
                    excluded_count = int(row[0] or 0)
    except Exception as error:
        logger.error(f"Не удалось подсчитать исключённые записи: {error}")

    meta = {
        "excluded_count": excluded_count,
        "total_count": total_count + excluded_count,
    }
    return records, meta


def build_analytics_summary_text(records: List[Dict[str, Any]], days: int) -> str:
    total = len(records)
    items = sum(rec.get("item_count", 0) for rec in records)
    value = sum(rec.get("total_evaluation", 0.0) for rec in records)
    durations = [rec.get("processing_time_seconds", 0.0) for rec in records if rec.get("processing_time_seconds")]
    avg_duration = format_duration(statistics.mean(durations)) if durations else "—"
    value_formatted = format_number(value)
    return (
        f"📈 Аналитика за {days} дней:\n"
        f"• Заключений: {total}\n"
        f"• Предметов: {items}\n"
        f"• Суммарная оценка: {value_formatted} руб.\n"
        f"• Среднее время оформления: {avg_duration}"
    )


def build_analytics_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📊 Свежая сводка", callback_data=f"{ANALYTICS_CALLBACK_PREFIX}summary")],
        [
            InlineKeyboardButton("📈 Тренд 30 дней", callback_data=f"{ANALYTICS_CALLBACK_PREFIX}trend:30"),
            InlineKeyboardButton("📈 Тренд 90 дней", callback_data=f"{ANALYTICS_CALLBACK_PREFIX}trend:90"),
        ],
        [InlineKeyboardButton("⏱ Время этапов", callback_data=f"{ANALYTICS_CALLBACK_PREFIX}steps")],
        [InlineKeyboardButton("Закрыть", callback_data=f"{ANALYTICS_CALLBACK_PREFIX}close")],
    ])


def generate_trend_chart(records: List[Dict[str, Any]], days: int) -> Optional[io.BytesIO]:
    if not records:
        return None
    end_date = datetime.now().date()
    start_date = end_date - timedelta(days=days - 1)
    day_cursor = start_date
    counts: Dict[datetime.date, int] = {}
    while day_cursor <= end_date:
        counts[day_cursor] = 0
        day_cursor += timedelta(days=1)
    for rec in records:
        day = rec["completed_at"].date()
        if start_date <= day <= end_date:
            counts[day] = counts.get(day, 0) + 1

    dates = sorted(counts.keys())
    values = [counts[day] for day in dates]
    if not any(values):
        return None

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(dates, values, marker="o", linewidth=2)
    ax.set_title(f"Заключения за {days} дней")
    ax.set_ylabel("Количество")
    ax.set_xlabel("Дата")
    ax.grid(True, linestyle="--", alpha=0.4)
    ax.xaxis.set_major_formatter(DateFormatter("%d.%m"))
    fig.autofmt_xdate()
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_step_duration_chart(records: List[Dict[str, Any]]) -> Optional[io.BytesIO]:
    durations_map: Dict[str, List[float]] = {}
    for rec in records:
        metrics = rec.get("step_metrics") or {}
        durations = metrics.get("durations") or {}
        for state_name, seconds in durations.items():
            try:
                seconds_val = float(seconds)
            except (TypeError, ValueError):
                continue
            durations_map.setdefault(state_name, []).append(seconds_val)

    if not durations_map:
        return None

    averages = [
        (state_name, statistics.mean(values))
        for state_name, values in durations_map.items()
        if values
    ]
    averages.sort(key=lambda item: item[1], reverse=True)
    top_entries = averages[:7]
    labels = [get_state_label(name) for name, _ in top_entries]
    values = [duration for _, duration in top_entries]

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.bar(labels, values, color="#4c8bf5")
    ax.set_title("Средняя длительность этапов (с)")
    ax.set_ylabel("Секунды")
    ax.set_xticklabels(labels, rotation=20, ha="right")
    ax.grid(True, axis="y", linestyle="--", alpha=0.3)
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf


async def analytics_handler(update: Update, context: CallbackContext) -> None:
    user = update.message.from_user
    if not is_admin(user.id):
        await safe_reply(update, "Команда доступна только администраторам.")
        return
    await safe_chat_action(context.bot, update.effective_chat.id, ChatAction.TYPING)
    records = await fetch_completion_stats(30)
    summary_text = build_analytics_summary_text(records, 30)
    message = await safe_reply(update, summary_text, reply_markup=build_analytics_keyboard())
    if message:
        context.user_data["analytics_message"] = {
            "chat_id": message.chat_id,
            "message_id": message.message_id,
        }


async def analytics_callback_handler(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(ANALYTICS_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return

    action = query.data[len(ANALYTICS_CALLBACK_PREFIX):]
    message_info = context.user_data.get("analytics_message")
    chat_id = query.message.chat_id if query.message else None
    thread_id = getattr(query.message, "message_thread_id", None)

    if action == "close":
        await query.answer("Панель закрыта.")
        if message_info:
            try:
                await context.bot.edit_message_text(
                    "Аналитика закрыта.",
                    chat_id=message_info.get("chat_id"),
                    message_id=message_info.get("message_id"),
                )
            except TelegramError:
                pass
        context.user_data.pop("analytics_message", None)
        return

    if action == "summary":
        if chat_id is not None:
            await safe_chat_action(context.bot, chat_id, ChatAction.TYPING, message_thread_id=thread_id)
        records = await fetch_completion_stats(30)
        summary_text = build_analytics_summary_text(records, 30)
        if message_info:
            try:
                await context.bot.edit_message_text(
                    summary_text,
                    chat_id=message_info.get("chat_id"),
                    message_id=message_info.get("message_id"),
                    reply_markup=build_analytics_keyboard(),
                )
            except TelegramError as error:
                logger.debug(f"Не удалось обновить сообщение аналитики: {error}")
        await query.answer("Сводка обновлена.")
        return

    if action.startswith("trend:"):
        try:
            days = int(action.split(":", 1)[1])
        except (ValueError, IndexError):
            await query.answer("Неизвестный параметр периода.", show_alert=True)
            return
        records = await fetch_completion_stats(days)
        chart = generate_trend_chart(records, days)
        if not chart:
            await query.answer("Недостаточно данных для графика.", show_alert=True)
            return
        if chat_id is not None:
            await safe_chat_action(context.bot, chat_id, ChatAction.UPLOAD_PHOTO, message_thread_id=thread_id)
            caption = f"📈 Заключения за {days} дней"
            await context.bot.send_photo(chat_id=chat_id, photo=chart, caption=caption)
        await query.answer("График отправлен.")
        return

    if action == "steps":
        records = await fetch_completion_stats(60)
        chart = generate_step_duration_chart(records)
        if not chart:
            await query.answer("Недостаточно данных по длительностям.", show_alert=True)
            return
        if chat_id is not None:
            await safe_chat_action(context.bot, chat_id, ChatAction.UPLOAD_PHOTO, message_thread_id=thread_id)
            await context.bot.send_photo(chat_id=chat_id, photo=chart, caption="⏱ Средняя длительность этапов")
        await query.answer("Диаграмма отправлена.")
        return

    await query.answer()


async def backup_handler(update: Update, context: CallbackContext) -> None:
    user = update.message.from_user
    if not is_admin(user.id):
        await safe_reply(update, "Команда доступна только администраторам.")
        return
    await safe_chat_action(context.bot, update.effective_chat.id, ChatAction.UPLOAD_DOCUMENT)
    backup_path = await create_backup_archive()
    if not backup_path:
        await safe_reply(update, "Не удалось сформировать резервную копию.")
        return
    try:
        await send_document_from_path(context.bot, chat_id=update.effective_chat.id, path=backup_path, caption="📦 Резервная копия")
    finally:
        try:
            if backup_path.exists():
                backup_path.unlink()
        except OSError:
            pass


def build_admin_main_text() -> str:
    return (
        "🛠 Панель администратора\n"
        "Выберите действие из меню."
    )


def build_admin_keyboard(view: str = "root") -> InlineKeyboardMarkup:
    if view == "root":
        return InlineKeyboardMarkup([
            [InlineKeyboardButton("👥 Пользователи", callback_data=f"{ADMIN_CALLBACK_PREFIX}view:users")],
            [InlineKeyboardButton("♻️ Удалённые записи", callback_data=f"{ADMIN_CALLBACK_PREFIX}view:deleted")],
            [InlineKeyboardButton("📦 Резервная копия", callback_data=f"{ADMIN_CALLBACK_PREFIX}backup")],
            [InlineKeyboardButton("📜 Логи", callback_data=f"{ADMIN_CALLBACK_PREFIX}view:logs")],
            [InlineKeyboardButton("Закрыть", callback_data=f"{ADMIN_CALLBACK_PREFIX}close")],
        ])
    return InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Назад", callback_data=f"{ADMIN_CALLBACK_PREFIX}back")],
                                 [InlineKeyboardButton("Закрыть", callback_data=f"{ADMIN_CALLBACK_PREFIX}close")]])


def format_user_entry(user: Dict[str, Any]) -> str:
    status = "🚫 заблокирован" if user.get("is_blocked") else "✅ активен"
    lines = [f"{user['user_id']} • {user.get('username') or '—'} • {status}"]
    if user.get("last_completed"):
        lines.append(f"  Последнее оформление: {user['last_completed'][:16]}")
    if user.get("notes"):
        lines.append(f"  Заметка: {user['notes']}")
    return "\n".join(lines)


def build_deleted_entry_text(entry: Dict[str, Any]) -> str:
    deleted_at = entry.get("deleted_at") or "—"
    return (
        f"{entry['id']} • {entry.get('username') or '—'}\n"
        f"  Билет {entry.get('ticket_number') or '—'}, № {entry.get('issue_number') or '—'}\n"
        f"  Дата {entry.get('date') or '—'}, {entry.get('region') or '—'}\n"
        f"  Удалено: {deleted_at}"
    )


def read_log_tail(lines: int = 200) -> str:
    if not LOG_FILE.exists():
        return "Лог-файл ещё не создан."
    try:
        with LOG_FILE.open("r", encoding="utf-8", errors="ignore") as fh:
            content = fh.readlines()
    except OSError as error:
        logger.error(f"Не удалось прочитать лог-файл: {error}")
        return "Не удалось прочитать лог-файл."
    tail = content[-lines:]
    text = "".join(tail)
    if len(text) > 3500:
        text = text[-3500:]
    return text or "Лог пуст."


async def create_backup_archive() -> Optional[Path]:
    sources: List[Path] = []
    for path in [DATABASE_FILE, EXCEL_FILE, DOCS_DIR, ARCHIVE_DIR, LOG_DIR]:
        if path.exists():
            sources.append(path)
    if not sources:
        return None

    BACKUP_ROOT.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_name = sanitize_filename(f"backup_{timestamp}.zip")
    backup_path = BACKUP_ROOT / backup_name

    def _build_backup() -> Path:
        with zipfile.ZipFile(backup_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for source in sources:
                if source.is_file():
                    arcname = source.name if source.parent == Path.cwd() else str(source.relative_to(Path.cwd()))
                    zf.write(source, arcname=arcname)
                else:
                    for file_path in source.rglob("*"):
                        if file_path.is_file():
                            try:
                                arcname = str(file_path.relative_to(Path.cwd()))
                            except ValueError:
                                arcname = file_path.name
                            zf.write(file_path, arcname=arcname)
        return backup_path

    try:
        await asyncio.to_thread(_build_backup)
        return backup_path
    except Exception as error:
        logger.error(f"Не удалось создать резервную копию: {error}")
        if backup_path.exists():
            try:
                backup_path.unlink()
            except OSError:
                pass
        return None


async def admin_handler(update: Update, context: CallbackContext) -> None:
    user = update.message.from_user
    if not is_admin(user.id):
        await safe_reply(update, "Команда доступна только администраторам.")
        return

    text = build_admin_main_text()
    keyboard = build_admin_keyboard("root")
    message = await safe_reply(update, text, reply_markup=keyboard)
    if message:
        context.user_data["admin_panel_message"] = {
            "chat_id": message.chat_id,
            "message_id": message.message_id,
        }
        context.user_data["admin_panel_view"] = "root"


async def admin_callback_handler(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(ADMIN_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return

    user = query.from_user
    if not is_admin(user.id):
        await query.answer("Недостаточно прав.", show_alert=True)
        return

    action = query.data[len(ADMIN_CALLBACK_PREFIX):]
    panel_info = context.user_data.get("admin_panel_message")

    async def _edit_panel(text: str, keyboard: InlineKeyboardMarkup) -> None:
        if panel_info:
            try:
                await context.bot.edit_message_text(
                    text,
                    chat_id=panel_info.get("chat_id"),
                    message_id=panel_info.get("message_id"),
                    reply_markup=keyboard,
                )
            except TelegramError as error:
                logger.debug(f"Не удалось обновить панель администратора: {error}")

    if action == "close":
        if panel_info:
            try:
                await context.bot.edit_message_text(
                    "Панель администратора закрыта.",
                    chat_id=panel_info.get("chat_id"),
                    message_id=panel_info.get("message_id"),
                )
            except TelegramError:
                pass
        context.user_data.pop("admin_panel_message", None)
        context.user_data.pop("admin_panel_view", None)
        await query.answer("Панель закрыта.")
        return

    if action == "back":
        await _edit_panel(build_admin_main_text(), build_admin_keyboard("root"))
        context.user_data["admin_panel_view"] = "root"
        await query.answer()
        return

    if action == "view:users":
        users = await fetch_recent_users(10)
        lines = ["👥 Последние активные пользователи:"]
        if not users:
            lines.append("Нет данных.")
        else:
            for entry in users:
                lines.append(format_user_entry(entry))
                lines.append("")
        keyboard_rows: List[List[InlineKeyboardButton]] = []
        for entry in users:
            button_text = "🔓 Разрешить" if entry.get("is_blocked") else "🚫 Заблокировать"
            action_key = "unblock" if entry.get("is_blocked") else "block"
            keyboard_rows.append([
                InlineKeyboardButton(
                    text=f"{button_text} {entry['user_id']}",
                    callback_data=f"{ADMIN_CALLBACK_PREFIX}user:{action_key}:{entry['user_id']}"
                )
            ])
        keyboard_rows.append([InlineKeyboardButton("⬅️ Назад", callback_data=f"{ADMIN_CALLBACK_PREFIX}back")])
        keyboard_rows.append([InlineKeyboardButton("Закрыть", callback_data=f"{ADMIN_CALLBACK_PREFIX}close")])
        await _edit_panel("\n".join(lines).strip(), InlineKeyboardMarkup(keyboard_rows))
        context.user_data["admin_panel_view"] = "users"
        await query.answer()
        return

    if action.startswith("user:block:") or action.startswith("user:unblock:"):
        parts = action.split(":")
        target_id = int(parts[-1])
        block = parts[1] == "block"
        await set_user_block_status(target_id, block)
        try:
            if block:
                await context.bot.send_message(target_id, "Ваш доступ к боту временно ограничен. Обратитесь к администратору.")
            else:
                await context.bot.send_message(target_id, "Доступ к боту восстановлен. Спасибо!")
        except TelegramError:
            pass
        await query.answer("Статус пользователя обновлён.")
        context.user_data["admin_panel_view"] = "users"
        users = await fetch_recent_users(10)
        lines = ["👥 Последние активные пользователи:"]
        if not users:
            lines.append("Нет данных.")
        else:
            for entry in users:
                lines.append(format_user_entry(entry))
                lines.append("")
        keyboard_rows: List[List[InlineKeyboardButton]] = []
        for entry in users:
            button_text = "🔓 Разрешить" if entry.get("is_blocked") else "🚫 Заблокировать"
            action_key = "unblock" if entry.get("is_blocked") else "block"
            keyboard_rows.append([
                InlineKeyboardButton(
                    text=f"{button_text} {entry['user_id']}",
                    callback_data=f"{ADMIN_CALLBACK_PREFIX}user:{action_key}:{entry['user_id']}"
                )
            ])
        keyboard_rows.append([InlineKeyboardButton("⬅️ Назад", callback_data=f"{ADMIN_CALLBACK_PREFIX}back")])
        keyboard_rows.append([InlineKeyboardButton("Закрыть", callback_data=f"{ADMIN_CALLBACK_PREFIX}close")])
        await _edit_panel("\n".join(lines).strip(), InlineKeyboardMarkup(keyboard_rows))
        return

    if action == "view:deleted":
        entries = await fetch_soft_deleted_completions(8)
        lines = ["♻️ Помеченные как удалённые заключения:"]
        if not entries:
            lines.append("Список пуст.")
        else:
            for entry in entries:
                lines.append(build_deleted_entry_text(entry))
                lines.append("")
        keyboard_rows = []
        for entry in entries:
            keyboard_rows.append([
                InlineKeyboardButton(
                    text=f"Восстановить {entry['id']}",
                    callback_data=f"{ADMIN_CALLBACK_PREFIX}restore:{entry['id']}"
                )
            ])
        keyboard_rows.append([InlineKeyboardButton("⬅️ Назад", callback_data=f"{ADMIN_CALLBACK_PREFIX}back")])
        keyboard_rows.append([InlineKeyboardButton("Закрыть", callback_data=f"{ADMIN_CALLBACK_PREFIX}close")])
        await _edit_panel("\n".join(lines).strip(), InlineKeyboardMarkup(keyboard_rows))
        context.user_data["admin_panel_view"] = "deleted"
        await query.answer()
        return

    if action.startswith("restore:"):
        try:
            target_id = int(action.split(":", 1)[1])
        except ValueError:
            await query.answer("Некорректный идентификатор.", show_alert=True)
            return
        restore_result = await restore_completion_record(target_id, restorer_id=user.id)
        if not restore_result.get("restored"):
            await query.answer("Не удалось восстановить запись.", show_alert=True)
            return
        restored_record = restore_result.get("record") or {}
        await query.answer("Запись восстановлена.")
        if restored_record.get("group_chat_id") and restored_record.get("group_message_id"):
            try:
                await context.bot.edit_message_reply_markup(
                    chat_id=restored_record["group_chat_id"],
                    message_id=restored_record["group_message_id"],
                    reply_markup=build_void_keyboard(restored_record["id"]),
                )
            except TelegramError:
                pass
        await send_personal_stats(context.bot, restored_record.get("user_id"))
        if restored_record.get("user_id"):
            await safe_bot_send_message(
                context.bot,
                restored_record["user_id"],
                "Одно из ваших заключений восстановлено в отчётах.",
                skip_notice_on_retry=True,
            )
        # Refresh view
        entries = await fetch_soft_deleted_completions(8)
        lines = ["♻️ Помеченные как удалённые заключения:"]
        if not entries:
            lines.append("Список пуст.")
        else:
            for entry in entries:
                lines.append(build_deleted_entry_text(entry))
                lines.append("")
        keyboard_rows = []
        for entry in entries:
            keyboard_rows.append([
                InlineKeyboardButton(
                    text=f"Восстановить {entry['id']}",
                    callback_data=f"{ADMIN_CALLBACK_PREFIX}restore:{entry['id']}"
                )
            ])
        keyboard_rows.append([InlineKeyboardButton("⬅️ Назад", callback_data=f"{ADMIN_CALLBACK_PREFIX}back")])
        keyboard_rows.append([InlineKeyboardButton("Закрыть", callback_data=f"{ADMIN_CALLBACK_PREFIX}close")])
        await _edit_panel("\n".join(lines).strip(), InlineKeyboardMarkup(keyboard_rows))
        return

    if action == "view:logs":
        log_text = read_log_tail(200)
        display_text = "📜 Последние строки логов:\n" + log_text
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("⬅️ Назад", callback_data=f"{ADMIN_CALLBACK_PREFIX}back")],
            [InlineKeyboardButton("📄 Скачать файл", callback_data=f"{ADMIN_CALLBACK_PREFIX}download_logs")],
            [InlineKeyboardButton("Закрыть", callback_data=f"{ADMIN_CALLBACK_PREFIX}close")],
        ])
        await _edit_panel(display_text, keyboard)
        context.user_data["admin_panel_view"] = "logs"
        await query.answer()
        return

    if action == "download_logs":
        if LOG_FILE.exists():
            await send_document_from_path(context.bot, chat_id=query.message.chat_id, path=LOG_FILE, caption="📜 Лог-файл")
            await query.answer("Лог отправлен.")
        else:
            await query.answer("Лог-файл отсутствует.", show_alert=True)
        return

    if action == "backup":
        if chat_id is not None:
            await safe_chat_action(
                context.bot,
                chat_id,
                ChatAction.UPLOAD_DOCUMENT,
                message_thread_id=getattr(query.message, "message_thread_id", None),
            )
        backup_path = await create_backup_archive()
        if backup_path:
            try:
                await send_document_from_path(context.bot, chat_id=query.message.chat_id, path=backup_path, caption="📦 Резервная копия")
            finally:
                try:
                    if backup_path.exists():
                        backup_path.unlink()
                except OSError:
                    pass
            await query.answer("Резервная копия подготовлена.")
        else:
            await query.answer("Не удалось собрать резервную копию.", show_alert=True)
        return

    await query.answer()


async def download_month_handler(update: Update, context: CallbackContext) -> None:
    user_id = update.message.from_user.id
    if not is_admin(user_id):
        await safe_reply(update, "Недостаточно прав для выгрузки данных.")
        return

    if not context.args:
        await safe_reply(update, "Использование: /download_month ММ.ГГГГ [Регион]")
        return

    month_text = context.args[0]
    bounds = get_month_bounds(month_text)
    if not bounds:
        await safe_reply(update, "❗ Неверный формат. Укажите месяц как ММ.ГГГГ, например 03.2025")
        return

    region = None
    if len(context.args) > 1:
        candidate = " ".join(context.args[1:])
        matched_region = match_region_name(candidate)
        if matched_region:
            region = matched_region
        else:
            await safe_reply(update, "❗ Неизвестный регион. Используйте вариант из справочника.")
            return

    start_date, end_date = bounds
    await send_month_archive(update, context, month_text, start_date, end_date, region)


async def stats_period_handler(update: Update, context: CallbackContext) -> None:
    user_id = update.message.from_user.id
    if not is_admin(user_id):
        await safe_reply(update, "Недостаточно прав для доступа к статистике.")
        return

    if len(context.args) < 2:
        await safe_reply(update, "Использование: /stats_period ДД.ММ.ГГГГ ДД.ММ.ГГГГ [Регион]")
        return

    start_date = parse_date_str(context.args[0])
    end_date = parse_date_str(context.args[1])
    if not start_date or not end_date:
        await safe_reply(update, "❗ Неверный формат дат. Используйте ДД.ММ.ГГГГ")
        return

    if start_date > end_date:
        start_date, end_date = end_date, start_date

    region = None
    if len(context.args) > 2:
        candidate = " ".join(context.args[2:])
        matched_region = match_region_name(candidate)
        if matched_region:
            region = matched_region
        else:
            await safe_reply(update, "❗ Неизвестный регион. Используйте вариант из справочника.")
            return

    await send_period_stats(update, start_date, end_date, region)

# -------------------- Диалог бота --------------------
async def menu_handler(update: Update, context: CallbackContext) -> None:
    user_id = update.message.from_user.id
    if not await ensure_user_not_blocked_message(update, context):
        return
    markup = build_main_menu(user_id)
    await safe_reply(update, "📋 Главное меню:", reply_markup=markup)

async def help_handler(update: Update, context: CallbackContext) -> None:
    if not await ensure_user_not_blocked_message(update, context):
        return
    message = (
        "Памятка по оформлению:\n"
        "• /start — введите подразделение, номер заключения, билет, дату и регион.\n"
        f"• Для каждого предмета: фото (JPG/PNG до {MAX_PHOTO_SIZE_MB} МБ), краткое описание и оценка.\n"
        "• После каждого фото решите, нужно ли добавить следующее — бот спросит отдельно.\n"
        "• Перед отправкой вы увидите сводку и сможете подтвердить или отменить.\n"
        "• Выберите режим: тестовый (файл только у вас) или окончательный (файл уходит в рабочую группу).\n"
        "• /cancel прерывает текущий сценарий, /leaders показывает рейтинг, /menu открывает главное меню."
        "\n"
        "• /achievements — уровни, награды и ваши ближайшие цели. Опыт начисляется за оформленные окончательные заключения."
    )
    await safe_reply(update, message)


async def webapp_handler(update: Update, context: CallbackContext) -> None:
    if not await ensure_user_not_blocked_message(update, context):
        return
    webapp_url = load_webapp_url()
    button = InlineKeyboardButton(
        "Открыть мини-приложение",
        web_app=WebAppInfo(url=webapp_url),
    )
    keyboard = InlineKeyboardMarkup([[button]])
    text = (
        "Откройте мини-приложение для оформления заключений и просмотра аналитики. "
        "В браузере появится веб-интерфейс с формой и отчётами."
    )
    await safe_reply(update, text, reply_markup=keyboard)


async def help_admin_handler(update: Update, context: CallbackContext) -> None:
    if not is_admin(update.message.from_user.id):
        await safe_reply(update, "Команда доступна только администраторам.")
        return

    if not await ensure_user_not_blocked_message(update, context):
        return

    admin_help = (
        "Инструменты администратора:\n"
        "• /history — последние 10 записей (по предметам).\n"
        "• /stats — общая статистика.\n"
        "• /download_month ММ.ГГГГ [Регион] — архив DOCX за месяц.\n"
        "• /void_ticket <номер билета> [дата] [№] — исключить окончательное заключение из отчётов (аналог кнопки в чате).\n"
        "• Кнопка «Удалить из отчётов» под сообщением в рабочей группе также доступна администраторам и выполняет /void_ticket.\n"
        "• /stats_period ДД.ММ.ГГГГ ДД.ММ.ГГГГ [Регион] — статистика за период.\n"
        "• /reports — мастер отчётов (архивы, Excel, сводки).\n"
        "• /leaders — рейтинг лидеров.\n"
        "• /add_admin ID — добавить администратора.\n"
        "• /menu — открыть админское меню, /help — пользовательская инструкция."
    )
    await safe_reply(update, admin_help)


async def drafts_handler(update: Update, context: CallbackContext) -> None:
    user = update.effective_user
    if not user:
        await safe_reply(update, "Не удалось определить пользователя.")
        return
    if await is_user_blocked(user.id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return
    draft = await fetch_draft(user.id)
    if not draft or not draft.get("data"):
        await safe_reply(update, "Сохранённых черновиков нет.")
        return

    context.user_data["awaiting_draft_choice"] = True
    context.user_data["draft_decision_context"] = "command"
    context.user_data["pending_draft"] = draft
    context.user_data["draft_discard_label"] = "Удалить черновик"

    message = await safe_reply(
        update,
        build_draft_summary_text(draft) + "\n\nВыберите действие:",
        reply_markup=build_draft_keyboard(discard_label=context.user_data["draft_discard_label"])
    )
    if message:
        context.user_data["draft_prompt_message"] = {
            "chat_id": message.chat_id,
            "message_id": message.message_id,
        }


async def add_admin_handler(update: Update, context: CallbackContext) -> None:
    requester_id = update.message.from_user.id
    if not is_admin(requester_id):
        await safe_reply(update, "Недостаточно прав для добавления администратора.")
        return

    if not context.args:
        await safe_reply(update, "Использование: /add_admin <ID пользователя>")
        return

    try:
        new_admin_id = int(context.args[0])
    except ValueError:
        await safe_reply(update, "❗ ID должен быть числом. Попробуйте ещё раз.")
        return

    if new_admin_id <= 0:
        await safe_reply(update, "❗ ID должен быть положительным числом.")
        return

    if new_admin_id == requester_id:
        await safe_reply(update, "Нельзя добавить самого себя.")
        return

    chat_member = None
    try:
        chat_member = await context.bot.get_chat(new_admin_id)
    except TelegramError as err:
        logger.warning(f"Не удалось получить информацию о пользователе {new_admin_id}: {err}")
    if chat_member and getattr(chat_member, "is_bot", False):
        await safe_reply(update, "Нельзя назначить бот-аккаунт администратором.")
        return

    if new_admin_id in admin_ids:
        await safe_reply(update, "Этот пользователь уже имеет права администратора.")
        return

    admin_ids.add(new_admin_id)
    save_admin_ids()
    await configure_bot_commands(context.bot)

    await safe_reply(update, f"👥 Пользователь {new_admin_id} добавлен в список администраторов.")
    try:
        await context.bot.send_message(new_admin_id, "🎉 Вам выданы права администратора в боте заключений.")
    except TelegramError as e:
        logger.warning(f"Не удалось уведомить нового администратора {new_admin_id}: {e}")


async def void_ticket_handler(update: Update, context: CallbackContext) -> None:
    user = update.effective_user
    if await is_user_blocked(user.id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return
    if not context.args:
        await safe_reply(update, "Использование: /void_ticket <номер билета> [ДД.ММ.ГГГГ] [номер заключения]")
        return

    ticket_number = context.args[0].strip()
    if not ticket_number:
        await safe_reply(update, "❗ Укажите номер билета.")
        return

    date_text: Optional[str] = None
    issue_number: Optional[str] = None
    for arg in context.args[1:]:
        candidate = arg.strip()
        if not candidate:
            continue
        if parse_date_str(candidate):
            date_text = candidate
        elif candidate.isdigit():
            issue_number = candidate

    records = await fetch_completions_by_ticket(ticket_number, date_text, issue_number)
    if not records:
        await safe_reply(update, "❗ Записей с таким билетом не найдено.")
        return

    if len(records) > 1 and not (date_text or issue_number):
        showcase = []
        for rec in records[:5]:
            issue_label = rec.get("issue_number") or "—"
            date_label = rec.get("date") or "не указана"
            created = (rec.get("completed_at") or "")[:16]
            showcase.append(f"• №{issue_label} от {date_label} (создано {created})")
        message = (
            "Найдено несколько заключений с таким билетом.\n"
            "Уточните дату (ДД.ММ.ГГГГ) или номер заключения из списка:\n"
            + "\n".join(showcase)
        )
        await safe_reply(update, message)
        return

    record = records[0]
    if record.get("user_id") != user.id and not is_admin(user.id):
        await safe_reply(update, "❗ Недостаточно прав для обнуления этого заключения.")
        return

    summary = await soft_delete_completion_record(record, initiator_id=user.id, note="void_ticket command")
    if summary.get("already_deleted"):
        await safe_reply(update, "Эта запись уже помечена как удалённая.")
        return
    if not summary.get("db_marked"):
        await safe_reply(update, "Не удалось пометить запись как удалённую. Проверьте журналы ошибок.")
        return

    status_line = "Статус: помечено как удалённое." if summary.get("db_marked") else "Статус: не изменилось."
    archive_line = "Архив обновлён." if summary.get("archive_marked") else "Архив: отметка не потребовалась."

    response_lines = [
        "♻️ Заключение исключено из отчётов.",
        f"Билет: {record.get('ticket_number') or 'не указан'}, заключение №{record.get('issue_number') or 'не указано'}, дата {record.get('date') or 'не указана'}.",
        status_line,
        archive_line,
        "Для восстановления используйте /admin → «Удалённые записи».",
    ]
    await safe_reply(update, "\n".join(response_lines))

    if record.get("user_id") != user.id:
        await safe_bot_send_message(
            context.bot,
            record["user_id"],
            "Одно из ваших заключений исключено из отчётов. При необходимости создайте новое.",
            skip_notice_on_retry=True,
        )

    if summary.get("db_marked"):
        await send_personal_stats(context.bot, record["user_id"])

    if record.get("group_chat_id") and record.get("group_message_id"):
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=record["group_chat_id"],
                message_id=record["group_message_id"],
                reply_markup=None,
            )
        except TelegramError:
            pass
        initiator_name = user.full_name if user else "Сотрудник"
        group_note = (
            f"♻️ Заключение по билету {record.get('ticket_number') or 'не указано'} исключено из отчётов (инициатор: {initiator_name})."
        )
        await safe_bot_send_message(
            context.bot,
            record["group_chat_id"],
            group_note,
            skip_notice_on_retry=True,
            message_thread_id=record.get("thread_id"),
        )


async def void_callback_handler(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    if not query or not query.data:
        return

    data = query.data
    if not data.startswith(VOID_CALLBACK_PREFIX):
        await query.answer()
        return

    try:
        completion_id = int(data[len(VOID_CALLBACK_PREFIX):])
    except ValueError:
        await query.answer("Некорректный идентификатор.", show_alert=True)
        return

    record = await fetch_completion_by_id(completion_id)
    if not record:
        await query.answer("Запись уже удалена.", show_alert=True)
        try:
            await query.edit_message_reply_markup(reply_markup=None)
        except TelegramError:
            pass
        return

    user = query.from_user
    initiator_id = user.id if user else None
    if initiator_id not in (record.get("user_id"),) and not is_admin(initiator_id or 0):
        await query.answer("Недостаточно прав.", show_alert=True)
        return
    if not await ensure_user_not_blocked_query(query, context):
        return

    summary = await soft_delete_completion_record(record, initiator_id=initiator_id, note="void_button")
    if summary.get("already_deleted"):
        await query.answer("Запись уже помечена как удалённая.", show_alert=True)
        return
    if not summary.get("db_marked"):
        await query.answer("Не удалось пометить запись как удалённую. Попробуйте позже.", show_alert=True)
        return

    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    if summary.get("db_marked"):
        await send_personal_stats(context.bot, record["user_id"])
    if record.get("user_id") != initiator_id:
        await safe_bot_send_message(
            context.bot,
            record["user_id"],
            "Одно из ваших заключений исключено из отчётов. При необходимости создайте новое.",
            skip_notice_on_retry=True,
        )

    initiator_name = user.full_name if user else "Сотрудник"
    group_note = (
        f"♻️ Заключение по билету {record.get('ticket_number') or 'не указано'} исключено из отчётов (инициатор: {initiator_name})."
    )
    if record.get("group_chat_id"):
        await safe_bot_send_message(
            context.bot,
            record["group_chat_id"],
            group_note,
            skip_notice_on_retry=True,
            message_thread_id=record.get("thread_id"),
        )

    await query.answer("Заключение помечено как удалённое.", show_alert=True)


async def deleted_message_handler(update: Update, context: CallbackContext) -> None:
    chat = update.effective_chat
    if not chat or chat.id != MAIN_GROUP_CHAT_ID:
        return

    message = update.effective_message
    if not message:
        return

    message_ids: List[int] = []
    if hasattr(message, "deleted_message_ids") and getattr(message, "deleted_message_ids", None):
        message_ids.extend(message.deleted_message_ids)
    else:
        msg_id = getattr(message, "message_id", None)
        if msg_id is not None:
            message_ids.append(msg_id)

    if not message_ids:
        return

    thread_id = getattr(message, "message_thread_id", None)
    for msg_id in message_ids:
        record = await fetch_completion_by_message(chat.id, msg_id, thread_id=thread_id)
        if not record:
            continue
        summary = await soft_delete_completion_record(record, initiator_id=None, note="message_deleted")
        if not summary.get("db_marked"):
            continue
        if record.get("group_chat_id") and record.get("group_message_id"):
            try:
                await context.bot.edit_message_reply_markup(
                    chat_id=record["group_chat_id"],
                    message_id=record["group_message_id"],
                    reply_markup=None,
                )
            except TelegramError:
                pass
        await send_personal_stats(context.bot, record["user_id"])
        await safe_bot_send_message(
            context.bot,
            record["user_id"],
            "Одно из ваших заключений удалено из группы и помечено как исключённое.",
            skip_notice_on_retry=True,
        )


def _report_data(context: CallbackContext) -> Dict[str, Any]:
    return context.user_data.setdefault("report", {})


async def _reports_finish(update: Update, context: CallbackContext, message: str) -> int:
    context.user_data.pop("report", None)
    context.user_data.pop("reports_menu_message", None)
    prompt = context.user_data.pop("report_region_prompt", None)
    if prompt:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=prompt.get("chat_id"),
                message_id=prompt.get("message_id"),
                reply_markup=None,
            )
        except (TelegramError, TypeError, KeyError):
            pass

    message_obj, chat_id, thread_id, user_id = _resolve_chat_context(update)
    if message_obj is None or chat_id is None:
        return ConversationHandler.END

    if isinstance(update, Update):
        await safe_reply(update, message, reply_markup=build_main_menu(user_id or 0))
    else:
        await safe_bot_send_message(
            context.bot,
            chat_id,
            message,
            reply_markup=build_main_menu(user_id or 0),
            message_thread_id=thread_id,
        )
        try:
            await update.answer()
        except TelegramError:
            pass
    return ConversationHandler.END


async def send_month_report(update: Update, context: CallbackContext, month_text: str, start_date: datetime, end_date: datetime, region: Optional[str]) -> None:
    filtered = await filter_records(start_date=start_date, end_date=end_date, region=region)
    if not filtered:
        await safe_reply(update, f"За {month_text} записей не найдено.")
        return

    await safe_chat_action(context.bot, update.effective_chat.id, ChatAction.UPLOAD_DOCUMENT)

    region_label = region or "Все регионы"
    try:
        filepath = await create_excel_snapshot(filtered, f"conclusions_{month_text}_{region_label}")
        caption = f"Заключения за {month_text} ({region_label})"
        await send_document_from_path(context.bot, chat_id=update.message.chat_id, path=filepath, caption=caption)
        await safe_reply(update, "📥 Файл с заключениями отправлен.")
    finally:
        try:
            if 'filepath' in locals() and filepath.exists():
                filepath.unlink()
        except Exception as cleanup_error:
            logger.debug(f"Не удалось удалить временный файл {filepath}: {cleanup_error}")


async def send_period_stats(update: Update, start_date: datetime, end_date: datetime, region: Optional[str]) -> None:
    filtered = await filter_records(start_date=start_date, end_date=end_date, region=region)
    if not filtered:
        await safe_reply(update, "За выбранный период записей не найдено.")
        return

    await safe_chat_action(context.bot, update.effective_chat.id, ChatAction.TYPING)

    total_items = len(filtered)
    total_eval = 0
    regions: Dict[str, int] = {}
    for row in filtered:
        region_name = row[4] or "Не указано"
        regions[region_name] = regions.get(region_name, 0) + 1
        try:
            total_eval += int(row[7] or 0)
        except (TypeError, ValueError):
            continue

    period_text = f"{start_date.strftime('%d.%m.%Y')} — {end_date.strftime('%d.%m.%Y')}"
    region_lines = "\n".join([f"  {r_name}: {count}" for r_name, count in sorted(regions.items(), key=lambda x: x[0])])
    if region:
        region_filter_text = f"Фильтр по региону: {region}\n"
    else:
        region_filter_text = ""
    await safe_reply(
        update,
        "📈 Статистика за период:\n"
        f"Диапазон: {period_text}\n"
        f"{region_filter_text}"
        f"Всего предметов: {total_items}\n"
        f"Суммарная оценка: {total_eval}\n\n"
        "Предметов по регионам:\n"
        f"{region_lines if region_lines else 'Нет данных'}"
    )


async def send_region_summary(update: Update, start_date: datetime, end_date: datetime) -> None:
    filtered = await filter_records(start_date=start_date, end_date=end_date, region=None)
    if not filtered:
        await safe_reply(update, "За выбранный период записей не найдено.")
        return

    await safe_chat_action(context.bot, update.effective_chat.id, ChatAction.TYPING)

    totals: Dict[str, Dict[str, float]] = {}
    for row in filtered:
        region_name = row[4] or "Не указано"
        entry = totals.setdefault(region_name, {"count": 0, "total": 0.0})
        entry["count"] += 1
        try:
            entry["total"] += float(row[7] or 0)
        except (TypeError, ValueError):
            continue

    period_text = f"{start_date.strftime('%d.%m.%Y')} — {end_date.strftime('%d.%m.%Y')}"
    lines = []
    for region_name, stats in sorted(totals.items(), key=lambda item: item[1]["total"], reverse=True):
        count = int(stats["count"])
        total_val = int(stats["total"])
        average = int(stats["total"] / stats["count"]) if stats["count"] else 0
        lines.append(f"  {region_name}: {count} предмет(ов), сумма {total_val}, средняя {average}")

    await safe_reply(
        update,
        "📊 Сводка по регионам:\n"
        f"Диапазон: {period_text}\n"
        f"\n".join(lines)
    )


async def reports_start_handler(update: Update, context: CallbackContext) -> int:
    if not is_admin(update.message.from_user.id):
        await safe_reply(update, "Недостаточно прав для работы с отчётами.")
        return ConversationHandler.END

    _report_data(context)
    message = await safe_reply(
        update,
        "Выберите действие для отчёта:",
        reply_markup=build_reports_action_keyboard()
    )
    if message:
        context.user_data["reports_menu_message"] = {
            "chat_id": update.effective_chat.id,
            "message_id": message.message_id,
        }
    return ReportState.ACTION


async def reports_action_handler(update: Update, context: CallbackContext) -> int:
    menu_prompt = context.user_data.pop("reports_menu_message", None)
    if menu_prompt:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=menu_prompt.get("chat_id"),
                message_id=menu_prompt.get("message_id"),
                reply_markup=None,
            )
        except (TelegramError, TypeError, KeyError):
            pass
    text = update.message.text.strip().lower()
    if "отмена" in text:
        return await _reports_finish(update, context, "❌ Отчёт отменён.")

    report = _report_data(context)
    if "архив" in text:
        report.clear()
        report["type"] = "archive"
        await safe_reply(
            update,
            "Введите месяц в формате 03.2025.",
            reply_markup=build_step_inline_keyboard()
        )
        return ReportState.MONTH_INPUT
    if "выгруз" in text or "меся" in text:
        report.clear()
        report["type"] = "month"
        await safe_reply(
            update,
            "Введите месяц в формате 03.2025.",
            reply_markup=build_step_inline_keyboard()
        )
        return ReportState.MONTH_INPUT
    if "свод" in text or ("регион" in text and "стат" not in text):
        report.clear()
        report["type"] = "region_summary"
        await safe_reply(
            update,
            "Введите дату начала периода (например 01.03.2025).",
            reply_markup=build_step_inline_keyboard()
        )
        return ReportState.PERIOD_START
    if "статист" in text or "период" in text:
        report.clear()
        report["type"] = "period"
        await safe_reply(
            update,
            "Введите дату начала периода (например 01.03.2025).",
            reply_markup=build_step_inline_keyboard()
        )
        return ReportState.PERIOD_START

    await safe_reply(
        update,
        "Выберите действие из меню отчётов.",
        reply_markup=build_reports_action_keyboard()
    )
    return ReportState.ACTION


async def reports_month_input_handler(update: Update, context: CallbackContext) -> int:
    text = update.message.text.strip()
    if "отмена" in text.lower():
        return await _reports_finish(update, context, "❌ Отчёт отменён.")

    bounds = get_month_bounds(text)
    if not bounds:
        await safe_reply(
            update,
            "Укажите месяц как 03.2025.",
            reply_markup=build_step_inline_keyboard()
        )
        return ReportState.MONTH_INPUT

    report = _report_data(context)
    report["month_text"] = text
    report["start_date"], report["end_date"] = bounds
    region_prompt = await safe_reply(
        update,
        "Выберите регион или 'Все регионы':",
        reply_markup=build_region_inline_keyboard(include_all=True, prefix=REPORT_REGION_CALLBACK_PREFIX)
    )
    if region_prompt:
        context.user_data["report_region_prompt"] = {
            "chat_id": update.effective_chat.id,
            "message_id": region_prompt.message_id,
        }
    return ReportState.MONTH_REGION


async def reports_month_region_handler(update: Update, context: CallbackContext) -> int:
    prompt = context.user_data.pop("report_region_prompt", None)
    if prompt:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=prompt.get("chat_id"),
                message_id=prompt.get("message_id"),
                reply_markup=None,
            )
        except (TelegramError, TypeError, KeyError):
            pass
    text = update.message.text.strip().lower()
    if "отмена" in text:
        return await _reports_finish(update, context, "❌ Отчёт отменён.")

    region: Optional[str]
    if "все" in text:
        region = None
    else:
        region = normalize_region_input(update.message.text)
        if not region:
            await safe_reply(
                update,
                "Выберите регион из списка или 'Все регионы'.",
                reply_markup=build_region_inline_keyboard(include_all=True, prefix=REPORT_REGION_CALLBACK_PREFIX)
            )
            return ReportState.MONTH_REGION

    report = _report_data(context)
    month_text = report.get("month_text")
    start_date = report.get("start_date")
    end_date = report.get("end_date")

    if not month_text or not start_date or not end_date:
        return await _reports_finish(update, context, "❌ Ошибка состояния отчёта. Попробуйте ещё раз.")

    report_type = report.get("type", "month")
    if report_type == "archive":
        await send_month_archive(update, context, month_text, start_date, end_date, region)
    else:
        await send_month_report(update, context, month_text, start_date, end_date, region)
    return await _reports_finish(update, context, "Готово. Возвращаюсь в меню.")


async def reports_period_start_handler(update: Update, context: CallbackContext) -> int:
    text = update.message.text.strip()
    if "отмена" in text.lower():
        return await _reports_finish(update, context, "❌ Отчёт отменён.")

    start_date = parse_date_str(text)
    if not start_date:
        await safe_reply(
            update,
            "❗ Неверный формат. Используйте ДД.ММ.ГГГГ (например 01.03.2025).",
            reply_markup=build_step_inline_keyboard()
        )
        return ReportState.PERIOD_START

    report = _report_data(context)
    report["start_date"] = start_date
    await safe_reply(
        update,
        "Введите дату окончания периода в формате ДД.ММ.ГГГГ:",
        reply_markup=build_step_inline_keyboard()
    )
    return ReportState.PERIOD_END


async def reports_period_end_handler(update: Update, context: CallbackContext) -> int:
    text = update.message.text.strip()
    if "отмена" in text.lower():
        return await _reports_finish(update, context, "❌ Отчёт отменён.")

    end_date = parse_date_str(text)
    if not end_date:
        await safe_reply(
            update,
            "❗ Неверный формат. Используйте ДД.ММ.ГГГГ (например 31.03.2025).",
            reply_markup=build_step_inline_keyboard()
        )
        return ReportState.PERIOD_END

    report = _report_data(context)
    report["end_date"] = end_date
    report_type = report.get("type")
    start_date = report.get("start_date")
    if not start_date:
        return await _reports_finish(update, context, "❌ Ошибка состояния отчёта. Попробуйте ещё раз.")
    if report_type == "region_summary":
        await send_region_summary(update, start_date, end_date)
        return await _reports_finish(update, context, "Готово. Возвращаюсь в меню.")

    region_prompt = await safe_reply(
        update,
        "Выберите регион для фильтра или 'Все регионы':",
        reply_markup=build_region_inline_keyboard(include_all=True, prefix=REPORT_REGION_CALLBACK_PREFIX)
    )
    if region_prompt:
        context.user_data["report_region_prompt"] = {
            "chat_id": update.effective_chat.id,
            "message_id": region_prompt.message_id,
        }
    return ReportState.PERIOD_REGION


async def reports_period_region_handler(update: Update, context: CallbackContext) -> int:
    text = update.message.text.strip().lower()
    if "отмена" in text:
        return await _reports_finish(update, context, "❌ Отчёт отменён.")

    region: Optional[str]
    if "все" in text:
        region = None
    else:
        region = normalize_region_input(update.message.text)
        if not region:
            await safe_reply(
                update,
                "Выберите регион из списка или 'Все регионы'.",
                reply_markup=build_region_inline_keyboard(include_all=True, prefix=REPORT_REGION_CALLBACK_PREFIX)
            )
            return ReportState.PERIOD_REGION

    report = _report_data(context)
    start_date = report.get("start_date")
    end_date = report.get("end_date")

    if not start_date or not end_date:
        return await _reports_finish(update, context, "❌ Ошибка состояния отчёта. Попробуйте ещё раз.")

    if start_date > end_date:
        start_date, end_date = end_date, start_date

    await send_period_stats(update, start_date, end_date, region)
    return await _reports_finish(update, context, "Готово. Возвращаюсь в меню.")


async def reports_cancel_handler(update: Update, context: CallbackContext) -> int:
    return await _reports_finish(update, context, "❌ Отчёт отменён.")


async def achievements_handler(update: Update, context: CallbackContext) -> None:
    user = update.effective_user
    if not user:
        await safe_reply(update, "Не удалось определить пользователя.")
        return
    if await is_user_blocked(user.id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return

    payload = await build_achievements_view(user.id, ACHIEVEMENTS_DEFAULT_VIEW)
    if not payload:
        await safe_reply(update, "Пока нет данных. Оформите первое заключение.")
        return

    message = await safe_reply(
        update,
        payload["text"],
        reply_markup=payload["keyboard"],
        disable_web_page_preview=True,
    )
    if message:
        context.user_data[ACHIEVEMENTS_PANEL_KEY] = {
            "chat_id": message.chat_id,
            "message_id": message.message_id,
            "view": payload["view"],
        }


async def achievements_callback_handler(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(ACHIEVEMENTS_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return

    action = query.data[len(ACHIEVEMENTS_CALLBACK_PREFIX):]
    if action == "close":
        await query.answer("Панель закрыта.")
        try:
            await query.edit_message_text("Панель достижений закрыта.")
        except TelegramError:
            pass
        context.user_data.pop(ACHIEVEMENTS_PANEL_KEY, None)
        return

    user = query.from_user
    if not user:
        await query.answer("Не удалось определить пользователя.", show_alert=True)
        return

    payload = await build_achievements_view(user.id, action or ACHIEVEMENTS_DEFAULT_VIEW)
    if not payload:
        await query.answer("Пока нет данных.", show_alert=True)
        return

    try:
        await query.edit_message_text(
            payload["text"],
            reply_markup=payload["keyboard"],
            disable_web_page_preview=True,
        )
    except TelegramError as error:
        logger.debug(f"Не удалось обновить панель достижений: {error}")
    finally:
        try:
            await query.answer()
        except TelegramError:
            pass

    context.user_data[ACHIEVEMENTS_PANEL_KEY] = {
        "chat_id": query.message.chat_id if query.message else None,
        "message_id": query.message.message_id if query.message else None,
        "view": payload["view"],
    }


async def begin_new_conclusion(target: Any, context: CallbackContext, user_id: int, username: str) -> int:
    existing_data = await load_user_data_from_db(user_id)
    if existing_data.photo_desc:
        await asyncio.to_thread(cleanup_user_photos, existing_data)

    await delete_user_data_from_db(user_id)
    blank = ConclusionData()
    await save_user_data_to_db(user_id, blank)
    await clear_draft(user_id)
    metrics_reset(context)
    metrics_enter_state(context, DialogState.DEPARTMENT)
    await persist_draft(context, user_id, blank, DialogState.DEPARTMENT)

    clear_state_stack(context)
    set_current_state(context, DialogState.DEPARTMENT)
    set_resume_state(context, None)
    context.user_data.pop("region_prompt", None)
    context.user_data.pop("summary_message_id", None)
    context.user_data.pop("add_photo_message", None)
    context.user_data.pop("mode_prompt", None)
    context.user_data.pop("draft_prompt_message", None)
    context.user_data.pop("draft_decision_context", None)
    context.user_data.pop("pending_draft", None)
    context.user_data.pop("draft_discard_label", None)
    context.user_data["awaiting_draft_choice"] = False

    await _send_via_target(
        target,
        context,
        f"Добро пожаловать, {username}.",
        reply_markup=ReplyKeyboardRemove(),
    )
    await _send_via_target(
        target,
        context,
        f"{format_progress('department')}\nВведите номер подразделения (например 385).\nНажмите «Отмена», чтобы прервать оформление.",
        reply_markup=build_step_inline_keyboard(context=context),
    )
    return DialogState.DEPARTMENT


async def start_handler(update: Update, context: CallbackContext) -> int:
    user_id = update.message.from_user.id
    username = update.message.from_user.full_name
    if not await ensure_user_not_blocked_message(update, context):
        return ConversationHandler.END
    draft = await fetch_draft(user_id)
    if draft and draft.get("data"):
        context.user_data["awaiting_draft_choice"] = True
        context.user_data["draft_decision_context"] = "start"
        context.user_data["pending_draft"] = draft
        context.user_data["draft_discard_label"] = "Удалить и начать заново"
        message = await safe_reply(
            update,
            build_draft_summary_text(draft) + "\n\nХотите продолжить или начать заново?",
            reply_markup=build_draft_keyboard(discard_label=context.user_data["draft_discard_label"])
        )
        if message:
            context.user_data["draft_prompt_message"] = {
                "chat_id": message.chat_id,
                "message_id": message.message_id,
            }
        return DialogState.DEPARTMENT

    return await begin_new_conclusion(update, context, user_id, username)

async def get_department(update: Update, context: CallbackContext) -> int:
    set_current_state(context, DialogState.DEPARTMENT)
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.DEPARTMENT
    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END
    department_number = update.message.text.strip()
    validation = validate_department(department_number)
    if not validation.ok:
        await emit_validation_error(update, context, validation, "department")
        return DialogState.DEPARTMENT
    data = await load_user_data_from_db(user_id)
    data.department_number = department_number
    await save_user_data_to_db(user_id, data)
    resume_state = get_resume_state(context)
    next_state = resume_state or DialogState.ISSUE_NUMBER
    metrics_complete_state(context, DialogState.DEPARTMENT)
    await persist_draft(context, user_id, data, next_state)
    push_state(context, DialogState.DEPARTMENT)
    await safe_reply(update, f"{format_progress('department')} Номер сохранён.")
    if resume_state:
        set_resume_state(context, None)
        await prompt_for_state(update, context, data, resume_state)
        return resume_state
    metrics_enter_state(context, DialogState.ISSUE_NUMBER)
    await _send_via_target(
        update,
        context,
        "Введите порядковый номер заключения за день (например 1).",
        reply_markup=build_step_inline_keyboard(context=context)
    )
    return DialogState.ISSUE_NUMBER

async def get_issue_number(update: Update, context: CallbackContext) -> int:
    set_current_state(context, DialogState.ISSUE_NUMBER)
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.ISSUE_NUMBER
    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END
    issue_number = update.message.text.strip()
    validation = validate_issue_number(issue_number)
    if not validation.ok:
        await emit_validation_error(update, context, validation, "issue")
        return DialogState.ISSUE_NUMBER
    data = await load_user_data_from_db(user_id)
    data.issue_number = issue_number
    await save_user_data_to_db(user_id, data)
    resume_state = get_resume_state(context)
    next_state = resume_state or DialogState.TICKET_NUMBER
    metrics_complete_state(context, DialogState.ISSUE_NUMBER)
    await persist_draft(context, user_id, data, next_state)
    push_state(context, DialogState.ISSUE_NUMBER)
    await safe_reply(update, f"{format_progress('issue')} Номер сохранён.")
    if resume_state:
        set_resume_state(context, None)
        await prompt_for_state(update, context, data, resume_state)
        return resume_state
    metrics_enter_state(context, DialogState.TICKET_NUMBER)
    await _send_via_target(
        update,
        context,
        f"Введите номер залогового билета (например 01230004567, {ticket_digits_phrase()}).",
        reply_markup=build_step_inline_keyboard(context=context)
    )
    return DialogState.TICKET_NUMBER

async def get_ticket_number(update: Update, context: CallbackContext) -> int:
    set_current_state(context, DialogState.TICKET_NUMBER)
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.TICKET_NUMBER
    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END
    ticket_number = update.message.text.strip()
    validation = validate_ticket_number(ticket_number)
    if not validation.ok:
        await emit_validation_error(update, context, validation, "ticket")
        return DialogState.TICKET_NUMBER
    data = await load_user_data_from_db(user_id)
    data.ticket_number = ticket_number
    await save_user_data_to_db(user_id, data)
    resume_state = get_resume_state(context)
    next_state = resume_state or DialogState.DATE
    metrics_complete_state(context, DialogState.TICKET_NUMBER)
    await persist_draft(context, user_id, data, next_state)
    push_state(context, DialogState.TICKET_NUMBER)
    await safe_reply(update, f"{format_progress('ticket')} Номер сохранён.")
    if resume_state:
        set_resume_state(context, None)
        await prompt_for_state(update, context, data, resume_state)
        return resume_state
    metrics_enter_state(context, DialogState.DATE)
    await _send_via_target(
        update,
        context,
        "Введите дату заключения в формате 01.03.2025.",
        reply_markup=build_step_inline_keyboard(context=context)
    )
    return DialogState.DATE

async def get_date(update: Update, context: CallbackContext) -> int:
    set_current_state(context, DialogState.DATE)
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.DATE
    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END
    date_text = update.message.text.strip()
    validation = validate_date_text(date_text)
    if not validation.ok:
        await emit_validation_error(update, context, validation, "date")
        return DialogState.DATE
    data = await load_user_data_from_db(user_id)
    data.date = date_text
    await save_user_data_to_db(user_id, data)
    resume_state = get_resume_state(context)
    next_state = resume_state or DialogState.REGION
    metrics_complete_state(context, DialogState.DATE)
    await persist_draft(context, user_id, data, next_state)
    push_state(context, DialogState.DATE)
    await safe_reply(update, f"{format_progress('date')} Дата сохранена.")
    if resume_state:
        set_resume_state(context, None)
        await prompt_for_state(update, context, data, resume_state)
        return resume_state
    region_message = await _send_via_target(
        update,
        context,
        "Выберите регион:",
        reply_markup=build_region_inline_keyboard(context=context)
    )
    metrics_enter_state(context, DialogState.REGION)
    if region_message:
        context.user_data["region_prompt"] = {
            "chat_id": region_message.chat_id,
            "message_id": region_message.message_id,
        }
    return DialogState.REGION

async def get_region(update: Update, context: CallbackContext) -> int:
    set_current_state(context, DialogState.REGION)
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.REGION
    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END
    prompt_info = context.user_data.pop("region_prompt", None)
    if prompt_info:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=prompt_info.get("chat_id"),
                message_id=prompt_info.get("message_id"),
                reply_markup=None,
            )
        except (TelegramError, TypeError, KeyError):
            pass
    region_text = update.message.text.strip().split(" ", 1)[-1]
    if region_text not in REGION_TOPICS:
        await safe_reply(
            update,
            "Выберите регион из списка.",
            reply_markup=build_region_inline_keyboard(context=context)
        )
        return DialogState.REGION
    data = await load_user_data_from_db(user_id)
    data.region = region_text
    await save_user_data_to_db(user_id, data)
    resume_state = get_resume_state(context)
    next_state = resume_state or DialogState.PHOTO
    metrics_complete_state(context, DialogState.REGION)
    await persist_draft(context, user_id, data, next_state)
    push_state(context, DialogState.REGION)
    photo_count = len(data.photo_desc)
    await safe_reply(update, f"{format_progress('region')} Регион сохранён.")
    if resume_state:
        set_resume_state(context, None)
        await prompt_for_state(update, context, data, resume_state)
        return resume_state
    await _send_via_target(
        update,
        context,
        f"Отправьте фото предмета.\n{PHOTO_REQUIREMENTS_MESSAGE}\n\n(Загружено: {photo_count}/{MAX_PHOTOS})",
        reply_markup=build_step_inline_keyboard(context=context)
    )
    metrics_enter_state(context, DialogState.PHOTO)
    return DialogState.PHOTO

async def photo_handler(update: Update, context: CallbackContext) -> int:
    set_current_state(context, DialogState.PHOTO)
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.PHOTO
    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END

    message = update.message
    file_entity = None
    if message.photo:
        file_entity = message.photo[-1]
    elif message.document and getattr(message.document, "mime_type", "").startswith("image/"):
        file_entity = message.document
    elif message.video or message.animation:
        await safe_reply(
            update,
            "❗ Видео или GIF не подходят. Пришлите фото в формате JPG/PNG.\n\n"
            f"{PHOTO_REQUIREMENTS_MESSAGE}",
            reply_markup=build_step_inline_keyboard(context=context)
        )
        return DialogState.PHOTO

    if not file_entity:
        await safe_reply(
            update,
            f"❗ На этом шаге принимаются только фото (JPG/PNG).\n\n{PHOTO_REQUIREMENTS_MESSAGE}",
            reply_markup=build_step_inline_keyboard(context=context)
        )
        return DialogState.PHOTO

    data = await load_user_data_from_db(user_id)

    if len(data.photo_desc) >= MAX_PHOTOS:
        await safe_reply(
            update,
            f"❗ Достигнут лимит в {MAX_PHOTOS} фото.",
            reply_markup=build_step_inline_keyboard(context=context)
        )
        return DialogState.PHOTO

    max_size_bytes = MAX_PHOTO_SIZE_MB * 1024 * 1024
    size_bytes = getattr(file_entity, 'file_size', None)
    if size_bytes and size_bytes > max_size_bytes:
        await safe_reply(
            update,
            f"❗ Фото весит {format_filesize(size_bytes)}, лимит — {MAX_PHOTO_SIZE_MB} МБ.\n\n{PHOTO_REQUIREMENTS_MESSAGE}",
            reply_markup=build_step_inline_keyboard(context=context)
        )
        return DialogState.PHOTO

    try:
        file = await file_entity.get_file()
        file_size = getattr(file, "file_size", None)
        if file_size and file_size > max_size_bytes:
            await safe_reply(
                update,
                f"❗ Фото весит {format_filesize(file_size)}, лимит — {MAX_PHOTO_SIZE_MB} МБ.\n\n{PHOTO_REQUIREMENTS_MESSAGE}",
                reply_markup=build_step_inline_keyboard(context=context)
            )
            return DialogState.PHOTO
    except TelegramError as e:
        logger.error(f"Не удалось получить файл изображения: {e}", exc_info=True)
        await safe_reply(
            update,
            "❌ Ошибка получения файла. Попробуйте отправить фото ещё раз.\n\n"
            f"{PHOTO_REQUIREMENTS_MESSAGE}",
            reply_markup=build_step_inline_keyboard(context=context)
        )
        return DialogState.PHOTO

    TEMP_PHOTOS_DIR.mkdir(parents=True, exist_ok=True)

    unique_name = generate_unique_filename()
    original_path = TEMP_PHOTOS_DIR / f"orig_{unique_name}"
    compressed_path = TEMP_PHOTOS_DIR / unique_name
    try:
        await file.download_to_drive(original_path)
        if is_image_too_large(original_path, max_size_mb=MAX_PHOTO_SIZE_MB):
            await safe_reply(
                update,
                f"❗ Фото весит {format_filesize(original_path.stat().st_size)}, лимит — {MAX_PHOTO_SIZE_MB} МБ.\n\n{PHOTO_REQUIREMENTS_MESSAGE}",
                reply_markup=build_step_inline_keyboard(context=context)
            )
            return DialogState.PHOTO
        await asyncio.to_thread(compress_image, original_path, compressed_path)
    except Exception as e:
        logger.error(f"Ошибка обработки изображения: {e}", exc_info=True)
        await safe_reply(
            update,
            "❌ Не удалось обработать изображение. Попробуйте снова.\n\n"
            f"{PHOTO_REQUIREMENTS_MESSAGE}",
            reply_markup=build_step_inline_keyboard(context=context)
        )
        if compressed_path.exists():
            try:
                compressed_path.unlink()
            except OSError as cleanup_error:
                logger.warning(f"Не удалось удалить временный файл {compressed_path}: {cleanup_error}")
        return DialogState.PHOTO
    finally:
        if original_path.exists():
            try:
                original_path.unlink()
            except OSError as cleanup_error:
                logger.warning(f"Не удалось удалить исходный файл {original_path}: {cleanup_error}")

    data.photo_desc.append({'photo': str(compressed_path), 'description': '', 'evaluation': ''})
    await save_user_data_to_db(user_id, data)
    metrics_complete_state(context, DialogState.PHOTO)
    await persist_draft(context, user_id, data, DialogState.DESCRIPTION)
    push_state(context, DialogState.PHOTO)

    await safe_reply(update, f"{format_progress('photo')} Фото сохранено.")
    await safe_reply(
        update,
        "Добавьте короткое описание предмета.",
        reply_markup=build_step_inline_keyboard(context=context)
    )
    metrics_enter_state(context, DialogState.DESCRIPTION)
    return DialogState.DESCRIPTION

async def description_handler(update: Update, context: CallbackContext) -> int:
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.DESCRIPTION
    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END
    description = (update.message.text or "").strip()
    if len(description) < 3:
        await safe_reply(
            update,
            f"{format_progress('description')} Описание слишком короткое. Пару слов о предмете — например «Кольцо из белого золота».",
            reply_markup=build_step_inline_keyboard(context=context)
        )
        return DialogState.DESCRIPTION
    data = await load_user_data_from_db(user_id)
    if data.photo_desc:
        data.photo_desc[-1]['description'] = description
    await save_user_data_to_db(user_id, data)
    metrics_complete_state(context, DialogState.DESCRIPTION)
    await persist_draft(context, user_id, data, DialogState.EVALUATION)
    await safe_reply(update, f"{format_progress('description')} Описание сохранено.")
    await safe_reply(
        update,
        "Введите оценку предмета (целое число, например 1500).",
        reply_markup=build_step_inline_keyboard()
    )
    metrics_enter_state(context, DialogState.EVALUATION)
    return DialogState.EVALUATION

async def evaluation_handler(update: Update, context: CallbackContext) -> int:
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.EVALUATION
    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END
    evaluation_text = update.message.text.strip()

    validation = validate_evaluation(evaluation_text)
    if not validation.ok:
        await emit_validation_error(update, context, validation, "evaluation", keyboard=build_step_inline_keyboard())
        return DialogState.EVALUATION

    data = await load_user_data_from_db(user_id)
    if data.photo_desc:
        data.photo_desc[-1]['evaluation'] = evaluation_text
    await save_user_data_to_db(user_id, data)

    metrics_complete_state(context, DialogState.EVALUATION)
    await persist_draft(context, user_id, data, DialogState.MORE_PHOTO)

    photo_count = len(data.photo_desc)
    buttons = [[
        InlineKeyboardButton("Добавить фото", callback_data=f"{ADD_PHOTO_PREFIX}yes"),
        InlineKeyboardButton("Перейти к сводке", callback_data=f"{ADD_PHOTO_PREFIX}no"),
    ]]
    message = await safe_reply(
        update,
        f"{format_progress('evaluation')} Оценка сохранена. Добавить ещё одно фото? ({photo_count}/{MAX_PHOTOS}).",
        reply_markup=build_step_inline_keyboard(buttons)
    )
    if message:
        context.user_data["add_photo_message"] = {"chat_id": update.effective_chat.id, "message_id": message.message_id}
    metrics_enter_state(context, DialogState.MORE_PHOTO)
    return DialogState.MORE_PHOTO

async def more_photo_handler(update: Update, context: CallbackContext) -> int:
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.MORE_PHOTO
    message_info = context.user_data.pop("add_photo_message", None)
    if message_info:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=message_info["chat_id"],
                message_id=message_info["message_id"],
                reply_markup=None,
            )
        except (TelegramError, KeyError, TypeError):
            pass

    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END

    user_response = (update.message.text or "").strip().lower()
    if "да" in user_response or "ещё" in user_response:
        data = await load_user_data_from_db(user_id)
        photo_count = len(data.photo_desc)
        metrics_complete_state(context, DialogState.MORE_PHOTO)
        await persist_draft(context, user_id, data, DialogState.PHOTO)
        await safe_reply(update, f"{format_progress('photo')} Добавим ещё одно фото.")
        await safe_reply(
            update,
            f"Отправьте следующее фото.\n{PHOTO_REQUIREMENTS_MESSAGE}\n\n(Загружено: {photo_count}/{MAX_PHOTOS})",
            reply_markup=build_step_inline_keyboard()
        )
        metrics_enter_state(context, DialogState.PHOTO)
        return DialogState.PHOTO

    data = await load_user_data_from_db(user_id)
    metrics_complete_state(context, DialogState.MORE_PHOTO)
    await persist_draft(context, user_id, data, DialogState.CONFIRMATION)
    await safe_reply(update, f"{format_progress('summary')} Данные собраны.")
    await show_summary(update, context, data)
    return DialogState.CONFIRMATION

async def confirmation_handler(update: Update, context: CallbackContext) -> int:
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.CONFIRMATION
    if not await ensure_user_not_blocked_message(update, context):
        return ConversationHandler.END
    summary_message_id = context.user_data.pop("summary_message_id", None)
    if summary_message_id:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=update.effective_chat.id,
                message_id=summary_message_id,
                reply_markup=None,
            )
        except TelegramError:
            pass
    user_response = update.message.text.strip().lower()
    if "да" in user_response:
        user = update.message.from_user
        data_snapshot = None
        if user:
            data_snapshot = await load_user_data_from_db(user.id)
        metrics_complete_state(context, DialogState.CONFIRMATION)
        if user and data_snapshot:
            await persist_draft(context, user.id, data_snapshot, DialogState.TESTING)
        await safe_reply(update, f"{format_progress('summary')} Данные подтверждены.")
        mode_message = await safe_reply(update, "Выберите режим создания заключения.", reply_markup=build_mode_keyboard())
        if mode_message:
            context.user_data["mode_prompt"] = {
                "chat_id": update.effective_chat.id,
                "message_id": mode_message.message_id,
            }
        metrics_enter_state(context, DialogState.TESTING)
        return DialogState.TESTING

    metrics_complete_state(context, DialogState.CONFIRMATION)
    await safe_reply(update, "Процесс остановлен.")
    await safe_reply(update, "Чтобы начать заново, введите /start.")
    return ConversationHandler.END


async def confirmation_callback_handler(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    if not query or not query.data:
        return DialogState.CONFIRMATION
    if context.user_data.get("awaiting_draft_choice"):
        await query.answer("Сначала выберите действие с черновиком.", show_alert=True)
        return DialogState.CONFIRMATION
    if not await ensure_user_not_blocked_query(query, context):
        return DialogState.CONFIRMATION
    if not await ensure_user_not_blocked_query(query, context):
        return DialogState.CONFIRMATION

    action = query.data
    if not action.startswith(CONFIRM_CALLBACK_PREFIX):
        await query.answer()
        return DialogState.CONFIRMATION

    action_suffix = action[len(CONFIRM_CALLBACK_PREFIX):]
    thread_id = getattr(query.message, "message_thread_id", None)
    chat_id = query.message.chat_id

    # Remove stored keyboard reference
    summary_message_id = context.user_data.pop("summary_message_id", None)
    if summary_message_id and summary_message_id != query.message.message_id:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=chat_id,
                message_id=summary_message_id,
                reply_markup=None,
            )
        except TelegramError:
            pass

    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    if action_suffix == "next":
        await query.answer("Данные подтверждены.")
        metrics_complete_state(context, DialogState.CONFIRMATION)
        if query.from_user:
            snapshot = await load_user_data_from_db(query.from_user.id)
            await persist_draft(context, query.from_user.id, snapshot, DialogState.TESTING)
        await safe_bot_send_message(
            context.bot,
            chat_id,
            f"{format_progress('summary')} Данные подтверждены.",
            message_thread_id=thread_id,
        )
        mode_prompt_message = await safe_bot_send_message(
            context.bot,
            chat_id,
            "Выберите режим создания заключения.",
            reply_markup=build_mode_keyboard(),
            message_thread_id=thread_id,
        )
        if mode_prompt_message:
            context.user_data["mode_prompt"] = {
                "chat_id": chat_id,
                "message_id": mode_prompt_message.message_id,
            }
        metrics_enter_state(context, DialogState.TESTING)
        return DialogState.TESTING

    if action_suffix == "cancel":
        await query.answer("Процесс остановлен.")
        user_id = query.from_user.id if query.from_user else None
        context.user_data.pop("mode_prompt", None)
        metrics_complete_state(context, DialogState.CONFIRMATION)
        if user_id is not None:
            data = await load_user_data_from_db(user_id)
            if data.photo_desc:
                await asyncio.to_thread(cleanup_user_photos, data)
            await delete_user_data_from_db(user_id)
        await safe_bot_send_message(
            context.bot,
            chat_id,
            "Процесс остановлен.",
            message_thread_id=thread_id,
        )
        await safe_bot_send_message(
            context.bot,
            chat_id,
            "Чтобы начать заново, введите /start.",
            message_thread_id=thread_id,
        )
        return ConversationHandler.END

    await query.answer()
    return DialogState.CONFIRMATION


async def edit_callback_handler(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(EDIT_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return DialogState.CONFIRMATION
    if context.user_data.get("awaiting_draft_choice"):
        await query.answer("Сначала выберите действие с черновиком.", show_alert=True)
        return DialogState.CONFIRMATION

    action = query.data[len(EDIT_CALLBACK_PREFIX):]
    message = query.message
    chat_id = message.chat_id if message else None
    thread_id = getattr(message, "message_thread_id", None) if message else None
    user = query.from_user
    user_id = user.id if user else None

    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    context.user_data.pop("summary_message_id", None)

    if action.startswith("field:"):
        field_key = action.split(":", 1)[1]
        config = EDIT_FIELD_TARGETS.get(field_key)
        if not config:
            await query.answer("Поле недоступно для редактирования.", show_alert=True)
            return DialogState.CONFIRMATION
        if chat_id is None:
            await query.answer()
            return DialogState.CONFIRMATION

        # подготовка к редактированию поля
        metrics_complete_state(context, DialogState.CONFIRMATION)
        set_resume_state(context, DialogState.CONFIRMATION)
        target_state: DialogState = config["state"]
        set_current_state(context, target_state)
        metrics_enter_state(context, target_state)
        prompt_text = f"{format_progress(config['stage'])} {config['prompt']}"

        if field_key == "region":
            region_message = await safe_bot_send_message(
                context.bot,
                chat_id,
                prompt_text,
                reply_markup=build_region_inline_keyboard(context=context),
                message_thread_id=thread_id,
            )
            if region_message:
                context.user_data["region_prompt"] = {
                    "chat_id": chat_id,
                    "message_id": region_message.message_id,
                }
        else:
            await safe_bot_send_message(
                context.bot,
                chat_id,
                prompt_text,
                reply_markup=build_step_inline_keyboard(context=context),
                message_thread_id=thread_id,
            )
        await query.answer("Введите новое значение.")
        return target_state

    if action == "delete_photo":
        if user_id is None:
            await query.answer()
            return DialogState.CONFIRMATION
        data = await load_user_data_from_db(user_id)
        if not data.photo_desc:
            await query.answer("Фото отсутствуют.", show_alert=True)
            return DialogState.CONFIRMATION

        removed_item = data.photo_desc.pop()
        photo_path = Path(removed_item.get("photo", "") or "")
        if photo_path.is_file():
            try:
                photo_path.unlink()
            except OSError as error:
                logger.warning(f"Не удалось удалить фото {photo_path}: {error}")
            parent_dir = photo_path.parent
            if parent_dir != TEMP_PHOTOS_DIR and parent_dir.exists():
                try:
                    if not any(parent_dir.iterdir()):
                        parent_dir.rmdir()
                except OSError:
                    pass

        await save_user_data_to_db(user_id, data)
        await persist_draft(context, user_id, data, DialogState.CONFIRMATION)
        await query.answer("Последнее фото удалено.")
        await safe_bot_send_message(
            context.bot,
            chat_id,
            f"{format_progress('photo')} Последнее фото удалено.",
            message_thread_id=thread_id,
        )
        await show_summary(query, context, data)
        return DialogState.CONFIRMATION

    await query.answer()
    return DialogState.CONFIRMATION


async def add_photo_callback_handler(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(ADD_PHOTO_PREFIX):
        if query:
            await query.answer()
        return DialogState.MORE_PHOTO

    if context.user_data.get("awaiting_draft_choice"):
        await query.answer("Сначала выберите действие с черновиком.", show_alert=True)
        return DialogState.MORE_PHOTO

    decision = query.data[len(ADD_PHOTO_PREFIX):]
    message_info = context.user_data.pop("add_photo_message", None)
    if message_info:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=message_info.get("chat_id"),
                message_id=message_info.get("message_id"),
                reply_markup=None,
            )
        except (TelegramError, TypeError, KeyError):
            pass

    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    user_id = query.from_user.id if query.from_user else None
    if user_id is None:
        await query.answer()
        return DialogState.MORE_PHOTO
    if not await ensure_user_not_blocked_query(query, context):
        return DialogState.MORE_PHOTO

    chat_id = query.message.chat_id
    thread_id = getattr(query.message, "message_thread_id", None)

    if decision == "yes":
        await query.answer("Добавим ещё одно фото.")
        data = await load_user_data_from_db(user_id)
        photo_count = len(data.photo_desc)
        metrics_complete_state(context, DialogState.MORE_PHOTO)
        await persist_draft(context, user_id, data, DialogState.PHOTO)
        await safe_bot_send_message(
            context.bot,
            chat_id,
            f"{format_progress('photo')} Добавим ещё одно фото.",
            message_thread_id=thread_id,
        )
        await safe_bot_send_message(
            context.bot,
            chat_id,
            f"Отправьте следующее фото.\n{PHOTO_REQUIREMENTS_MESSAGE}\n\n(Загружено: {photo_count}/{MAX_PHOTOS})",
            message_thread_id=thread_id,
        )
        metrics_enter_state(context, DialogState.PHOTO)
        return DialogState.PHOTO

    if decision == "no":
        await query.answer("Показываю сводку.")
        data = await load_user_data_from_db(user_id)
        metrics_complete_state(context, DialogState.MORE_PHOTO)
        await persist_draft(context, user_id, data, DialogState.CONFIRMATION)
        await safe_bot_send_message(
            context.bot,
            chat_id,
            f"{format_progress('summary')} Данные собраны.",
            message_thread_id=thread_id,
        )
        await show_summary(query, context, data)
        return DialogState.CONFIRMATION

    await query.answer()
    return DialogState.MORE_PHOTO


async def navigation_callback_handler(update: Update, context: CallbackContext) -> int:
    """Handle generic navigation callbacks such as cancel/back/menu within conversations."""
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(NAVIGATION_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return ConversationHandler.END

    action = query.data[len(NAVIGATION_CALLBACK_PREFIX):]
    chat_id = query.message.chat_id if query.message else None
    thread_id = getattr(query.message, "message_thread_id", None)
    user_id = query.from_user.id if query.from_user else None

    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    if action == "cancel":
        if context.user_data.get("report"):
            await query.answer("❌ Отчёт отменён.")
            return await _reports_finish(query, context, "❌ Отчёт отменён.")

        await query.answer("Процесс остановлен.")
        if user_id is not None:
            data = await load_user_data_from_db(user_id)
            if data.photo_desc:
                await asyncio.to_thread(cleanup_user_photos, data)
            await delete_user_data_from_db(user_id)
            await clear_draft(user_id)
        context.user_data.pop("region_prompt", None)
        context.user_data.pop("summary_message_id", None)
        context.user_data.pop("add_photo_message", None)
        context.user_data.pop("mode_prompt", None)
        context.user_data["awaiting_draft_choice"] = False
        if chat_id is not None:
            await safe_bot_send_message(
                context.bot,
                chat_id,
                "Процесс остановлен. Чтобы начать заново, используйте /start.",
                message_thread_id=thread_id,
            )
        return ConversationHandler.END

    if action == "back":
        if context.user_data.get("awaiting_draft_choice"):
            await query.answer(_("Сначала завершите работу с черновиком."), show_alert=True)
            return get_current_state(context) or ConversationHandler.END
        if not await ensure_user_not_blocked_query(query, context):
            return get_current_state(context) or ConversationHandler.END
        current_state = get_current_state(context)
        previous_state = pop_state(context)
        if previous_state is None:
            await query.answer(_("Предыдущий шаг недоступен."), show_alert=True)
            return current_state or ConversationHandler.END

        data = await load_user_data_from_db(user_id) if user_id is not None else ConclusionData()
        origin_state_name = current_state.name if isinstance(current_state, DialogState) else None
        if current_state:
            set_resume_state(context, current_state)

        if previous_state in TEXTUAL_BACK_STATES:
            value = extract_state_value(data, previous_state) or "не указано"
            state_label = get_state_label(previous_state.name)
            context.user_data[PENDING_BACK_DECISION_KEY] = {
                "previous": previous_state.name,
                "origin": origin_state_name,
            }
            set_current_state(context, previous_state)
            decision_keyboard = InlineKeyboardMarkup([
                [
                    InlineKeyboardButton(
                        _("Оставить без изменений"),
                        callback_data=f"{BACK_NAV_CALLBACK_PREFIX}keep:{previous_state.name}",
                    )
                ],
                [
                    InlineKeyboardButton(
                        _("Изменить значение"),
                        callback_data=f"{BACK_NAV_CALLBACK_PREFIX}edit:{previous_state.name}",
                    )
                ],
                [
                    InlineKeyboardButton(
                        _("Отмена"),
                        callback_data=f"{NAVIGATION_CALLBACK_PREFIX}cancel",
                    )
                ],
            ])
            await query.answer(_("Возвращаемся на предыдущий шаг."))
            if chat_id is not None:
                await safe_bot_send_message(
                    context.bot,
                    chat_id,
                    _("{}: текущее значение — {}.\nВыберите действие ниже.").format(state_label, value),
                    reply_markup=decision_keyboard,
                    message_thread_id=thread_id,
                )
            return previous_state

        set_current_state(context, previous_state)
        await query.answer(_("Возвращаемся на предыдущий шаг."))
        await prompt_for_state(query, context, data, previous_state)
        return previous_state

    if action == "menu":
        await query.answer("Откройте главное меню.")
        if chat_id is not None:
            await safe_bot_send_message(
                context.bot,
                chat_id,
                "Используйте команду /menu, чтобы открыть главное меню.",
                message_thread_id=thread_id,
            )
        return ConversationHandler.END

    await query.answer()
    return ConversationHandler.END


async def back_navigation_decision_handler(update: Update, context: CallbackContext) -> int:
    """Handle inline decisions when the user navigates back to a previous state."""
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(BACK_NAV_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return get_current_state(context) or ConversationHandler.END

    if context.user_data.get("awaiting_draft_choice"):
        await query.answer(_("Сначала завершите работу с черновиком."), show_alert=True)
        return get_current_state(context) or ConversationHandler.END

    if not await ensure_user_not_blocked_query(query, context):
        return get_current_state(context) or ConversationHandler.END

    payload = query.data[len(BACK_NAV_CALLBACK_PREFIX):]
    try:
        decision, state_name = payload.split(":", 1)
    except ValueError:
        await query.answer(_("Некорректное действие."), show_alert=True)
        return get_current_state(context) or ConversationHandler.END

    try:
        target_state = DialogState[state_name]
    except KeyError:
        await query.answer(_("Не удалось определить шаг."), show_alert=True)
        return get_current_state(context) or ConversationHandler.END

    pending = context.user_data.get(PENDING_BACK_DECISION_KEY)
    if not pending or pending.get("previous") != state_name:
        await query.answer(_("Действие устарело."), show_alert=True)
        return get_current_state(context) or ConversationHandler.END

    origin_state_name = pending.get("origin")
    origin_state = None
    if origin_state_name:
        try:
            origin_state = DialogState[origin_state_name]
        except KeyError:
            origin_state = None

    user_id = query.from_user.id if query.from_user else None
    data = await load_user_data_from_db(user_id) if user_id is not None else ConclusionData()
    chat_id = query.message.chat_id if query.message else None
    thread_id = getattr(query.message, "message_thread_id", None)

    if decision == "keep":
        push_state(context, target_state)
        set_resume_state(context, None)
        context.user_data.pop(PENDING_BACK_DECISION_KEY, None)
        await query.answer(_("Значение оставлено без изменений."))
        followup_state = origin_state or target_state
        if followup_state:
            await prompt_for_state(query, context, data, followup_state)
            return followup_state
        return get_current_state(context) or ConversationHandler.END

    if decision == "edit":
        context.user_data.pop(PENDING_BACK_DECISION_KEY, None)
        if origin_state:
            set_resume_state(context, origin_state)
        set_current_state(context, target_state)
        await query.answer(_("Измените значение."))
        await prompt_for_state(query, context, data, target_state)
        return target_state

    if chat_id is not None:
        await safe_bot_send_message(
            context.bot,
            chat_id,
            _("Действие не поддерживается."),
            message_thread_id=thread_id,
        )
    await query.answer()
    return get_current_state(context) or ConversationHandler.END


async def draft_callback_handler(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(DRAFT_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return DialogState.DEPARTMENT

    user = query.from_user
    if not user:
        await query.answer("Не удалось определить пользователя.", show_alert=True)
        return DialogState.DEPARTMENT

    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    prompt_message = context.user_data.pop("draft_prompt_message", None)
    if prompt_message:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=prompt_message.get("chat_id"),
                message_id=prompt_message.get("message_id"),
                reply_markup=None,
            )
        except (TypeError, TelegramError):
            pass

    action = query.data[len(DRAFT_CALLBACK_PREFIX):]
    decision_context = context.user_data.get("draft_decision_context")
    chat_id = query.message.chat_id if query.message else None
    thread_id = getattr(query.message, "message_thread_id", None)
    user_id = user.id

    if not await ensure_user_not_blocked_query(query, context):
        return DialogState.DEPARTMENT

    if action == "resume":
        draft = await fetch_draft(user_id)
        if not draft or not draft.get("data"):
            await query.answer("Черновик не найден.", show_alert=True)
            context.user_data["awaiting_draft_choice"] = False
            context.user_data.pop("draft_decision_context", None)
            context.user_data.pop("draft_discard_label", None)
            return DialogState.DEPARTMENT

        context.user_data["awaiting_draft_choice"] = False
        context.user_data.pop("draft_decision_context", None)
        context.user_data.pop("draft_discard_label", None)
        context.user_data["pending_draft"] = draft

        data = draft.get("conclusion") or ConclusionData()
        await save_user_data_to_db(user_id, data)

        state_name = draft.get("state")
        try:
            next_state = DialogState[state_name] if state_name else DialogState.DEPARTMENT
        except (KeyError, TypeError):
            next_state = DialogState.DEPARTMENT

        metrics_restore(context, draft.get("metrics"), next_state)
        await persist_draft(context, user_id, data, next_state)
        clear_state_stack(context)
        set_current_state(context, next_state)
        context.user_data.pop("region_prompt", None)
        context.user_data.pop("summary_message_id", None)
        context.user_data.pop("add_photo_message", None)
        context.user_data.pop("mode_prompt", None)

        await query.answer("Черновик открыт.")
        await prompt_for_state(query, context, data, next_state)
        return next_state

    if action == "discard":
        await clear_draft(user_id)
        context.user_data["awaiting_draft_choice"] = False
        context.user_data.pop("pending_draft", None)
        context.user_data.pop("draft_decision_context", None)
        context.user_data.pop("draft_discard_label", None)
        await query.answer("Черновик удалён.")
        metrics_reset(context)

        if decision_context == "start":
            username = user.full_name
            return await begin_new_conclusion(query, context, user_id, username)

        if chat_id is not None:
            await safe_bot_send_message(
                context.bot,
                chat_id,
                "Черновик удалён.",
                message_thread_id=thread_id,
            )
        return DialogState.DEPARTMENT

    await query.answer()
    return DialogState.DEPARTMENT


async def history_callback_handler(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(HISTORY_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return

    payload = query.data[len(HISTORY_CALLBACK_PREFIX):]
    cache = context.user_data.get("history_cache")
    if not cache:
        await query.answer("Исторические данные отсутствуют. Запросите /history заново.", show_alert=True)
        return

    records: List[List[Any]] = cache.get("records", [])
    page_size: int = cache.get("page_size", 6)
    current_page: int = cache.get("page", 1)

    if payload == "close":
        try:
            await query.edit_message_text("История закрыта.")
        except TelegramError:
            pass
        context.user_data.pop("history_cache", None)
        context.user_data.pop("history_message", None)
        await query.answer("История закрыта.")
        return

    try:
        requested_page = int(payload)
    except ValueError:
        await query.answer()
        return

    text, keyboard, actual_page = build_history_page(records, requested_page, page_size)
    try:
        await query.edit_message_text(text, reply_markup=keyboard)
    except TelegramError as error:
        logger.debug(f"Не удалось обновить вывод истории: {error}")
    else:
        cache["page"] = actual_page
    finally:
        try:
            await query.answer()
        except TelegramError:
            pass


async def region_callback_handler(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(REGION_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return DialogState.REGION

    suffix = query.data[len(REGION_CALLBACK_PREFIX):]
    if suffix == "all":
        await query.answer("Выберите конкретный регион.", show_alert=True)
        return DialogState.REGION

    try:
        index = int(suffix)
        region_name = REGION_CHOICES[index]
    except (ValueError, IndexError):
        await query.answer("Регион не найден.", show_alert=True)
        return DialogState.REGION

    context.user_data.pop("region_prompt", None)
    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    user = query.from_user
    if not user:
        await query.answer("Не удалось определить пользователя.", show_alert=True)
        return DialogState.REGION

    user_id = user.id
    data = await load_user_data_from_db(user_id)
    data.region = region_name
    await save_user_data_to_db(user_id, data)

    photo_count = len(data.photo_desc)
    chat_id = query.message.chat_id
    thread_id = getattr(query.message, "message_thread_id", None)

    await query.answer(f"Выбран: {region_name}")
    await safe_bot_send_message(
        context.bot,
        chat_id,
        f"{format_progress('region')} Регион сохранён.",
        message_thread_id=thread_id,
    )
    await safe_bot_send_message(
        context.bot,
        chat_id,
        f"Отправьте фото предмета.\n{PHOTO_REQUIREMENTS_MESSAGE}\n\n(Загружено: {photo_count}/{MAX_PHOTOS})",
        message_thread_id=thread_id,
        reply_markup=build_step_inline_keyboard()
    )
    return DialogState.PHOTO


async def report_action_callback(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(REPORT_ACTION_PREFIX):
        if query:
            await query.answer()
        return ReportState.ACTION

    prompt = context.user_data.pop("reports_menu_message", None)
    if prompt:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=prompt.get("chat_id"),
                message_id=prompt.get("message_id"),
                reply_markup=None,
            )
        except (TelegramError, TypeError, KeyError):
            pass

    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    action_code = query.data[len(REPORT_ACTION_PREFIX):]
    chat_id = query.message.chat_id
    thread_id = getattr(query.message, "message_thread_id", None)
    report = _report_data(context)
    report.clear()

    if action_code == "cancel":
        return await _reports_finish(query, context, "❌ Отчёт отменён.")

    if action_code == "archive":
        report["type"] = "archive"
        await query.answer("Архив за месяц.")
        await safe_bot_send_message(
            context.bot,
            chat_id,
            "Введите месяц в формате 03.2025.",
            message_thread_id=thread_id,
        )
        return ReportState.MONTH_INPUT

    if action_code == "month":
        report["type"] = "month"
        await query.answer("Выгрузка за месяц.")
        await safe_bot_send_message(
            context.bot,
            chat_id,
            "Введите месяц в формате 03.2025.",
            message_thread_id=thread_id,
        )
        return ReportState.MONTH_INPUT

    if action_code == "period":
        report["type"] = "period"
        await query.answer("Статистика за период.")
        await safe_bot_send_message(
            context.bot,
            chat_id,
            "Введите дату начала периода (например 01.03.2025).",
            message_thread_id=thread_id,
        )
        return ReportState.PERIOD_START

    if action_code == "summary":
        report["type"] = "region_summary"
        await query.answer("Сводка по регионам.")
        await safe_bot_send_message(
            context.bot,
            chat_id,
            "Введите дату начала периода (например 01.03.2025).",
            message_thread_id=thread_id,
        )
        return ReportState.PERIOD_START

    await query.answer()
    return ReportState.ACTION


async def report_region_callback_handler(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(REPORT_REGION_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return ReportState.MONTH_REGION

    prompt = context.user_data.pop("report_region_prompt", None)
    if prompt:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=prompt.get("chat_id"),
                message_id=prompt.get("message_id"),
                reply_markup=None,
            )
        except (TelegramError, TypeError, KeyError):
            pass

    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    suffix = query.data[len(REPORT_REGION_CALLBACK_PREFIX):]
    if suffix == "all":
        region = None
        await query.answer("Выбраны все регионы.")
    else:
        try:
            region = REGION_CHOICES[int(suffix)]
            await query.answer(f"Регион: {region}")
        except (ValueError, IndexError):
            await query.answer("Регион не найден.", show_alert=True)
            return ReportState.MONTH_REGION

    report = _report_data(context)
    month_text = report.get("month_text")
    start_date = report.get("start_date")
    end_date = report.get("end_date")
    report_type = report.get("type", "month")

    if not month_text or not start_date or not end_date:
        return await _reports_finish(query, context, "❌ Ошибка состояния отчёта. Попробуйте ещё раз.")

    if report_type == "archive":
        await send_month_archive(query, context, month_text, start_date, end_date, region)
    else:
        await send_month_report(query, context, month_text, start_date, end_date, region)
    return await _reports_finish(query, context, "Готово. Возвращаюсь в меню.")

async def mode_callback_handler(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    if not query or not query.data or not query.data.startswith(MODE_CALLBACK_PREFIX):
        if query:
            await query.answer()
        return DialogState.TESTING
    if context.user_data.get("awaiting_draft_choice"):
        await query.answer("Сначала выберите действие с черновиком.", show_alert=True)
        return DialogState.TESTING
    if not await ensure_user_not_blocked_query(query, context):
        return DialogState.TESTING

    mode_prompt = context.user_data.pop("mode_prompt", None)
    if mode_prompt:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=mode_prompt.get("chat_id"),
                message_id=mode_prompt.get("message_id"),
                reply_markup=None,
            )
        except (TypeError, TelegramError):
            pass

    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except TelegramError:
        pass

    payload = query.data[len(MODE_CALLBACK_PREFIX):]
    if payload not in ("test", "final"):
        await query.answer("Недоступный режим.", show_alert=True)
        return DialogState.TESTING

    if not query.from_user:
        await query.answer("Не удалось определить пользователя.", show_alert=True)
        return DialogState.TESTING

    is_final = payload == "final"
    mode_label = "Окончательное" if is_final else "Тестовое"
    user_id = query.from_user.id
    username = query.from_user.full_name
    chat_id = query.message.chat_id
    thread_id = getattr(query.message, "message_thread_id", None)

    await query.answer(f"Режим: {mode_label}.")
    await safe_bot_send_message(
        context.bot,
        chat_id,
        f"{format_progress('mode')} Режим выбран: {mode_label}.",
        message_thread_id=thread_id,
    )
    await safe_bot_send_message(
        context.bot,
        chat_id,
        "Создаю документ. Это займёт несколько секунд.",
        message_thread_id=thread_id,
    )

    metrics_complete_state(context, DialogState.TESTING)

    return await mode_flow(context, user_id, username, is_final, chat_id, thread_id)


async def mode_flow(
    context: CallbackContext,
    user_id: int,
    username: str,
    is_final: bool,
    chat_id: int,
    thread_id: Optional[int] = None,
) -> int:
    context.user_data.pop("mode_prompt", None)
    data_for_cleanup = await load_user_data_from_db(user_id)
    metrics_data = metrics_finalize(context)
    processing_time = metrics_processing_time(metrics_data)

    try:
        await safe_chat_action(
            context.bot,
            chat_id,
            ChatAction.UPLOAD_DOCUMENT,
            message_thread_id=thread_id,
        )
        filename_path = await create_document(user_id, username)
        await send_document_from_path(context.bot, chat_id=user_id, path=filename_path)

        if is_final:
            data = data_for_cleanup
            region = data.region if data else ""
            if data and region and region in REGION_TOPICS:
                topic_id = REGION_TOPICS[region]
                await safe_chat_action(
                    context.bot,
                    MAIN_GROUP_CHAT_ID,
                    ChatAction.UPLOAD_DOCUMENT,
                    message_thread_id=topic_id,
                )
                caption = (
                    f"Заключение от п. {data.department_number or 'N/A'}, "
                    f"билет: {data.ticket_number or 'N/A'}, от {data.date or 'N/A'}"
                )
                group_message = await send_document_from_path(
                    context.bot,
                    chat_id=MAIN_GROUP_CHAT_ID,
                    path=filename_path,
                    caption=caption,
                    message_thread_id=topic_id
                )
                try:
                    await update_excel(data)
                except Exception as excel_error:
                    logger.error(f"Ошибка при обновлении Excel: {excel_error}", exc_info=True)
                    await safe_bot_send_message(
                        context.bot,
                        chat_id,
                        "Документ отправлен, но запись в отчёте не обновлена.",
                        message_thread_id=thread_id,
                    )
                else:
                    await safe_bot_send_message(
                        context.bot,
                        chat_id,
                        "Документ отправлен в рабочую группу.",
                        message_thread_id=thread_id,
                    )
                    archive_target: Optional[Path] = None
                    try:
                        archive_target = await archive_document(filename_path, data)
                    except Exception as archive_error:
                        logger.error(f"Не удалось архивировать документ {filename_path}: {archive_error}", exc_info=True)
                    archive_relative: Optional[str] = None
                    if archive_target:
                        try:
                            archive_relative = str(archive_target.relative_to(ARCHIVE_DIR))
                        except ValueError:
                            archive_relative = str(archive_target)
                    try:
                        completion_stats = await record_completion_entry(
                            user_id,
                            username,
                            data,
                            group_chat_id=MAIN_GROUP_CHAT_ID,
                            group_message_id=getattr(group_message, "message_id", None),
                            thread_id=getattr(group_message, "message_thread_id", None),
                            archive_path=Path(archive_relative) if archive_relative else None,
                            processing_time=processing_time,
                            step_metrics=metrics_data,
                        )
                    except Exception as stats_error:
                        logger.error(f"Не удалось зафиксировать завершённое заключение: {stats_error}", exc_info=True)
                        completion_stats = None
                    await register_completion_meta(user_id, data)
                    if completion_stats:
                        await handle_new_achievements(
                            context.bot,
                            user_id,
                            username,
                            completion_stats,
                            data.region,
                        )
                        await send_personal_stats(context.bot, user_id)
                        completion_id = completion_stats.get("completion_id")
                        if completion_id and group_message and getattr(group_message, "message_id", None):
                            try:
                                await context.bot.edit_message_reply_markup(
                                    chat_id=MAIN_GROUP_CHAT_ID,
                                    message_id=group_message.message_id,
                                    reply_markup=build_void_keyboard(completion_id),
                                )
                            except TelegramError as edit_error:
                                logger.debug(
                                    f"Не удалось добавить кнопку удаления к сообщению {group_message.message_id}: {edit_error}"
                                )
            elif data and region:
                await safe_bot_send_message(
                    context.bot,
                    chat_id,
                    "Не найдена тема для региона. Документ не отправлен.",
                    message_thread_id=thread_id,
                )
            else:
                await safe_bot_send_message(
                    context.bot,
                    chat_id,
                    "Данные пользователя не найдены. Документ не отправлен.",
                    message_thread_id=thread_id,
                )
        else:
            await safe_bot_send_message(
                context.bot,
                chat_id,
                "Тестовое заключение отправлено только вам.",
                message_thread_id=thread_id,
            )

        try:
            filename_path.unlink()
        except OSError as cleanup_error:
            logger.warning(f"Не удалось удалить файл документа {filename_path}: {cleanup_error}")

    except FileNotFoundError as error:
        await safe_bot_send_message(
            context.bot,
            chat_id,
            f"Не удалось создать документ: {error}",
            message_thread_id=thread_id,
        )
    except Exception as error:
        logger.error(f"Критическая ошибка при создании/отправке документа: {error}", exc_info=True)
        await safe_bot_send_message(
            context.bot,
            chat_id,
            f"Произошла непредвиденная ошибка: {error}",
            message_thread_id=thread_id,
        )
    finally:
        if data_for_cleanup and data_for_cleanup.photo_desc:
            await asyncio.to_thread(cleanup_user_photos, data_for_cleanup)
        else:
            residual_data = await load_user_data_from_db(user_id)
            if residual_data.photo_desc:
                await asyncio.to_thread(cleanup_user_photos, residual_data)
        await delete_user_data_from_db(user_id)
        await clear_draft(user_id)
        metrics_reset(context)

    await safe_bot_send_message(
        context.bot,
        chat_id,
        f"{format_progress('mode')} Оформление завершено. Для нового заключения используйте /start.",
        message_thread_id=thread_id,
    )
    return ConversationHandler.END

async def test_choice_handler(update: Update, context: CallbackContext) -> int:
    if context.user_data.get("awaiting_draft_choice"):
        await remind_draft_decision(update, context)
        return DialogState.TESTING
    user_id = update.message.from_user.id
    if await is_user_blocked(user_id):
        await safe_reply(update, "Ваш доступ к боту ограничен. Обратитесь к администратору.")
        return ConversationHandler.END
    username = update.message.from_user.full_name
    choice_text = (update.message.text or "").strip().lower()

    mode_prompt = context.user_data.pop("mode_prompt", None)
    if mode_prompt:
        try:
            await context.bot.edit_message_reply_markup(
                chat_id=mode_prompt.get("chat_id"),
                message_id=mode_prompt.get("message_id"),
                reply_markup=None,
            )
        except (TypeError, TelegramError):
            pass

    if "тест" in choice_text:
        mode_key = "test"
    elif "оконч" in choice_text or "финал" in choice_text:
        mode_key = "final"
    else:
        await safe_reply(
            update,
            "Выберите режим кнопкой или введите 'тестовое' / 'окончательное'.",
            reply_markup=build_mode_keyboard()
        )
        return DialogState.TESTING

    is_final = mode_key == "final"
    mode_label = "Окончательное" if is_final else "Тестовое"
    await safe_reply(update, f"{format_progress('mode')} Режим выбран: {mode_label}.")
    await safe_reply(update, "Создаю документ. Это займёт несколько секунд.")

    metrics_complete_state(context, DialogState.TESTING)

    return await mode_flow(
        context,
        user_id,
        username,
        is_final,
        chat_id=update.effective_chat.id,
    )

async def cancel_handler(update: Update, context: CallbackContext) -> int:
    user_id = update.message.from_user.id
    data = await load_user_data_from_db(user_id)
    if data.photo_desc:
        await asyncio.to_thread(cleanup_user_photos, data)
    await delete_user_data_from_db(user_id)
    await clear_draft(user_id)
    context.user_data["awaiting_draft_choice"] = False
    await safe_reply(update, "Процесс остановлен.")
    await safe_reply(update, "Чтобы начать заново, введите /start.")
    return ConversationHandler.END


def build_reports_conversation_handler() -> ConversationHandler:
    """Create the reports conversation handler."""
    return ConversationHandler(
        entry_points=[CommandHandler("reports", reports_start_handler)],
        states={
            ReportState.ACTION: [
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                CallbackQueryHandler(report_action_callback, pattern=f"^{REPORT_ACTION_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, reports_action_handler),
            ],
            ReportState.MONTH_INPUT: [
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, reports_month_input_handler),
            ],
            ReportState.MONTH_REGION: [
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                CallbackQueryHandler(report_region_callback_handler, pattern=f"^{REPORT_REGION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, reports_month_region_handler),
            ],
            ReportState.PERIOD_START: [
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, reports_period_start_handler),
            ],
            ReportState.PERIOD_END: [
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, reports_period_end_handler),
            ],
            ReportState.PERIOD_REGION: [
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, reports_period_region_handler),
            ],
        },
        fallbacks=[
            CommandHandler("cancel", reports_cancel_handler),
            CommandHandler("menu", reports_cancel_handler),
            MessageHandler(filters.Regex("^❌ Отмена$"), reports_cancel_handler),
        ],
        allow_reentry=True,
    )


def build_conclusion_conversation_handler() -> ConversationHandler:
    """Create the primary conclusion conversation handler with enhanced back navigation."""
    return ConversationHandler(
        entry_points=[CommandHandler("start", start_handler)],
        states={
            DialogState.DEPARTMENT: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, get_department),
            ],
            DialogState.ISSUE_NUMBER: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, get_issue_number),
            ],
            DialogState.TICKET_NUMBER: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, get_ticket_number),
            ],
            DialogState.DATE: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, get_date),
            ],
            DialogState.REGION: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                CallbackQueryHandler(region_callback_handler, pattern=f"^{REGION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, get_region),
            ],
            DialogState.PHOTO: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler((filters.PHOTO | filters.Document.IMAGE | filters.VIDEO | filters.ANIMATION), photo_handler),
            ],
            DialogState.DESCRIPTION: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, description_handler),
            ],
            DialogState.EVALUATION: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, evaluation_handler),
            ],
            DialogState.MORE_PHOTO: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                CallbackQueryHandler(add_photo_callback_handler, pattern=f"^{ADD_PHOTO_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, more_photo_handler),
            ],
            DialogState.CONFIRMATION: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                CallbackQueryHandler(edit_callback_handler, pattern=f"^{EDIT_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, confirmation_handler),
                CallbackQueryHandler(confirmation_callback_handler, pattern=f"^{CONFIRM_CALLBACK_PREFIX}"),
            ],
            DialogState.TESTING: [
                CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"),
                CallbackQueryHandler(navigation_callback_handler, pattern=f"^{NAVIGATION_CALLBACK_PREFIX}"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, test_choice_handler),
                CallbackQueryHandler(mode_callback_handler, pattern=f"^{MODE_CALLBACK_PREFIX}"),
            ],
        },
        fallbacks=[CommandHandler("cancel", cancel_handler)],
        allow_reentry=True,
    )


def register_handlers(application: Application) -> None:
    """Register command, callback, and conversation handlers with the application."""
    application.add_handler(build_reports_conversation_handler())
    application.add_handler(build_conclusion_conversation_handler())

    deleted_filter = None
    for attr_name in ("DELETED_MESSAGES", "DELETED"):
        deleted_filter = getattr(filters.StatusUpdate, attr_name, None)
        if deleted_filter is not None:
            break

    if deleted_filter is not None:
        application.add_handler(MessageHandler(deleted_filter, deleted_message_handler))
    else:
        logger.info(
            "Поддержка уведомлений о удалённых сообщениях недоступна в текущей версии API. "
            "Используйте команду /void_ticket для ручного отката."
        )

    application.add_handler(CallbackQueryHandler(back_navigation_decision_handler, pattern=f"^{BACK_NAV_CALLBACK_PREFIX}"))
    application.add_handler(CallbackQueryHandler(void_callback_handler, pattern=f"^{VOID_CALLBACK_PREFIX}"))
    application.add_handler(CommandHandler("menu", menu_handler))
    application.add_handler(CommandHandler("help", help_handler))
    application.add_handler(CommandHandler("webapp", webapp_handler))
    application.add_handler(CommandHandler("help_admin", help_admin_handler))
    application.add_handler(CommandHandler("history", history_handler))
    application.add_handler(CommandHandler("stats", stats_handler))
    application.add_handler(CommandHandler("leaders", leaders_handler))
    application.add_handler(CommandHandler("achievements", achievements_handler))
    application.add_handler(CommandHandler("analytics", analytics_handler))
    application.add_handler(CommandHandler("drafts", drafts_handler))
    application.add_handler(CommandHandler("admin", admin_handler))
    application.add_handler(CommandHandler("backup", backup_handler))
    application.add_handler(CommandHandler("search_archive", search_archive_handler))
    application.add_handler(CommandHandler("download_month", download_month_handler))
    application.add_handler(CommandHandler("stats_period", stats_period_handler))
    application.add_handler(CommandHandler("void_ticket", void_ticket_handler))
    application.add_handler(CommandHandler("add_admin", add_admin_handler))
    application.add_handler(CallbackQueryHandler(history_callback_handler, pattern=f"^{HISTORY_CALLBACK_PREFIX}"))
    application.add_handler(CallbackQueryHandler(draft_callback_handler, pattern=f"^{DRAFT_CALLBACK_PREFIX}"))
    application.add_handler(CallbackQueryHandler(achievements_callback_handler, pattern=f"^{ACHIEVEMENTS_CALLBACK_PREFIX}"))
    application.add_handler(CallbackQueryHandler(analytics_callback_handler, pattern=f"^{ANALYTICS_CALLBACK_PREFIX}"))
    application.add_handler(CallbackQueryHandler(admin_callback_handler, pattern=f"^{ADMIN_CALLBACK_PREFIX}"))

async def main() -> None:
    # Создаём папки при старте
    TEMP_PHOTOS_DIR.mkdir(exist_ok=True)
    DOCS_DIR.mkdir(exist_ok=True)
    ARCHIVE_DIR.mkdir(exist_ok=True)
    ACHIEVEMENT_MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    if not ARCHIVE_INDEX_FILE.exists():
        ARCHIVE_INDEX_FILE.write_text("[]", encoding="utf-8")

    load_admin_ids()
    
    await init_db()

    try:
        bot_token = load_bot_token()
    except RuntimeError as token_error:
        logger.critical(token_error)
        raise

    application = Application.builder().token(bot_token).post_shutdown(close_db).build()

    # Добавляем периодическую задачу очистки временных файлов
    job_queue = application.job_queue
    job_queue.run_repeating(clean_temp_files_job, interval=3600, first=60)
    job_queue.run_repeating(network_recovery_job, interval=60, first=60)

    await application.bot.delete_webhook(drop_pending_updates=True)
    application.add_error_handler(error_handler)

    await configure_bot_commands(application.bot)

    register_handlers(application)

    logger.info("Бот запускается...")
    await application.run_polling()

if __name__ == "__main__":
    asyncio.run(main())
